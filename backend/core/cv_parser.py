"""CV Parser — Intelligent extraction from PDF, DOCX, and images.

Hybrid OCR approach:
- PyMuPDF for digital PDFs (FREE, instant)
- PaddleOCR for scanned/typed pages (FREE, local, no GPU)
- Vision LLM for handwritten pages ONLY (paid, high accuracy)
- Chandra-inspired page classification + layout preservation
- Quality scoring per page
"""
from __future__ import annotations

import os
import logging
import threading
from pathlib import Path

logger = logging.getLogger(__name__)

# Concurrency cap for PaddleOCR — singleton model is not thread-safe under
# heavy parallel load and each call holds ~500MB. Default 2 = comfortable
# on a 6GB container. Tune via OCR_CONCURRENCY env.
_OCR_GATE = threading.Semaphore(int(os.getenv("OCR_CONCURRENCY", "2")))

# Lazy-loaded PaddleOCR singleton
_paddle_ocr = None

# Target minimum dimension for OCR-bound images. Below this, upscale via
# LANCZOS (best quality bicubic-family resampler). 2000px ≈ 200 DPI on
# A4 — sweet spot for PaddleOCR/Tesseract accuracy.
_OCR_TARGET_MIN_PX = int(os.getenv("OCR_UPSCALE_MIN_PX", "2000"))
_OCR_TARGET_MAX_PX = int(os.getenv("OCR_UPSCALE_MAX_PX", "4000"))


def enhance_image_for_ocr(src_path: str, force: bool = False) -> str:
    """Pre-OCR enhancement pipeline. Returns path to enhanced image (may be
    a temp file) or the original path if no enhancement needed/possible.

    Steps (all free, CPU-only, ~50–200ms per image):
      1. EXIF auto-rotate
      2. Convert to RGB if needed
      3. Upscale to min 2000px (LANCZOS) if smaller
      4. Auto-deskew via projection-profile (skipped if PIL only)
      5. Sharpen (UnsharpMask)
      6. Auto-contrast + brightness normalize
      7. Convert to grayscale + Otsu binarize when very low contrast
         (printed text only — skipped for color photos / handwriting)

    Skipped entirely if image already > 2000px AND well-contrasted unless
    `force=True`. Failures fall back to original path silently.
    """
    try:
        from PIL import Image, ImageOps, ImageFilter, ImageEnhance, ImageStat
        import tempfile
        img = Image.open(src_path)
        img = ImageOps.exif_transpose(img)
        if img.mode not in ("RGB", "L"):
            img = img.convert("RGB")
        w, h = img.size
        min_dim = min(w, h)

        # Quick contrast stat — skip enhancement on already-good images
        gray_for_stat = img.convert("L") if img.mode != "L" else img
        contrast_std = ImageStat.Stat(gray_for_stat).stddev[0]
        if not force and min_dim >= _OCR_TARGET_MIN_PX and contrast_std >= 55:
            logger.info(f"OCR enhance: skip ({w}x{h}, std={contrast_std:.1f} — already good)")
            return src_path

        # 1. Upscale if needed
        if min_dim < _OCR_TARGET_MIN_PX:
            scale = _OCR_TARGET_MIN_PX / min_dim
            new_w = min(int(w * scale), _OCR_TARGET_MAX_PX)
            new_h = min(int(h * scale), _OCR_TARGET_MAX_PX)
            img = img.resize((new_w, new_h), Image.LANCZOS)
            logger.info(f"OCR enhance: upscaled {w}x{h} -> {new_w}x{new_h} (LANCZOS)")

        # 2. Auto-contrast (clip 1% tails, stretch histogram)
        img = ImageOps.autocontrast(img, cutoff=1)

        # 3. Sharpen — gentle unsharp mask, biased for text
        img = img.filter(ImageFilter.UnsharpMask(radius=1.5, percent=140, threshold=3))

        # 4. Brightness boost if image is dark
        gray = img.convert("L")
        avg_brightness = ImageStat.Stat(gray).mean[0]
        if avg_brightness < 110:
            factor = min(1.6, 130 / max(avg_brightness, 30))
            img = ImageEnhance.Brightness(img).enhance(factor)
            logger.info(f"OCR enhance: brightness boosted (avg={avg_brightness:.0f}, factor={factor:.2f})")

        # 5. Save enhanced copy as PNG (lossless) to temp
        out_fd, out_path = tempfile.mkstemp(suffix="_enh.png", prefix="pulse_ocr_")
        os.close(out_fd)
        img.save(out_path, "PNG", optimize=False)
        return out_path
    except Exception as e:
        logger.warning(f"OCR enhance failed ({e}); using original")
        return src_path


def _get_paddle_ocr(lang='en'):
    """Get or initialize PaddleOCR instance for specified language (lazy, singleton)."""
    global _paddle_ocr
    if os.getenv("DISABLE_PADDLE_OCR", "").lower() == "true":
        return None
    if _paddle_ocr is None or (hasattr(_paddle_ocr, '_lang') and _paddle_ocr._lang != lang):
        try:
            from paddleocr import PaddleOCR
            _paddle_ocr = PaddleOCR(use_angle_cls=True, lang=lang, show_log=False)
            _paddle_ocr._lang = lang
            logger.info(f"PaddleOCR initialized (lang={lang}, free local OCR)")
        except ImportError:
            logger.warning("PaddleOCR not installed — falling back to Vision LLM")
            _paddle_ocr = False  # mark as unavailable
    return _paddle_ocr if _paddle_ocr is not False else None


def preload_paddle_ocr():
    """Pre-warm PaddleOCR on startup. Call from main.py if PRELOAD_OCR=true."""
    if os.getenv("PRELOAD_OCR", "").lower() == "true":
        logger.info("Pre-loading PaddleOCR model...")
        ocr = _get_paddle_ocr()
        if ocr:
            logger.info("PaddleOCR ready")
        else:
            logger.warning("PaddleOCR not available")


def detect_language(text: str) -> str:
    """Simple language detection from text characters.

    Burmese note: Myanmar Unicode block is U+1000\u2013U+109F (covers Burmese,
    Shan, Mon, Karen, Kayah, S'gaw Karen, Eastern + Western Pwo Karen,
    Geba Karen, Pa'O Karen, Rumai Palaung, Khamti Shan). Both legitimate
    Unicode and Zawgyi-encoded text fall in this same block \u2014 disambiguation
    happens in `normalize_burmese()` via myanmar-tools heuristic.
    """
    if not text:
        return 'en'
    # Check character ranges
    for char in text[:500]:
        if '\u1000' <= char <= '\u109f':
            return 'my'      # Myanmar (Burmese, Shan, Mon, Karen, etc.)
        if '\uaa60' <= char <= '\uaa7f':
            return 'my'      # Myanmar Extended-A
        if '\ua9e0' <= char <= '\ua9ff':
            return 'my'      # Myanmar Extended-B
        if '\u0900' <= char <= '\u097f':
            return 'hi'      # Hindi
        if '\u0600' <= char <= '\u06ff':
            return 'ar'      # Arabic
        if '\u4e00' <= char <= '\u9fff':
            return 'ch'      # Chinese
        if '\u3040' <= char <= '\u309f' or '\u30a0' <= char <= '\u30ff':
            return 'japan'   # Japanese
        if '\uac00' <= char <= '\ud7af':
            return 'korean'  # Korean
    return 'en'


# Lazy Zawgyi\u2192Unicode converter singleton
_zawgyi_converter = None
_zawgyi_detector = None


def _get_zawgyi_tools():
    """Lazy-load myanmar-tools (Google's Zawgyi detector + ICU converter)."""
    global _zawgyi_converter, _zawgyi_detector
    if _zawgyi_converter is None:
        try:
            from myanmartools import ZawgyiDetector
            from icu import Transliterator
            _zawgyi_detector = ZawgyiDetector()
            _zawgyi_converter = Transliterator.createInstance("Zawgyi-my")
        except ImportError:
            logger.warning("myanmar-tools / PyICU not installed \u2014 Zawgyi text will not be normalized")
            _zawgyi_detector = False
            _zawgyi_converter = False
    return _zawgyi_detector, _zawgyi_converter


def normalize_burmese(text: str) -> str:
    """Normalize any Burmese-family text to Unicode.

    - Detects Zawgyi (legacy non-Unicode encoding common in Myanmar pre-2019)
      using Google's myanmar-tools probability model (>0.95 = Zawgyi).
    - Converts Zawgyi \u2192 Unicode via ICU's `Zawgyi-my` transliterator.
    - Pass-through for already-Unicode Burmese, Shan, Mon, Karen, etc.
    - No-op for non-Myanmar text (cheap early-return on script check).
    """
    if not text or detect_language(text) != 'my':
        return text
    detector, converter = _get_zawgyi_tools()
    if not detector or not converter:
        return text
    try:
        if detector.get_zawgyi_probability(text) > 0.95:
            return converter.transliterate(text)
    except Exception as e:
        logger.warning(f"Burmese normalization failed: {e}")
    return text


# ---------------------------------------------------------------------------
# File Classification
# ---------------------------------------------------------------------------
def classify_file(file_path: str) -> str:
    """Classify file type: pdf, docx, image, xlsx, txt, unknown."""
    ext = Path(file_path).suffix.lower()
    if ext == ".pdf":
        return "pdf"
    elif ext in (".docx", ".doc"):
        return "docx"
    elif ext in (".png", ".jpg", ".jpeg", ".tiff", ".bmp", ".webp"):
        return "image"
    elif ext in (".xlsx", ".xlsm"):
        return "xlsx"
    elif ext in (".txt", ".md", ".rtf"):
        return "txt"
    return "unknown"


def classify_page(page_text: str, page_image_path: str = None) -> str:
    """Classify a page as: digital, scanned_typed, handwritten, mixed.

    Learned from Chandra: classify BEFORE choosing extraction method.
    """
    char_count = len(page_text.strip()) if page_text else 0

    # Digital text: PyMuPDF extracted meaningful text
    if char_count > 200:
        return "digital"

    # Very little or no text extracted — it's scanned or handwritten
    if char_count < 50:
        # If we have an image, try to detect handwriting
        if page_image_path:
            return _detect_handwriting(page_image_path)
        return "scanned_typed"

    # Some text but not much — possibly mixed
    return "mixed"


def _peek_image_language(image_path: str) -> str:
    """Cheap Vision LLM probe — asks model to name the dominant script.

    Returns a `detect_language()`-compatible code (en/my/hi/ar/ch/japan/korean).
    Used to route Burmese/Myanmar images away from PaddleOCR (no model)
    BEFORE wasting an OCR pass on garbage output.
    Fails open to 'en' — caller still gets normal fallback chain.
    """
    from backend.core.config import is_ai_available
    if not is_ai_available():
        return 'en'
    import base64
    from backend.core.config import get_llm_client, LITE_MODEL
    try:
        with open(image_path, "rb") as f:
            img_data = base64.b64encode(f.read()).decode()
        ext = Path(image_path).suffix.lower().lstrip(".")
        mime = {"png": "image/png", "jpg": "image/jpeg", "jpeg": "image/jpeg",
                "tiff": "image/tiff", "bmp": "image/bmp", "webp": "image/webp"}.get(ext, "image/png")
        client = get_llm_client()
        resp = client.chat.completions.create(
            model=LITE_MODEL,
            messages=[{
                "role": "user",
                "content": [
                    {"type": "text", "text": "What is the dominant script in this document? Reply ONE word: ENGLISH, BURMESE, HINDI, ARABIC, CHINESE, JAPANESE, KOREAN, or OTHER."},
                    {"type": "image_url", "image_url": {"url": f"data:{mime};base64,{img_data}"}},
                ],
            }],
            max_tokens=10,
            temperature=0.0,
        )
        ans = resp.choices[0].message.content.strip().upper()
        m = {"BURMESE": "my", "MYANMAR": "my", "HINDI": "hi", "ARABIC": "ar",
             "CHINESE": "ch", "JAPANESE": "japan", "KOREAN": "korean", "ENGLISH": "en"}
        for k, v in m.items():
            if k in ans:
                return v
        return 'en'
    except Exception as e:
        logger.warning(f"Language peek failed: {e}")
        return 'en'


def _detect_handwriting(image_path: str) -> str:
    """Use Vision LLM to detect if an image contains handwriting.

    Returns: 'handwritten', 'scanned_typed', or 'mixed'
    """
    from backend.core.config import is_ai_available
    if not is_ai_available():
        return "scanned_typed"  # default when AI not available

    import base64
    from backend.core.config import get_llm_client, LITE_MODEL

    try:
        with open(image_path, "rb") as f:
            img_data = base64.b64encode(f.read()).decode()

        ext = Path(image_path).suffix.lower().lstrip(".")
        mime = {"png": "image/png", "jpg": "image/jpeg", "jpeg": "image/jpeg",
                "tiff": "image/tiff", "bmp": "image/bmp"}.get(ext, "image/png")

        client = get_llm_client()
        resp = client.chat.completions.create(
            model=LITE_MODEL,
            messages=[{
                "role": "user",
                "content": [
                    {"type": "text", "text": "Is this document handwritten, typed/printed, or a mix of both? Reply with ONLY one word: HANDWRITTEN, TYPED, or MIXED"},
                    {"type": "image_url", "image_url": {"url": f"data:{mime};base64,{img_data}"}},
                ],
            }],
            max_tokens=10,
            temperature=0.0,
        )
        answer = resp.choices[0].message.content.strip().upper()
        if "HANDWRITTEN" in answer:
            return "handwritten"
        elif "MIXED" in answer:
            return "mixed"
        return "scanned_typed"
    except Exception as e:
        logger.warning(f"Handwriting detection failed: {e}")
        return "scanned_typed"


# ---------------------------------------------------------------------------
# PDF Extraction
# ---------------------------------------------------------------------------
def extract_pdf(file_path: str) -> dict:
    """Extract text and metadata from PDF using PyMuPDF.

    Classifies each page and routes to the best extraction method.
    """
    import fitz

    doc = fitz.open(file_path)
    pages = []
    full_text = ""
    page_types = []

    for page_num in range(len(doc)):
        page = doc[page_num]
        text = page.get_text("text")
        page_type = classify_page(text)
        page_types.append(page_type)

        pages.append({
            "page": page_num + 1,
            "text": text,
            "char_count": len(text),
            "page_type": page_type,
        })
        full_text += text + "\n"

    # Overall document classification
    has_scanned = any(t in ("scanned_typed", "handwritten", "mixed") for t in page_types)
    has_handwritten = any(t == "handwritten" for t in page_types)
    is_scanned = all(p["char_count"] < 100 for p in pages) if pages else False

    doc.close()

    full_text = normalize_burmese(full_text)
    for p in pages:
        p["text"] = normalize_burmese(p["text"])
        p["char_count"] = len(p["text"])

    return {
        "text": full_text.strip(),
        "pages": pages,
        "page_count": len(pages),
        "is_scanned": is_scanned,
        "has_handwritten": has_handwritten,
        "page_types": page_types,
        "extraction_method": "pymupdf",
    }


# ---------------------------------------------------------------------------
# DOCX Extraction (layout-aware)
# ---------------------------------------------------------------------------
def extract_docx(file_path: str) -> dict:
    """Extract text from DOCX with layout awareness.

    Preserves: headings, sections, bullet points, tables.
    """
    from docx import Document

    doc = Document(file_path)
    sections = []
    full_text = ""
    current_section = ""

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        # Detect section headers by style
        style = para.style.name.lower() if para.style else ""
        is_heading = "heading" in style or para.runs and any(r.bold for r in para.runs)

        if is_heading and current_section:
            sections.append(current_section)
            current_section = ""

        if is_heading:
            full_text += f"\n## {text}\n"
            current_section = text
        elif any(text.startswith(b) for b in ("•", "-", "·", "▪", "●")):
            full_text += f"- {text.lstrip('•-·▪● ')}\n"
        else:
            full_text += text + "\n"

    if current_section:
        sections.append(current_section)

    # Extract tables with structure preserved (proper markdown: leading + trailing pipes)
    for table in doc.tables:
        full_text += "\n"
        for row_idx, row in enumerate(table.rows):
            cells = [cell.text.strip().replace("|", "\\|").replace("\n", " ") for cell in row.cells]
            row_text = "| " + " | ".join(cells) + " |"
            if row_idx == 0:
                full_text += row_text + "\n"
                full_text += "| " + " | ".join(["---"] * len(cells)) + " |\n"
            else:
                full_text += row_text + "\n"
        full_text += "\n"

    full_text = normalize_burmese(full_text)
    return {
        "text": full_text.strip(),
        "pages": [{"page": 1, "text": full_text.strip(), "char_count": len(full_text), "page_type": "digital"}],
        "page_count": 1,
        "is_scanned": False,
        "has_handwritten": False,
        "sections_detected": sections,
        "extraction_method": "python-docx",
    }


# ---------------------------------------------------------------------------
# Image/Scanned Extraction (Chandra-inspired approach)
# ---------------------------------------------------------------------------
def extract_image_ocr(file_path: str) -> dict:
    """Extract text from image using hybrid OCR.

    Routing:
    - Typed/scanned → PaddleOCR (FREE) → Tesseract fallback
    - Handwritten → Vision LLM (PAID, high quality) → PaddleOCR fallback
    """
    # Step -1: Enhance low-quality images BEFORE OCR (upscale + sharpen +
    # contrast + brightness). All downstream stages use enhanced path.
    enhanced_path = enhance_image_for_ocr(file_path)
    enhanced_is_temp = enhanced_path != file_path

    # Step 0: Quick Vision LLM peek for script detection on enhanced image
    lang_hint = _peek_image_language(enhanced_path)

    # Step 1: Try PaddleOCR first (free) — unless Burmese (no model)
    paddle_text = _paddle_ocr_extract(enhanced_path, lang_hint=lang_hint)

    try:
        if paddle_text and len(paddle_text) > 100:
            paddle_text = normalize_burmese(paddle_text)
            return {
                "text": paddle_text,
                "pages": [{"page": 1, "text": paddle_text, "char_count": len(paddle_text), "page_type": "scanned_typed"}],
                "page_count": 1,
                "is_scanned": True,
                "has_handwritten": False,
                "extraction_method": "paddleocr",
            }

        # Step 2: PaddleOCR got little text (or skipped for Burmese)
        page_type = _detect_handwriting(enhanced_path)

        if page_type == "handwritten":
            # Handwriting: use ORIGINAL (sharpening can degrade stroke nuance)
            text = _vision_extract_handwritten(file_path)
            method = "vision_llm_handwritten"
        else:
            # Printed/scanned typed: enhanced version helps Vision read crisper
            text = _vision_extract_structured(enhanced_path)
            method = "vision_llm_structured"

        text = normalize_burmese(text or "")
        if text and len(text) > 50:
            return {
                "text": text,
                "pages": [{"page": 1, "text": text, "char_count": len(text), "page_type": page_type}],
                "page_count": 1,
                "is_scanned": True,
                "has_handwritten": page_type == "handwritten",
                "extraction_method": method,
            }

        # Step 3: Last fallback — Tesseract (with Burmese trained data if available)
        text = _tesseract_extract(enhanced_path, lang_hint=lang_hint) or paddle_text or ""
        text = normalize_burmese(text)
        return {
            "text": text,
            "pages": [{"page": 1, "text": text, "char_count": len(text), "page_type": page_type}],
            "page_count": 1,
            "is_scanned": True,
            "has_handwritten": page_type == "handwritten",
            "extraction_method": "tesseract_fallback" if text else "failed",
        }
    finally:
        if enhanced_is_temp:
            try:
                os.remove(enhanced_path)
            except OSError:
                pass


def extract_scanned_pdf(file_path: str) -> dict:
    """Extract text from scanned/handwritten PDF pages.

    Improved: Per-page classification + specialized prompts.
    """
    import fitz

    doc = fitz.open(file_path)
    pages = []
    full_text = ""
    has_handwritten = False

    for page_num in range(len(doc)):
        page = doc[page_num]

        # Render page to image at 300 DPI
        pix = page.get_pixmap(dpi=300)
        img_path = f"/tmp/hire_page_{page_num}.png"
        pix.save(img_path)

        # Enhance the rendered page (no-op when already crisp at 300 DPI)
        enh_path = enhance_image_for_ocr(img_path)
        enh_is_temp = enh_path != img_path

        try:
            # Classify this page
            pymupdf_text = page.get_text("text")
            page_type = classify_page(pymupdf_text, enh_path)

            page_lang = detect_language(pymupdf_text) if pymupdf_text else _peek_image_language(enh_path)

            if page_type == "digital" and len(pymupdf_text.strip()) > 200:
                page_text = pymupdf_text
            elif page_type == "handwritten":
                has_handwritten = True
                # Handwriting → original render (preserve stroke nuance)
                page_text = _vision_extract_handwritten(img_path) or _paddle_ocr_extract(enh_path, lang_hint=page_lang) or _tesseract_extract(enh_path, lang_hint=page_lang) or ""
            else:
                page_text = _paddle_ocr_extract(enh_path, lang_hint=page_lang)
                if not page_text or len(page_text) < 50:
                    page_text = _vision_extract_structured(enh_path) or _tesseract_extract(enh_path, lang_hint=page_lang) or ""

            page_text = normalize_burmese(page_text)

            pages.append({
                "page": page_num + 1,
                "text": page_text,
                "char_count": len(page_text),
                "page_type": page_type,
            })
            full_text += page_text + "\n"
        finally:
            # Cleanup temp images
            try:
                os.remove(img_path)
            except OSError:
                pass
            if enh_is_temp:
                try:
                    os.remove(enh_path)
                except OSError:
                    pass

    doc.close()

    return {
        "text": full_text.strip(),
        "pages": pages,
        "page_count": len(pages),
        "is_scanned": True,
        "has_handwritten": has_handwritten,
        "extraction_method": "vision_llm_multipage",
    }


# ---------------------------------------------------------------------------
# Screenshots
# ---------------------------------------------------------------------------
def generate_screenshots(file_path: str, candidate_id: int, output_dir: str) -> list[dict]:
    """Generate 300 DPI page screenshots for a PDF, compressed to WebP."""
    import fitz
    from PIL import Image as PILImage

    screenshots = []
    ext = Path(file_path).suffix.lower()

    if ext != ".pdf":
        # Non-PDF (docx, doc, txt) — skip screenshot generation
        if ext in (".docx", ".doc", ".txt", ".md", ".rtf"):
            logger.info(f"Skipping screenshot generation for {ext} file")
            return []
        # Image file — copy + convert to webp
        from shutil import copy2
        dest = Path(output_dir) / f"cand_{candidate_id}_page_1.png"
        dest.parent.mkdir(parents=True, exist_ok=True)
        copy2(file_path, dest)
        try:
            img = PILImage.open(file_path)
            w, h = img.size
            img.close()
        except Exception as e:
            logger.warning(f"Could not open as image: {e}")
            return []
        webp_dest = dest.with_suffix('.webp')
        try:
            img = PILImage.open(str(dest))
            img.save(str(webp_dest), 'WEBP', quality=85)
            os.remove(str(dest))
            dest = webp_dest
        except Exception as comp_err:
            logger.warning(f"WebP compression failed, keeping PNG: {comp_err}")
        return [{"page": 1, "img_index": 0, "path": str(dest), "width": w, "height": h}]

    doc = fitz.open(file_path)
    for page_num in range(len(doc)):
        page = doc[page_num]
        pix = page.get_pixmap(dpi=300)
        dest = Path(output_dir) / f"cand_{candidate_id}_page_{page_num + 1}.png"
        dest.parent.mkdir(parents=True, exist_ok=True)
        pix.save(str(dest))
        w, h = pix.width, pix.height
        # Convert to WebP for smaller file size
        webp_dest = dest.with_suffix('.webp')
        try:
            img = PILImage.open(str(dest))
            img.save(str(webp_dest), 'WEBP', quality=85)
            os.remove(str(dest))  # remove original PNG
            dest = webp_dest
            logger.info(f"Compressed screenshot to WebP: {webp_dest.name}")
        except Exception as comp_err:
            logger.warning(f"WebP compression failed, keeping PNG: {comp_err}")
            # Keep PNG as fallback
        screenshots.append({
            "page": page_num + 1,
            "img_index": 0,
            "path": str(dest),
            "width": w,
            "height": h,
        })
    doc.close()
    return screenshots


# ---------------------------------------------------------------------------
# Vision LLM Prompts (Chandra-inspired specialized prompts)
# ---------------------------------------------------------------------------
def _vision_extract_structured(image_path: str) -> str | None:
    """Extract text from a typed/printed document image.

    Uses layout-aware prompt to preserve document structure.
    """
    import base64
    from backend.core.config import is_ai_available, get_llm_client, CHAT_MODEL

    if not is_ai_available():
        return None

    try:
        with open(image_path, "rb") as f:
            img_data = base64.b64encode(f.read()).decode()

        ext = Path(image_path).suffix.lower().lstrip(".")
        mime = {"png": "image/png", "jpg": "image/jpeg", "jpeg": "image/jpeg",
                "tiff": "image/tiff", "bmp": "image/bmp", "webp": "image/webp"}.get(ext, "image/png")

        client = get_llm_client()
        resp = client.chat.completions.create(
            model=CHAT_MODEL,
            messages=[{
                "role": "user",
                "content": [
                    {"type": "text", "text": """Extract ALL text from this CV/resume document with LAYOUT PRESERVATION.

OUTPUT FORMAT (use Markdown):
- Use ## for section headers (e.g., ## Experience, ## Education, ## Skills)
- Use bullet points (- ) for list items
- Preserve the hierarchical structure
- Keep dates, locations, and titles on their original lines
- For tables, use Markdown table format (| col1 | col2 |)
- Include ALL text — do not skip or summarize anything
- If there are contact details at the top, list them clearly
- LANGUAGE: preserve the ORIGINAL script. If text is in Burmese/Myanmar
  (Burmese, Shan, Mon, Karen, Kayah, Pa'O, etc.), output in Myanmar Unicode
  (U+1000–U+109F). DO NOT transliterate to Latin. DO NOT output Zawgyi —
  use proper Unicode codepoints (e.g., က ခ ဂ ဃ င, not Zawgyi visual lookalikes).

Return ONLY the extracted text in Markdown format, no commentary."""},
                    {"type": "image_url", "image_url": {"url": f"data:{mime};base64,{img_data}"}},
                ],
            }],
            max_tokens=4000,
            temperature=0.1,
        )
        return resp.choices[0].message.content
    except Exception as e:
        logger.warning(f"Vision LLM structured extraction failed: {e}")
        return None


def _vision_extract_handwritten(image_path: str) -> str | None:
    """Extract text from a HANDWRITTEN document image.

    Specialized prompt for handwriting — learned from Chandra's approach.
    """
    import base64
    from backend.core.config import is_ai_available, get_llm_client, CHAT_MODEL

    if not is_ai_available():
        return None

    try:
        with open(image_path, "rb") as f:
            img_data = base64.b64encode(f.read()).decode()

        ext = Path(image_path).suffix.lower().lstrip(".")
        mime = {"png": "image/png", "jpg": "image/jpeg", "jpeg": "image/jpeg",
                "tiff": "image/tiff", "bmp": "image/bmp", "webp": "image/webp"}.get(ext, "image/png")

        client = get_llm_client()
        resp = client.chat.completions.create(
            model=CHAT_MODEL,
            messages=[{
                "role": "user",
                "content": [
                    {"type": "text", "text": """This is a HANDWRITTEN CV/resume. Carefully read and transcribe ALL handwritten text.

INSTRUCTIONS FOR HANDWRITING:
1. Read each word carefully — handwriting may be cursive or print
2. If a word is unclear, make your best guess based on context (it's a CV)
3. Common CV sections to look for: Name, Contact, Education, Experience, Skills, References
4. Preserve the document structure using Markdown:
   - ## for section headers
   - Bullet points for list items
   - Keep dates and locations
5. For checkboxes: ☑ for checked, ☐ for unchecked
6. For crossed-out text: ~~crossed out~~
7. Include ALL text — even marginal notes or annotations
8. LANGUAGE: preserve the ORIGINAL script. Burmese/Myanmar handwriting
   (Burmese, Shan, Mon, Karen, Kayah, Pa'O, etc.) → output Myanmar Unicode
   (U+1000–U+109F). Never transliterate to Latin. Never output Zawgyi.

Return the transcribed text in clean Markdown format. Do NOT add any commentary or notes about legibility."""},
                    {"type": "image_url", "image_url": {"url": f"data:{mime};base64,{img_data}"}},
                ],
            }],
            max_tokens=4000,
            temperature=0.1,
        )
        return resp.choices[0].message.content
    except Exception as e:
        logger.warning(f"Vision LLM handwriting extraction failed: {e}")
        return None


def verify_critical_fields(image_path: str) -> dict:
    """Run a high-fidelity vision model on critical fields only.

    Returns dict with possibly-empty: name, dob, phone, national_id, email,
    gender, marital_status, nationality, religion, height, weight, father_name.
    Used as second-pass verification on handwritten docs where critical
    fields (digits, exact spelling) matter more than full extraction.

    Configurable via env: VISION_VERIFIER_MODEL (default opus-4.7).
    Returns {} on error or when AI unavailable.
    """
    import base64
    import json as _json
    import os
    from backend.core.config import is_ai_available, get_llm_client

    if not is_ai_available():
        return {}

    verifier_model = os.getenv("VISION_VERIFIER_MODEL", "anthropic/claude-opus-4.7")

    try:
        with open(image_path, "rb") as f:
            img_data = base64.b64encode(f.read()).decode()
        ext = Path(image_path).suffix.lower().lstrip(".")
        mime = {"png": "image/png", "jpg": "image/jpeg", "jpeg": "image/jpeg",
                "tiff": "image/tiff", "bmp": "image/bmp", "webp": "image/webp"}.get(ext, "image/png")

        client = get_llm_client()
        resp = client.chat.completions.create(
            model=verifier_model,
            messages=[{
                "role": "user",
                "content": [
                    {"type": "text", "text": """Re-extract ONLY these fields from this image with maximum care for digits and spelling:
- name (exact spelling, every letter)
- dob (DD.MM.YYYY or whatever format is shown, else null)
- phone (all digits, all numbers, comma-separated if multiple)
- national_id (exact characters incl. last digit; for Myanmar NRC, e.g. "12/Tha Ga Ka (N) 189824")
- email (if visible)
- gender (Male / Female / Other, else null)
- marital_status (Single / Married / Divorced / Widowed, else null)
- nationality (e.g. Myanmar, Indian, etc., else null)
- religion (e.g. Buddhist, Christian, Muslim, Hindu, etc., else null)
- height (e.g. "5'8\"" or "175cm", else null)
- weight (e.g. "130 lb" or "60kg", else null)
- father_name (if a "father" / "father's name" field is on the form, else null)

Return STRICT JSON only with EXACTLY these keys:
{"name": "...", "dob": "...", "phone": "...", "national_id": "...", "email": "...", "gender": "...", "marital_status": "...", "nationality": "...", "religion": "...", "height": "...", "weight": "...", "father_name": "..."}
Use null for any field not visible. NO commentary."""},
                    {"type": "image_url", "image_url": {"url": f"data:{mime};base64,{img_data}"}},
                ],
            }],
            max_tokens=600,
            temperature=0.0,
        )
        text = resp.choices[0].message.content.strip()
        if text.startswith("```"):
            text = text.split("```", 2)[1]
            if text.startswith("json"):
                text = text[4:]
            text = text.rsplit("```", 1)[0]
        try:
            data = _json.loads(text)
        except Exception:
            i, j = text.find("{"), text.rfind("}")
            if i >= 0 and j > i:
                data = _json.loads(text[i:j+1])
            else:
                return {}
        # Strip nulls and empty strings
        return {k: v for k, v in data.items() if v and str(v).strip().lower() not in ("null", "none", "")}
    except Exception as e:
        logger.warning(f"Critical-field verifier ({verifier_model}) failed: {e}")
        return {}


def _html_table_to_markdown(html: str) -> str:
    """Convert simple HTML table to Markdown."""
    import re
    rows = re.findall(r'<tr>(.*?)</tr>', html, re.DOTALL)
    if not rows:
        return ""
    md_rows = []
    for i, row in enumerate(rows):
        cells = re.findall(r'<t[dh]>(.*?)</t[dh]>', row, re.DOTALL)
        cells = [c.strip() for c in cells]
        md_rows.append("| " + " | ".join(cells) + " |")
        if i == 0:
            md_rows.append("| " + " | ".join(["---"] * len(cells)) + " |")
    return "\n".join(md_rows)


def _paddle_ocr_extract(image_path: str, lang_hint: str | None = None) -> str | None:
    """Extract text using PaddleOCR (FREE, local, no GPU needed).

    Best for: typed/printed/scanned documents.
    Not great for: handwriting (use Vision LLM instead).
    Skipped for Burmese/Myanmar — PaddleOCR has no Myanmar model;
    caller falls back to Vision LLM which reads Burmese well.
    """
    if lang_hint == 'my':
        logger.info("PaddleOCR skipped: lang=my unsupported, deferring to Vision LLM")
        return None
    ocr = _get_paddle_ocr()
    if not ocr:
        return None

    # Cap parallel OCR jobs (RAM + thread-safety on shared model)
    _OCR_GATE.acquire()
    try:
        result = ocr.ocr(image_path, cls=True)
        if not result or not result[0]:
            return None

        # Sort by vertical position (top to bottom), then horizontal (left to right)
        lines = []
        for line in result[0]:
            box = line[0]  # [[x1,y1],[x2,y2],[x3,y3],[x4,y4]]
            text = line[1][0]  # text string
            confidence = line[1][1]  # confidence score
            y_pos = (box[0][1] + box[2][1]) / 2  # average y position
            x_pos = box[0][0]  # left x position
            lines.append({"text": text, "confidence": confidence, "y": y_pos, "x": x_pos})

        # Group by lines (similar y position = same line)
        lines.sort(key=lambda l: (l["y"], l["x"]))
        grouped = []
        current_line = []
        current_y = -1

        for item in lines:
            if current_y == -1 or abs(item["y"] - current_y) < 15:  # 15px threshold
                current_line.append(item)
                current_y = item["y"] if current_y == -1 else (current_y + item["y"]) / 2
            else:
                if current_line:
                    current_line.sort(key=lambda l: l["x"])
                    grouped.append(" ".join(l["text"] for l in current_line))
                current_line = [item]
                current_y = item["y"]

        if current_line:
            current_line.sort(key=lambda l: l["x"])
            grouped.append(" ".join(l["text"] for l in current_line))

        text = "\n".join(grouped)
        avg_confidence = sum(l["confidence"] for l in lines) / len(lines) if lines else 0
        logger.info(f"PaddleOCR: extracted {len(lines)} text blocks, avg confidence: {avg_confidence:.2%}")

        if avg_confidence < 0.7:
            logger.warning(f"PaddleOCR low confidence: {avg_confidence:.2%} — may need Vision LLM")

        # Try table detection with ppstructure (if available)
        try:
            from paddleocr import PPStructure
            table_engine = PPStructure(show_log=False, table=True)
            table_result = table_engine(image_path)
            for item in table_result:
                if item.get('type') == 'table':
                    # Convert table HTML to markdown
                    html = item.get('res', {}).get('html', '')
                    if html:
                        text += "\n" + _html_table_to_markdown(html)
        except ImportError:
            pass  # ppstructure not available
        except Exception as e:
            logger.warning(f"Table extraction failed: {e}")

        return text if text.strip() else None

    except Exception as e:
        logger.warning(f"PaddleOCR extraction failed: {e}")
        return None
    finally:
        _OCR_GATE.release()


_TESS_LANG_MAP = {
    'my': 'mya', 'hi': 'hin', 'ar': 'ara', 'ch': 'chi_sim',
    'japan': 'jpn', 'korean': 'kor', 'en': 'eng',
}


def _tesseract_extract(image_path: str, lang_hint: str | None = None) -> str | None:
    """Fallback OCR using Tesseract.

    Tesseract supports Burmese via the `mya` traineddata pack
    (install: `apt-get install tesseract-ocr-mya`). Always tries
    the hinted language first, then falls back to English.
    """
    try:
        import subprocess
        tess_lang = _TESS_LANG_MAP.get(lang_hint or 'en', 'eng')
        if tess_lang != 'eng':
            tess_lang = f"{tess_lang}+eng"  # multi-language mode
        result = subprocess.run(
            ["tesseract", image_path, "-", "-l", tess_lang, "--oem", "3", "--psm", "6"],
            capture_output=True, text=True, timeout=30,
        )
        text = result.stdout.strip()
        return text if text else None
    except (FileNotFoundError, subprocess.TimeoutExpired) as e:
        logger.warning(f"Tesseract not available or timed out: {e}")
        return None


# ---------------------------------------------------------------------------
# Extraction Quality Scoring
# ---------------------------------------------------------------------------
def score_extraction_quality(text: str) -> dict:
    """Score the quality of extracted text.

    Learned from Chandra: measure extraction confidence.
    """
    if not text:
        return {"score": 0, "issues": ["No text extracted"]}

    issues = []
    score = 100
    words = text.split()
    word_count = len(words)

    # Too short
    if word_count < 30:
        score -= 40
        issues.append(f"Very short ({word_count} words)")
    elif word_count < 100:
        score -= 20
        issues.append(f"Short ({word_count} words)")

    # Check for garbled text (high ratio of non-alpha characters)
    alpha_count = sum(1 for c in text if c.isalpha())
    total_count = len(text.replace(" ", "").replace("\n", ""))
    if total_count > 0:
        alpha_ratio = alpha_count / total_count
        if alpha_ratio < 0.5:
            score -= 30
            issues.append(f"Possibly garbled (alpha ratio: {alpha_ratio:.1%})")

    # Check for common CV sections
    text_lower = text.lower()
    cv_sections = ["experience", "education", "skills", "work", "qualifications", "summary", "contact"]
    found_sections = sum(1 for s in cv_sections if s in text_lower)
    if found_sections == 0:
        score -= 20
        issues.append("No CV sections detected")
    elif found_sections >= 3:
        score += 10  # bonus for good structure

    # Check for email/phone (signs of good extraction)
    import re
    has_email = bool(re.search(r'[\w.+-]+@[\w-]+\.[\w.-]+', text))
    has_phone = bool(re.search(r'[\+\d][\d\s\-\(\)]{7,}', text))
    if has_email:
        score += 5
    if has_phone:
        score += 5

    return {
        "score": max(0, min(100, score)),
        "word_count": word_count,
        "has_email": has_email,
        "has_phone": has_phone,
        "sections_found": found_sections,
        "issues": issues,
    }
