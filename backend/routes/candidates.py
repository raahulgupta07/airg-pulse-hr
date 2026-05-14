"""Candidate routes — CV upload, list, search, CRUD, LinkedIn/GitHub/text import, smart search."""
from __future__ import annotations

import re
import os
import csv as _csv
import io as _io
import json
import uuid
import hmac
import hashlib
import secrets as _secrets_mod
import time as _time_mod
import logging
from datetime import datetime, timezone
from pathlib import Path
from typing import Optional

import httpx
from fastapi import APIRouter, Request, UploadFile, File, Query, HTTPException, Form
from fastapi.responses import JSONResponse, Response
from pydantic import BaseModel

from backend.core.auth import get_current_user, validate_token
from backend.core.config import CHAT_MODEL, LITE_MODEL, llm_call, EMBEDDING_MODELS
from backend.core.database import (
    list_candidates, get_candidate, count_candidates, insert_candidate, get_cursor,
)
from backend.core.ids import resolve_id_sync
from backend.core.permissions import (
    has_min_role,
    require_role,
    ensure_owner_or_min,
)
from backend.core.rate_limit import limiter
from fastapi import Depends


def _resolve_cid(candidate_id) -> Optional[int]:
    """Resolve int|str (digits or public_id like cv_...) -> internal int id, or None."""
    return resolve_id_sync("candidates", candidate_id)

EMBED_MODEL_FOR_ARTIFACTS = EMBEDDING_MODELS[0] if EMBEDDING_MODELS else "google/gemini-embedding-001"

logger = logging.getLogger(__name__)
router = APIRouter()

# Track active pipeline tasks for cancellation. Keyed by candidate_id.
import asyncio as _asyncio_mod
_active_tasks: "dict[int, _asyncio_mod.Task]" = {}

# Queue + worker for bulk pipeline runs. Caps concurrency to MAX_PARALLEL.
# Queue entry: {"cid": int, "pdf_path": str, "run_id": str, "queued_at": float}
_pipeline_queue: "list[dict]" = []
_queue_running: "set[int]" = set()
MAX_PARALLEL_PIPELINES = int(os.getenv("MAX_PARALLEL_PIPELINES", "8"))
_queue_workers_started = False

def _track(cid: int, task):
    _active_tasks[cid] = task
    def _done(_):
        if _active_tasks.get(cid) is task:
            _active_tasks.pop(cid, None)
        _queue_running.discard(cid)
    task.add_done_callback(_done)
    return task


async def _queue_worker():
    """Pulls jobs from _pipeline_queue, runs at most MAX_PARALLEL_PIPELINES concurrent pipelines."""
    while True:
        if not _pipeline_queue or len(_queue_running) >= MAX_PARALLEL_PIPELINES:
            await _asyncio_mod.sleep(0.2)
            continue
        job = _pipeline_queue.pop(0)
        cid = job["cid"]
        if cid in _queue_running:
            continue
        _queue_running.add(cid)
        async def _run(j=job):
            try:
                from backend.core.cv_pipeline import process_cv
                await process_cv(j["cid"], j["pdf_path"], run_id=j["run_id"])
            except Exception as e:
                logger.error(f"Queue pipeline failed cid={j['cid']}: {e}")
        _track(cid, _asyncio_mod.create_task(_run()))


def _ensure_queue_workers():
    global _queue_workers_started
    if _queue_workers_started:
        return
    try:
        for _ in range(MAX_PARALLEL_PIPELINES):
            _asyncio_mod.create_task(_queue_worker())
        _queue_workers_started = True
    except RuntimeError:
        pass

_PROJECT_ROOT = Path(__file__).parent.parent.parent
DATA_DIR = Path(os.getenv("DATA_DIR", str(_PROJECT_ROOT / "data")))
CV_DIR = DATA_DIR / "cvs"
CV_DIR.mkdir(parents=True, exist_ok=True)

# ---------------------------------------------------------------------------
# File signing (HMAC short-lived URLs for /file endpoint)
# ---------------------------------------------------------------------------
FILE_SIGN_TTL = 5 * 60  # 5 minutes


def _file_sign_secret() -> bytes:
    secret = os.getenv("FILE_SIGN_SECRET") or os.getenv("JWT_SECRET")
    if not secret:
        # generate-and-warn (process-stable for current run)
        global _GENERATED_FILE_SIGN_SECRET
        try:
            _GENERATED_FILE_SIGN_SECRET
        except NameError:
            _GENERATED_FILE_SIGN_SECRET = _secrets_mod.token_hex(32)
            logger.warning(
                "FILE_SIGN_SECRET / JWT_SECRET not set — generated ephemeral secret. "
                "Signed URLs will not survive process restart. Set FILE_SIGN_SECRET in env."
            )
        secret = _GENERATED_FILE_SIGN_SECRET
    return secret.encode() if isinstance(secret, str) else secret


def _compute_file_sig(candidate_id: int, exp: int) -> str:
    msg = f"{candidate_id}:{exp}".encode()
    return hmac.new(_file_sign_secret(), msg, hashlib.sha256).hexdigest()


def _cand_scope_filter(scope: str, user: dict):
    uid = user.get("user_id") or 1
    sector_id = user.get("sector_id")
    if scope == "mine":
        return "owner_id = %s", [uid]
    if scope == "sector":
        if sector_id is None:
            return "(visibility = 'sector' OR shared_sector = TRUE)", []
        return "(visibility = 'sector' OR shared_sector = TRUE) AND sector_id = %s", [sector_id]
    if scope == "pool":
        return "(visibility = 'global' OR shared_global = TRUE)", []
    if sector_id is None:
        return "(owner_id = %s OR visibility = 'global' OR shared_global = TRUE)", [uid]
    return ("(owner_id = %s "
            " OR ((visibility='sector' OR shared_sector=TRUE) AND sector_id=%s) "
            " OR visibility='global' OR shared_global=TRUE)"), [uid, sector_id]


def _user_with_sector(user: dict) -> dict:
    if user.get("sector_id") is not None:
        return user
    from backend.core.database import get_cursor as _gc
    with _gc() as cur:
        cur.execute("SELECT sector_id FROM users WHERE id = %s", (user.get("user_id") or 1,))
        row = cur.fetchone()
        if row:
            user["sector_id"] = row[0]
    return user


SORT_COLS_CV = {
    "created_at": "c.created_at",
    "updated_at": "c.updated_at",
    "name": "c.name",
    "experience": "c.total_experience_years",
    "quality": "c.quality_score",
}


@router.get("/")
async def list_candidates_route(
    request: Request,
    search: Optional[str] = Query(None),
    skills: Optional[str] = Query(None),
    companies: Optional[str] = Query(None, description="csv of company canonicals"),
    locations: Optional[str] = Query(None, description="csv of location canonicals"),
    languages: Optional[str] = Query(None, description="csv of language canonicals"),
    certs: Optional[str] = Query(None, description="csv of certification canonicals"),
    education: Optional[str] = Query(None, description="csv of education canonicals"),
    seniority: Optional[str] = Query(None),
    scope: Optional[str] = Query(None, description="mine | sector | pool"),
    created_by: Optional[int] = Query(None),
    modified_by: Optional[int] = Query(None),
    created_after: Optional[str] = Query(None),
    created_before: Optional[str] = Query(None),
    modified_after: Optional[str] = Query(None),
    modified_before: Optional[str] = Query(None),
    sort: str = Query("created_at"),
    dir: str = Query("desc"),
    with_assignments: bool = Query(False),
    force_mask: bool = Query(False),
    limit: int = Query(50, le=200),
    offset: int = Query(0),
):
    """List candidates with scope, filters, sort."""
    user = _user_with_sector(get_current_user(request))

    if scope:
        from backend.core.database import get_cursor as _gc
        scope_sql, scope_params = _cand_scope_filter(scope, user)
        conds = ["c.status = 'active'"]
        params = list(scope_params)
        if scope_sql:
            conds.append(scope_sql.replace("owner_id", "c.owner_id")
                                  .replace("visibility", "c.visibility")
                                  .replace("sector_id", "c.sector_id"))
        if search:
            conds.append("(c.name ILIKE %s OR c.email ILIKE %s OR c.current_role ILIKE %s)")
            params.extend([f"%{search}%", f"%{search}%", f"%{search}%"])
        if seniority:
            conds.append("c.seniority_level = %s")
            params.append(seniority)
        # AI-mined facet filters (CSV → array overlap / ILIKE)
        if skills:
            sk = [s.strip() for s in skills.split(",") if s.strip()]
            if sk:
                conds.append("(c.skills_technical && %s OR c.skills_soft && %s OR c.tools && %s)")
                params.extend([sk, sk, sk])
        if companies:
            cs = [c.strip() for c in companies.split(",") if c.strip()]
            if cs:
                # match current_company OR any experience[].company (case-insensitive substring)
                ors = []
                for c0 in cs:
                    ors.append("(c.current_company ILIKE %s OR c.experience::text ILIKE %s)")
                    params.extend([f"%{c0}%", f"%\"company\": \"{c0}%"])
                conds.append("(" + " OR ".join(ors) + ")")
        if locations:
            ls = [s.strip() for s in locations.split(",") if s.strip()]
            if ls:
                ors = []
                for l0 in ls:
                    ors.append("c.location ILIKE %s")
                    params.append(f"%{l0}%")
                conds.append("(" + " OR ".join(ors) + ")")
        if languages:
            lg = [s.strip() for s in languages.split(",") if s.strip()]
            if lg:
                conds.append("c.languages && %s")
                params.append(lg)
        if certs:
            cr = [s.strip() for s in certs.split(",") if s.strip()]
            if cr:
                ors = []
                for c0 in cr:
                    ors.append("c.certifications::text ILIKE %s")
                    params.append(f"%{c0}%")
                conds.append("(" + " OR ".join(ors) + ")")
        if education:
            ed = [s.strip() for s in education.split(",") if s.strip()]
            if ed:
                ors = []
                for e0 in ed:
                    ors.append("c.education::text ILIKE %s")
                    params.append(f"%{e0}%")
                conds.append("(" + " OR ".join(ors) + ")")
        if created_by:    conds.append("c.owner_id = %s");    params.append(created_by)
        if modified_by:   conds.append("c.updated_by = %s");  params.append(modified_by)
        if created_after: conds.append("c.created_at >= %s::timestamptz"); params.append(created_after)
        if created_before:conds.append("c.created_at <  %s::timestamptz"); params.append(created_before)
        if modified_after:conds.append("c.updated_at >= %s::timestamptz"); params.append(modified_after)
        if modified_before:conds.append("c.updated_at < %s::timestamptz"); params.append(modified_before)

        where = ("WHERE " + " AND ".join(conds)) if conds else ""
        sort_col = SORT_COLS_CV.get(sort, "c.created_at")
        sort_dir = "ASC" if dir.lower() == "asc" else "DESC"

        with _gc() as cur:
            cur.execute(f"SELECT COUNT(*) FROM candidates c {where}", params)
            total = cur.fetchone()[0]
            cur.execute(
                f"""SELECT c.id, c.public_id, c.name, c.email, c.current_role, c.current_company,
                           c.total_experience_years, c.seniority_level, c.skills_technical, c.tags,
                           c.quality_score, c.source, c.visibility, c.sector_id, c.owner_id,
                           c.is_processed, c.processing_error,
                           c.created_at, c.updated_at, c.updated_by,
                           uo.display_name AS owner_name,
                           um.display_name AS updated_by_name,
                           c.created_by_id, c.updated_by_id, c.expires_at,
                           uc.display_name AS created_by_name,
                           umi.display_name AS updated_by_id_name,
                           (c.expires_at IS NOT NULL AND c.expires_at < NOW()) AS is_expired,
                           CASE WHEN c.expires_at IS NULL THEN NULL
                                ELSE EXTRACT(DAY FROM (c.expires_at - NOW()))::int
                           END AS days_until_expiry
                    FROM candidates c
                    LEFT JOIN users uo  ON uo.id  = c.owner_id
                    LEFT JOIN users um  ON um.id  = c.updated_by
                    LEFT JOIN users uc  ON uc.id  = c.created_by_id
                    LEFT JOIN users umi ON umi.id = c.updated_by_id
                    {where}
                    ORDER BY {sort_col} {sort_dir} NULLS LAST
                    LIMIT %s OFFSET %s""",
                (*params, limit, offset),
            )
            cols = [desc[0] for desc in cur.description]
            candidates = [dict(zip(cols, r)) for r in cur.fetchall()]
            # Backfill created_by_name fallback to owner_name + updated_by_name preference
            for c in candidates:
                if not c.get("created_by_name"):
                    c["created_by_name"] = c.get("owner_name")
                if not c.get("updated_by_name"):
                    c["updated_by_name"] = c.get("updated_by_id_name")
            # Optional: aggregate active assignments per candidate
            if with_assignments and candidates:
                cand_ids = [c["id"] for c in candidates]
                cur.execute("""
                    SELECT pc.candidate_id, pc.position_id, p.slug, p.title, pc.stage,
                           pc.match_score_composite,
                           pc.created_at AS added_at
                    FROM position_candidates pc
                    JOIN positions p ON p.id = pc.position_id
                    WHERE pc.candidate_id = ANY(%s)
                      AND COALESCE(pc.dismissed, FALSE) = FALSE
                    ORDER BY pc.created_at DESC
                """, (cand_ids,))
                cols2 = [d[0] for d in cur.description]
                rows = [dict(zip(cols2, r)) for r in cur.fetchall()]
                by_cand: dict = {}
                for r in rows:
                    if r.get("added_at"):
                        r["added_at"] = r["added_at"].isoformat()
                    by_cand.setdefault(r["candidate_id"], []).append(r)
                for c in candidates:
                    c["assignments"] = by_cand.get(c["id"], [])
            # Dedup candidates by (name, email) — collapse duplicate uploads;
            # keep highest quality_score, merge assignments
            if candidates:
                deduped: dict = {}
                for c in candidates:
                    key = ((c.get("name") or "").strip().lower(),
                           (c.get("email") or "").strip().lower())
                    if not key[0] and not key[1]:
                        deduped[("__noid__", c["id"])] = c
                        continue
                    existing = deduped.get(key)
                    if existing is None:
                        c["duplicate_candidate_ids"] = []
                        deduped[key] = c
                        continue
                    cur_q = existing.get("quality_score") or 0
                    new_q = c.get("quality_score") or 0
                    winner, loser = (existing, c) if cur_q >= new_q else (c, existing)
                    a_existing = winner.get("assignments") or []
                    a_loser = loser.get("assignments") or []
                    merged = a_existing + a_loser
                    seen = set()
                    out_assign = []
                    for a in merged:
                        ak = (a.get("position_id"), a.get("stage"))
                        if ak in seen: continue
                        seen.add(ak); out_assign.append(a)
                    winner["assignments"] = out_assign
                    winner.setdefault("duplicate_candidate_ids", [])
                    winner["duplicate_candidate_ids"] = list(set(
                        (existing.get("duplicate_candidate_ids") or []) +
                        (c.get("duplicate_candidate_ids") or []) +
                        [loser["id"]]
                    ))
                    deduped[key] = winner
                candidates = list(deduped.values())
            # PII mask for pool view
            pool_masked = False
            if scope == "pool":
                uid_self = user.get("user_id")
                role = user.get("role")
                is_privileged = role in ("admin", "group_hr")
                for c in candidates:
                    should_mask = force_mask or (not is_privileged) or (c.get("owner_id") != uid_self)
                    if should_mask:
                        pool_masked = True
                        c["name"] = f"Anonymous#{c['id']}"
                        c["email"] = None
                        c["phone"] = None
                        c["linkedin_url"] = None
                        c["owner_name"] = None
                        c["updated_by_name"] = None
                        loc = c.get("location")
                        if loc:
                            # Strip city — keep last comma-separated segment as country
                            parts = [p.strip() for p in str(loc).split(",") if p.strip()]
                            c["location"] = parts[-1] if parts else None
                # Audit-log pool access
                if candidates:
                    try:
                        cur.execute(
                            "INSERT INTO audit_log (user_id, action, resource_type, details) VALUES (%s,%s,%s,%s)",
                            (uid_self, "POOL_VIEW", "candidates",
                             json.dumps({"count": len(candidates), "force_mask": force_mask})),
                        )
                    except Exception:
                        pass
            # Counts
            uid = user.get("user_id") or 1
            sector_id = user.get("sector_id")
            cur.execute("SELECT COUNT(*) FROM candidates WHERE owner_id = %s", (uid,))
            n_mine = cur.fetchone()[0]
            if sector_id is not None:
                cur.execute("SELECT COUNT(*) FROM candidates "
                            "WHERE (visibility='sector' OR shared_sector=TRUE) AND sector_id = %s", (sector_id,))
            else:
                cur.execute("SELECT COUNT(*) FROM candidates "
                            "WHERE visibility='sector' OR shared_sector=TRUE")
            n_sector = cur.fetchone()[0]
            cur.execute("SELECT COUNT(*) FROM candidates "
                        "WHERE visibility='global' OR shared_global=TRUE")
            n_pool = cur.fetchone()[0]
        return {
            "candidates": candidates,
            "total": total,
            "limit": limit,
            "offset": offset,
            "pool_masked": pool_masked,
            "counts": {"mine": n_mine, "sector": n_sector, "pool": n_pool},
        }

    # Default path (legacy, no scope)
    skills_list = skills.split(",") if skills else None
    candidates = list_candidates(
        search=search,
        skills=skills_list,
        seniority=seniority,
        limit=limit,
        offset=offset,
    )
    total = count_candidates()
    return {"candidates": candidates, "total": total, "limit": limit, "offset": offset}


# ---------------------------------------------------------------------------
# Smart NLP Search
# ---------------------------------------------------------------------------
@router.post("/smart-search")
async def smart_search(request: Request):
    """Natural language search: parse query into structured filters and search candidates."""
    get_current_user(request)
    body = await request.json()

    query = body.get("query", "").strip()
    if not query:
        raise HTTPException(status_code=400, detail="query is required")

    # Use LLM to parse the natural language query into structured filters
    parse_prompt = f"""Parse this natural language hiring search query into structured filters.

Query: "{query}"

Return ONLY valid JSON:
{{
  "skills": ["skill1", "skill2"],
  "min_experience_years": null or number,
  "max_experience_years": null or number,
  "companies": ["company1", "company2"],
  "seniority": null or "junior|mid|senior|staff|principal|lead|manager|director",
  "location": null or "string",
  "search_text": "any remaining free-text search terms",
  "interpretation": "Brief human-readable interpretation of the query"
}}

Examples:
- "Python developer with 5+ years" -> {{"skills": ["python"], "min_experience_years": 5, "seniority": null, "search_text": "developer"}}
- "Senior data scientist in NYC" -> {{"skills": ["data science"], "seniority": "senior", "location": "NYC"}}
- "Find backend engineers who worked at Google or Meta" -> {{"skills": ["backend"], "companies": ["google", "meta"], "search_text": "engineer"}}"""

    parse_result = llm_call(parse_prompt, model=LITE_MODEL, temperature=0.1, max_tokens=500)

    parsed_filters = {}
    if parse_result:
        try:
            json_match = re.search(r'\{[\s\S]*\}', parse_result)
            if json_match:
                parsed_filters = json.loads(json_match.group())
        except (json.JSONDecodeError, ValueError):
            logger.warning(f"Failed to parse smart search filters")

    if not parsed_filters:
        # Fallback: use the raw query as text search
        parsed_filters = {"search_text": query, "interpretation": query}

    # Build and execute the structured search
    conditions = ["c.status = 'active'", "c.is_processed = TRUE"]
    params = []

    skills = parsed_filters.get("skills", [])
    if skills:
        conditions.append("c.skills_technical && %s")
        params.append(skills)

    seniority = parsed_filters.get("seniority")
    if seniority:
        conditions.append("c.seniority_level = %s")
        params.append(seniority)

    min_exp = parsed_filters.get("min_experience_years")
    if min_exp is not None:
        conditions.append("c.total_experience_years >= %s")
        params.append(min_exp)

    max_exp = parsed_filters.get("max_experience_years")
    if max_exp is not None:
        conditions.append("c.total_experience_years <= %s")
        params.append(max_exp)

    location = parsed_filters.get("location")
    if location:
        conditions.append("c.location ILIKE %s")
        params.append(f"%{location}%")

    companies = parsed_filters.get("companies", [])
    if companies:
        company_conditions = []
        for comp in companies:
            company_conditions.append('(c.experience::text ILIKE %s OR c."current_company" ILIKE %s)')
            params.append(f"%{comp}%")
            params.append(f"%{comp}%")
        conditions.append(f"({' OR '.join(company_conditions)})")

    search_text = parsed_filters.get("search_text")
    if search_text:
        conditions.append("(c.name ILIKE %s OR c.\"current_role\" ILIKE %s OR c.raw_text ILIKE %s)")
        params.append(f"%{search_text}%")
        params.append(f"%{search_text}%")
        params.append(f"%{search_text}%")

    where = " AND ".join(conditions)
    params.append(50)  # limit

    with get_cursor() as cur:
        cur.execute(f"""
            SELECT c.id, c.public_id, c.name, c.email, c."current_role", c."current_company",
                   c.total_experience_years, c.seniority_level, c.skills_technical,
                   c.location, c.tags, c.quality_score, c.source, c.created_at
            FROM candidates c
            WHERE {where}
            ORDER BY c.quality_score DESC NULLS LAST, c.created_at DESC
            LIMIT %s
        """, params)
        cols = [desc[0] for desc in cur.description]
        candidates = [dict(zip(cols, row)) for row in cur.fetchall()]

    return {
        "candidates": candidates,
        "total": len(candidates),
        "query": query,
        "parsed_filters": parsed_filters,
        "interpretation": parsed_filters.get("interpretation", query),
    }


@router.get("/pending")
async def list_pending_candidates(request: Request, include_recent: bool = Query(False)):
    """List candidates uploaded but not yet pipeline-processed.

    When ``include_recent=true``, also returns recently-processed candidates
    (``is_processed=TRUE`` AND ``updated_at`` within the last 24 hours) so
    the user can see their work history in the pending UI.
    """
    user = _user_with_sector(get_current_user(request))
    uid = user.get("user_id") or 1
    extra_recent = (
        " OR (c.is_processed = TRUE AND c.updated_at > NOW() - INTERVAL '24 hours')"
        if include_recent else ""
    )
    from backend.core.database import get_cursor as _gc
    with _gc() as cur:
        cur.execute(f"""
            SELECT c.id, c.public_id, c.name, c.file_name, c.file_type, c.pdf_path,
                   c.created_at, c.updated_at, c.processing_error, c.is_processed,
                   COALESCE(
                     (SELECT json_build_object(
                       'run_id', t.run_id,
                       'total_steps', COUNT(*),
                       'done_steps', COUNT(*) FILTER (WHERE status IN ('done','skipped')),
                       'has_error', BOOL_OR(status = 'error'),
                       'has_running', BOOL_OR(status = 'running'),
                       'total_cost', COALESCE(SUM(cost_usd), 0),
                       'total_latency_ms', COALESCE(SUM(latency_ms), 0),
                       'last_step', (SELECT step_name FROM pipeline_trace WHERE candidate_id=c.id AND run_id=t.run_id ORDER BY step_order DESC LIMIT 1),
                       'last_status', (SELECT status FROM pipeline_trace WHERE candidate_id=c.id AND run_id=t.run_id ORDER BY step_order DESC LIMIT 1)
                     )
                     FROM pipeline_trace t
                     WHERE t.candidate_id = c.id
                     GROUP BY t.run_id
                     ORDER BY MAX(t.started_at) DESC
                     LIMIT 1), '{{}}'::json) AS latest_run
            FROM candidates c
            WHERE c.status = 'active' AND c.owner_id = %s
              AND (c.is_processed = FALSE{extra_recent})
            ORDER BY c.created_at DESC
            LIMIT 200
        """, (uid,))
        cols = [d[0] for d in cur.description]
        rows = [dict(zip(cols, r)) for r in cur.fetchall()]
        for r in rows:
            for k in ('created_at', 'updated_at'):
                if r.get(k): r[k] = r[k].isoformat()
            try:
                if r.get("pdf_path"):
                    p = Path(r["pdf_path"])
                    r["file_size"] = p.stat().st_size if p.exists() else None
                else:
                    r["file_size"] = None
            except Exception:
                r["file_size"] = None
    return {"candidates": rows, "total": len(rows)}


@router.get("/pipeline-events")
async def pipeline_events(request: Request, since_id: int = Query(0), limit: int = Query(200, le=500)):
    """Recent pipeline_trace rows for live CLI feed. since_id = last event id seen (monotonic cursor)."""
    get_current_user(request)
    from backend.core.database import get_cursor as _gc
    with _gc() as cur:
        cur.execute(
            """SELECT t.id, t.candidate_id, c.name AS cand_name, c.file_name AS cand_file,
                      t.run_id, t.step_order,
                      t.step_name, t.model, t.status, t.latency_ms, t.cost_usd,
                      t.started_at, t.finished_at, t.error_msg
               FROM pipeline_trace t
               LEFT JOIN candidates c ON c.id = t.candidate_id
               WHERE t.id > %s
               ORDER BY t.id ASC
               LIMIT %s""",
            (since_id, limit),
        )
        cols = [d[0] for d in cur.description]
        rows = [dict(zip(cols, r)) for r in cur.fetchall()]
        # Always report DB-wide max id so frontend can skip historical events on first probe
        cur.execute("SELECT COALESCE(MAX(id), 0) FROM pipeline_trace")
        db_max = cur.fetchone()[0]
        max_id = rows[-1]["id"] if rows else int(db_max)
        for r in rows:
            for k in ("started_at", "finished_at"):
                if r.get(k): r[k] = r[k].isoformat()
            if r.get("run_id"): r["run_id"] = str(r["run_id"])
            if r.get("cost_usd") is not None: r["cost_usd"] = float(r["cost_usd"])
    return {"events": rows, "max_id": max_id}


class BulkProcessReq(BaseModel):
    candidate_ids: list[int]
    force: bool = False  # when True, also re-process already-done candidates


@router.post("/bulk_process", dependencies=[Depends(require_role("recruiter"))])
async def bulk_process(body: BulkProcessReq, request: Request):
    """Enqueue pipeline runs. Returns immediately. Skips done ones unless force=True.
    Concurrent execution capped at MAX_PARALLEL_PIPELINES — rest queued.
    """
    get_current_user(request)
    import uuid as _uuid, time as _time
    _ensure_queue_workers()
    started = {}
    skipped = []
    now = _time.time()
    # Already-queued ids (avoid double-queue on rapid clicks)
    queued_ids = {j["cid"] for j in _pipeline_queue}
    for cid in body.candidate_ids:
        if cid in queued_ids or cid in _queue_running:
            skipped.append({"id": cid, "reason": "already_queued"})
            continue
        candidate = get_candidate(cid)
        if not candidate or not candidate.get("pdf_path"):
            skipped.append({"id": cid, "reason": "no_file"})
            continue
        if candidate.get("is_processed") and not body.force:
            skipped.append({"id": cid, "reason": "already_processed"})
            continue
        run_id = str(_uuid.uuid4())
        _pipeline_queue.append({
            "cid": cid,
            "pdf_path": candidate["pdf_path"],
            "run_id": run_id,
            "queued_at": now,
        })
        started[cid] = run_id
    return {"started": started, "count": len(started), "skipped": skipped, "queue_depth": len(_pipeline_queue)}


@router.get("/queue-status")
async def queue_status(request: Request):
    """Return current queue state: running cids (queue + standalone tasks) + queued cids in order."""
    get_current_user(request)
    # Merge: queue-worker-pulled running cids AND standalone-tracked tasks (run/reprocess/upload)
    active_alive = {cid for cid, t in _active_tasks.items() if t and not t.done()}
    running = list(_queue_running | active_alive)
    return {
        "running": running,
        "queued": [{"cid": j["cid"], "position": i + 1} for i, j in enumerate(_pipeline_queue)],
        "max_parallel": MAX_PARALLEL_PIPELINES,
        "queue_depth": len(_pipeline_queue),
    }


@router.get("/{candidate_id}/artifacts")
async def get_candidate_artifacts(candidate_id: str, request: Request):
    """Return counts/summaries of all derivative pipeline artifacts for a candidate."""
    get_current_user(request)
    candidate_id = _resolve_cid(candidate_id)
    if candidate_id is None:
        raise HTTPException(status_code=404, detail="Candidate not found")
    candidate = get_candidate(candidate_id)
    if not candidate:
        raise HTTPException(status_code=404, detail="Candidate not found")

    out: dict = {
        "candidate_id": candidate_id,
        "raw_text_chars": len(candidate.get("raw_text") or ""),
        "embeddings": {"count": 0, "by_chunk_type": {}, "model": EMBED_MODEL_FOR_ARTIFACTS, "dim": None},
        "qa_pairs": {"count": 0, "sample": []},
        "screenshots": {"count": 0, "paths": []},
        "pipeline_runs": [],
        "verified_critical_present": {
            "name": bool(candidate.get("name") and candidate.get("name") != "Unknown"),
            "dob": bool(candidate.get("dob")),
            "phone": bool(candidate.get("phone")),
            "nrc": bool(candidate.get("national_id")),
            "email": bool(candidate.get("email")),
        },
    }

    with get_cursor() as cur:
        # Embeddings: count, by chunk_type, vector dim
        cur.execute("""
            SELECT chunk_type, COUNT(*) AS n
            FROM candidate_embeddings
            WHERE candidate_id = %s
            GROUP BY chunk_type
        """, (candidate_id,))
        rows = cur.fetchall()
        by_type = {r[0]: int(r[1]) for r in rows}
        out["embeddings"]["count"] = sum(by_type.values())
        out["embeddings"]["by_chunk_type"] = by_type

        # vector dim — fetch first non-null embedding
        try:
            cur.execute("""
                SELECT vector_dims(embedding)
                FROM candidate_embeddings
                WHERE candidate_id = %s AND embedding IS NOT NULL
                LIMIT 1
            """, (candidate_id,))
            row = cur.fetchone()
            if row:
                out["embeddings"]["dim"] = int(row[0])
        except Exception:
            pass

        # QA pairs
        cur.execute("SELECT COUNT(*) FROM candidate_qa_pairs WHERE candidate_id = %s", (candidate_id,))
        out["qa_pairs"]["count"] = int(cur.fetchone()[0] or 0)
        cur.execute("""
            SELECT question, answer FROM candidate_qa_pairs
            WHERE candidate_id = %s
            ORDER BY id DESC LIMIT 5
        """, (candidate_id,))
        out["qa_pairs"]["sample"] = [
            {"q": r[0], "a": r[1]} for r in cur.fetchall()
        ]

        # Screenshots
        cur.execute("""
            SELECT path FROM candidate_screenshots
            WHERE candidate_id = %s
            ORDER BY page ASC, img_index ASC
        """, (candidate_id,))
        paths = [r[0] for r in cur.fetchall()]
        out["screenshots"]["count"] = len(paths)
        out["screenshots"]["paths"] = paths

        # Pipeline runs (group by run_id)
        cur.execute("""
            SELECT run_id::text,
                   MIN(started_at)              AS started_at,
                   COUNT(*)                     AS total_steps,
                   COUNT(*) FILTER (WHERE status IN ('done','skipped')) AS done_steps,
                   COALESCE(SUM(cost_usd), 0)   AS total_cost_usd,
                   COALESCE(SUM(latency_ms), 0) AS total_latency_ms,
                   BOOL_OR(status = 'error')    AS has_error,
                   BOOL_OR(status = 'running')  AS has_running
            FROM pipeline_trace
            WHERE candidate_id = %s
            GROUP BY run_id
            ORDER BY MIN(started_at) DESC
        """, (candidate_id,))
        for r in cur.fetchall():
            run_id_s, started_at, total_steps, done_steps, cost, latency, has_err, has_run = r
            if has_err:
                status = "error"
            elif has_run:
                status = "running"
            elif done_steps == total_steps:
                status = "done"
            else:
                status = "partial"
            out["pipeline_runs"].append({
                "run_id": run_id_s,
                "started_at": started_at.isoformat() if started_at else None,
                "total_steps": int(total_steps or 0),
                "done_steps": int(done_steps or 0),
                "total_cost_usd": float(cost or 0),
                "total_latency_ms": int(latency or 0),
                "status": status,
            })

    return out


@router.get("/{candidate_id}")
async def get_candidate_route(
    candidate_id: str,
    request: Request,
    scope: Optional[str] = Query(None),
    force_mask: bool = Query(False),
):
    """Get a single candidate with full details."""
    user = _user_with_sector(get_current_user(request))

    candidate_id = _resolve_cid(candidate_id)
    if candidate_id is None:
        raise HTTPException(status_code=404, detail="Candidate not found")
    candidate = get_candidate(candidate_id)
    if not candidate:
        raise HTTPException(status_code=404, detail="Candidate not found")

    # PII mask if pool scope and user not owner (or force_mask)
    if scope == "pool":
        uid_self = user.get("user_id")
        role = user.get("role")
        is_privileged = role in ("admin", "group_hr")
        should_mask = force_mask or (not is_privileged) or (candidate.get("owner_id") != uid_self)
        if should_mask:
            candidate["name"] = f"Anonymous#{candidate.get('id', candidate_id)}"
            candidate["email"] = None
            candidate["phone"] = None
            candidate["linkedin_url"] = None
            if candidate.get("location"):
                parts = [p.strip() for p in str(candidate["location"]).split(",") if p.strip()]
                candidate["location"] = parts[-1] if parts else None
            candidate["pool_masked"] = True
        else:
            candidate["pool_masked"] = False
    return candidate


MAX_UPLOAD_BYTES = 20 * 1024 * 1024  # 20MB hard cap
ALLOWED_UPLOAD_MIMES = {
    "application/pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "application/msword",
    "text/plain",
    # legacy image MIMEs allowed since CV scans are common
    "image/png", "image/jpeg",
}


@router.post("/upload")
@limiter.limit("5/minute")
async def upload_cvs(
    request: Request,
    files: list[UploadFile] = File(...),
    auto_process: bool = Form(True),
):
    """Upload one or more CVs. If auto_process=False, skip pipeline (user runs manually)."""
    user = get_current_user(request)
    if not has_min_role(user, "recruiter"):
        from backend.core.permissions import _log_denied
        _log_denied(user, request, "recruiter")
        raise HTTPException(403, "requires recruiter+ to upload CVs")

    # Early Content-Length guard (rejects oversized payloads before any pipeline work)
    cl = request.headers.get("content-length")
    if cl and cl.isdigit() and int(cl) > MAX_UPLOAD_BYTES:
        raise HTTPException(status_code=413, detail=f"Payload too large. Max {MAX_UPLOAD_BYTES // (1024*1024)}MB.")

    results = []
    for file in files:
        # MIME whitelist (early reject before reading body)
        ctype = (file.content_type or "").lower()
        if ctype and ctype not in ALLOWED_UPLOAD_MIMES:
            results.append({
                "filename": file.filename,
                "status": "error",
                "error": f"Unsupported MIME type: {ctype}",
            })
            continue

        # Validate file type
        ext = Path(file.filename or "").suffix.lower()
        if ext not in {".pdf", ".docx", ".doc", ".png", ".jpg", ".jpeg", ".txt"}:
            results.append({
                "filename": file.filename,
                "status": "error",
                "error": f"Unsupported file type: {ext}",
            })
            continue

        # Read and validate file content
        content = await file.read()

        if len(content) > MAX_UPLOAD_BYTES:
            results.append({"filename": file.filename, "status": "error",
                            "error": f"File too large ({len(content) // 1024 // 1024}MB). Max {MAX_UPLOAD_BYTES // (1024*1024)}MB."})
            continue

        # Validate file content (not just extension)
        if content[:4] == b'%PDF':
            pass  # valid PDF
        elif content[:2] == b'PK':
            pass  # valid DOCX (ZIP-based)
        elif content[:8] == b'\x89PNG\r\n\x1a\n' or content[:3] == b'\xff\xd8\xff':
            pass  # valid PNG / JPEG
        else:
            results.append({"filename": file.filename, "status": "error", "error": "Invalid file content — PDF, DOCX, PNG, JPG accepted"})
            continue

        # Save file
        file_id = str(uuid.uuid4())
        save_path = CV_DIR / f"{file_id}{ext}"
        save_path.write_bytes(content)

        # Create candidate record (unprocessed)
        try:
            candidate_id = insert_candidate({
                "name": Path(file.filename or "unknown").stem.replace("_", " ").replace("-", " ").title(),
                "email": None,
                "phone": None,
                "location": None,
                "linkedin_url": None,
                "current_role": None,
                "current_company": None,
                "total_experience_years": None,
                "seniority_level": None,
                "skills_technical": [],
                "skills_soft": [],
                "tools": [],
                "languages": [],
                "experience": "[]",
                "education": "[]",
                "certifications": "[]",
                "projects": "[]",
                "summary_short": None,
                "summary_detailed": None,
                "raw_text": None,
                "pdf_path": str(save_path),
                "file_type": ext.lstrip("."),
                "file_name": file.filename,
                "page_count": 0,
                "quality_score": 0,
                "tags": [],
                "source": "upload",
                "status": "active",
                "is_processed": False,
                "owner_id": _user_with_sector(get_current_user(request)).get("user_id") or 1,
                "sector_id": _user_with_sector(get_current_user(request)).get("sector_id"),
                "visibility": "private",
            })

            results.append({
                "filename": file.filename,
                "candidate_id": candidate_id,
                "file_url": f"/api/candidates/{candidate_id}/file",
                "status": "uploaded",
                "auto_process": auto_process,
                "message": "Queued for processing" if auto_process else "Uploaded. Click RUN PIPELINE to process.",
            })

            # Check for duplicates by email and log if found
            try:
                from backend.core.database import get_cursor as _get_cursor
                email_val = None  # email is None at upload time, checked after processing
                # We'll check by name (file-derived) for now
                with _get_cursor() as dup_cur:
                    cand_name = Path(file.filename or "unknown").stem.replace("_", " ").replace("-", " ").title()
                    try:
                        dup_cur.execute(
                            "SELECT id, name FROM candidates WHERE similarity(name, %s) > 0.6 AND id != %s AND status = 'active' LIMIT 5",
                            (cand_name, candidate_id),
                        )
                        dupes = dup_cur.fetchall()
                        if dupes:
                            dup_ids = [d[0] for d in dupes]
                            dup_cur.execute(
                                "INSERT INTO candidate_activity (candidate_id, event_type, description) VALUES (%s, %s, %s)",
                                (candidate_id, "duplicate_warning", f"Potential duplicates detected: candidate IDs {dup_ids}"),
                            )
                            logger.info(f"Duplicate warning for candidate {candidate_id}: matches {dup_ids}")
                    except Exception:
                        pass  # pg_trgm may not be available
            except Exception as dup_err:
                logger.debug(f"Duplicate check skipped: {dup_err}")

            # Trigger async processing only if auto_process=True
            if auto_process:
                import asyncio
                _track(candidate_id, asyncio.create_task(_process_cv_background(candidate_id, str(save_path), file.filename)))
                logger.info(f"CV uploaded: {file.filename} -> candidate_id={candidate_id} (auto-process)")
            else:
                logger.info(f"CV uploaded: {file.filename} -> candidate_id={candidate_id} (manual process)")

        except Exception as e:
            logger.error(f"Failed to save candidate: {e}")
            results.append({
                "filename": file.filename,
                "status": "error",
                "error": str(e),
            })

    return {
        "uploaded": len([r for r in results if r["status"] == "uploaded"]),
        "errors": len([r for r in results if r["status"] == "error"]),
        "results": results,
    }


# ---------------------------------------------------------------------------
# Bulk CV upload — multi-file with email-dedup + per-file error isolation
# ---------------------------------------------------------------------------

BULK_UPLOAD_CAP = 50

_EMAIL_RE = re.compile(r"[A-Za-z0-9._%+\-]+@[A-Za-z0-9.\-]+\.[A-Za-z]{2,}")


def _quick_extract_email(content: bytes, ext: str) -> Optional[str]:
    """Best-effort sync email scan. Falls back to None on any error."""
    try:
        text = ""
        if ext == ".txt":
            text = content.decode("utf-8", errors="ignore")
        elif ext == ".pdf":
            try:
                import io as _bio
                import fitz  # PyMuPDF
                with fitz.open(stream=content, filetype="pdf") as doc:
                    parts = []
                    for i, page in enumerate(doc):
                        if i >= 3:  # first 3 pages is plenty for an email
                            break
                        parts.append(page.get_text("text") or "")
                    text = "\n".join(parts)
            except Exception:
                text = ""
        elif ext in {".docx"}:
            try:
                import io as _bio
                from docx import Document  # python-docx
                doc = Document(_bio.BytesIO(content))
                text = "\n".join(p.text for p in doc.paragraphs)
            except Exception:
                text = ""
        else:
            text = ""
        if not text:
            return None
        m = _EMAIL_RE.search(text)
        if not m:
            return None
        em = m.group(0).strip().lower()
        # filter out obvious noise / fake placeholders
        if em.endswith(".png") or em.endswith(".jpg"):
            return None
        return em
    except Exception:
        return None


def _email_exists(email: str) -> Optional[dict]:
    if not email:
        return None
    try:
        with get_cursor() as cur:
            cur.execute(
                "SELECT id, public_id, name, email FROM candidates "
                "WHERE LOWER(email) = %s AND status = 'active' LIMIT 1",
                (email.lower(),),
            )
            row = cur.fetchone()
            if not row:
                return None
            return {"id": row[0], "public_id": row[1], "name": row[2], "email": row[3]}
    except Exception:
        return None


@router.post("/bulk-upload")
@limiter.limit("5/minute")
async def bulk_upload_cvs(
    request: Request,
    files: list[UploadFile] = File(...),
):
    """Bulk-upload CVs. Caps at 50 files. Skips duplicates by email; isolates per-file errors."""
    user = _user_with_sector(get_current_user(request))
    if not has_min_role(user, "recruiter"):
        from backend.core.permissions import _log_denied
        _log_denied(user, request, "recruiter")
        raise HTTPException(403, "requires recruiter+ to upload CVs")

    if len(files) > BULK_UPLOAD_CAP:
        raise HTTPException(
            status_code=400,
            detail=f"Too many files. Max {BULK_UPLOAD_CAP} per request (got {len(files)}).",
        )

    total = len(files)
    created: list[dict] = []
    skipped: list[dict] = []
    errors: list[dict] = []

    for file in files:
        fname = file.filename or "unknown"
        try:
            ctype = (file.content_type or "").lower()
            if ctype and ctype not in ALLOWED_UPLOAD_MIMES:
                errors.append({"filename": fname, "error": f"Unsupported MIME: {ctype}"})
                continue

            ext = Path(fname).suffix.lower()
            if ext not in {".pdf", ".docx", ".doc", ".txt"}:
                errors.append({"filename": fname, "error": f"Unsupported file type: {ext}"})
                continue

            content = await file.read()

            if len(content) > MAX_UPLOAD_BYTES:
                errors.append({
                    "filename": fname,
                    "error": f"File too large ({len(content)//1024//1024}MB). Max {MAX_UPLOAD_BYTES//(1024*1024)}MB.",
                })
                continue

            # content sniff
            if content[:4] == b"%PDF":
                pass
            elif content[:2] == b"PK":
                pass
            elif ext == ".txt":
                pass
            elif ext == ".doc":
                pass
            else:
                errors.append({"filename": fname, "error": "Invalid file content"})
                continue

            # Content-hash dedupe (catches renamed/identical files; works on scanned PDFs)
            import hashlib as _hl
            content_hash = _hl.sha256(content).hexdigest()
            try:
                with get_cursor() as cur:
                    cur.execute(
                        "SELECT id, public_id, name FROM candidates "
                        "WHERE cv_content_hash=%s AND status='active' LIMIT 1",
                        (content_hash,),
                    )
                    hrow = cur.fetchone()
                    if hrow:
                        skipped.append({
                            "filename": fname,
                            "reason": "duplicate_content",
                            "existing_id": hrow[0],
                            "existing_public_id": hrow[1],
                            "existing_name": hrow[2],
                        })
                        continue
            except Exception as e:
                logger.warning(f"bulk-upload content-hash dedup failed for {fname}: {e}")

            # Quick parse → email
            email = _quick_extract_email(content, ext)

            # Email-based dedupe
            if email:
                dup = _email_exists(email)
                if dup:
                    skipped.append({
                        "filename": fname,
                        "email": email,
                        "reason": "duplicate_email",
                        "existing_id": dup["id"],
                        "existing_public_id": dup["public_id"],
                        "existing_name": dup["name"],
                    })
                    continue

            # Save file
            file_id = str(uuid.uuid4())
            save_path = CV_DIR / f"{file_id}{ext}"
            save_path.write_bytes(content)

            cand_name = Path(fname).stem.replace("_", " ").replace("-", " ").title()
            candidate_id = insert_candidate({
                "name": cand_name,
                "email": email,
                "phone": None,
                "location": None,
                "linkedin_url": None,
                "current_role": None,
                "current_company": None,
                "total_experience_years": None,
                "seniority_level": None,
                "skills_technical": [],
                "skills_soft": [],
                "tools": [],
                "languages": [],
                "experience": "[]",
                "education": "[]",
                "certifications": "[]",
                "projects": "[]",
                "summary_short": None,
                "summary_detailed": None,
                "raw_text": None,
                "pdf_path": str(save_path),
                "file_type": ext.lstrip("."),
                "file_name": fname,
                "page_count": 0,
                "quality_score": 0,
                "tags": [],
                "source": "bulk_upload",
                "status": "active",
                "is_processed": False,
                "owner_id": user.get("user_id") or 1,
                "sector_id": user.get("sector_id"),
                "visibility": "private",
            })

            # Stamp content hash + fetch public_id
            pub = None
            try:
                with get_cursor() as cur:
                    cur.execute(
                        "UPDATE candidates SET cv_content_hash=%s WHERE id=%s RETURNING public_id",
                        (content_hash, candidate_id),
                    )
                    r = cur.fetchone()
                    if r:
                        pub = r[0]
            except Exception:
                pass

            created.append({
                "id": candidate_id,
                "public_id": pub,
                "name": cand_name,
                "email": email,
                "filename": fname,
            })

            # Kick off pipeline async (matches /upload behaviour)
            try:
                import asyncio
                _track(candidate_id, asyncio.create_task(_process_cv_background(candidate_id, str(save_path), fname)))
            except Exception as bg_err:
                logger.warning(f"bulk-upload: failed to start pipeline for {fname}: {bg_err}")

        except Exception as e:
            logger.error(f"bulk-upload error on {fname}: {e}")
            errors.append({"filename": fname, "error": str(e)[:300]})

    return {
        "total": total,
        "created": len(created),
        "skipped_duplicates": len(skipped),
        "errors": errors,
        "skipped": skipped,
        "candidates": [{k: v for k, v in c.items() if k != "filename"} for c in created],
    }


async def _process_cv_background(candidate_id: int, file_path: str, filename: str):
    """Background task to run the 11-step CV pipeline."""
    try:
        from backend.core.cv_pipeline import process_cv
        result = await process_cv(candidate_id, file_path)
        logger.info(f"CV processed: {filename} -> {result.get('status')}")
    except Exception as e:
        logger.error(f"CV processing failed for {filename}: {e}")
        from backend.core.database import get_cursor
        with get_cursor() as cur:
            cur.execute(
                "UPDATE candidates SET processing_error = %s, updated_at = NOW() WHERE id = %s",
                (str(e)[:500], candidate_id),
            )


@router.post("/{candidate_id}/reprocess")
async def reprocess_cv(candidate_id: str, request: Request):
    """Re-run the CV processing pipeline for a candidate."""
    user = get_current_user(request)
    candidate_id = _resolve_cid(candidate_id)
    if candidate_id is None:
        raise HTTPException(status_code=404, detail="Candidate not found")
    candidate = get_candidate(candidate_id)
    if not candidate:
        raise HTTPException(status_code=404, detail="Candidate not found")
    if not candidate.get("pdf_path"):
        raise HTTPException(status_code=400, detail="No file to process")
    ensure_owner_or_min(user, candidate.get("owner_id"), "admin", request)

    import asyncio
    _track(candidate_id, asyncio.create_task(_process_cv_background(
        candidate_id, candidate["pdf_path"], candidate.get("file_name", "unknown"),
    )))
    return {"message": "Reprocessing started", "candidate_id": candidate_id}


@router.post("/{candidate_id}/process")
async def process_cv_manual(candidate_id: str, request: Request):
    """Manually trigger pipeline for an uploaded-but-unprocessed candidate."""
    user = get_current_user(request)
    candidate_id = _resolve_cid(candidate_id)
    if candidate_id is None:
        raise HTTPException(404, "Candidate not found")
    candidate = get_candidate(candidate_id)
    if not candidate:
        raise HTTPException(404, "Candidate not found")
    if not candidate.get("pdf_path"):
        raise HTTPException(400, "No file to process")
    if candidate.get("is_processed"):
        raise HTTPException(409, "Already processed. Use /reprocess to re-run.")
    ensure_owner_or_min(user, candidate.get("owner_id"), "admin", request)

    import asyncio, uuid as _uuid
    run_id = str(_uuid.uuid4())

    async def _run():
        try:
            from backend.core.cv_pipeline import process_cv
            await process_cv(candidate_id, candidate["pdf_path"], run_id=run_id)
        except Exception as e:
            logger.error(f"Manual process failed for cid={candidate_id}: {e}")
            from backend.core.database import get_cursor
            with get_cursor() as cur:
                cur.execute(
                    "UPDATE candidates SET processing_error = %s, updated_at = NOW() WHERE id = %s",
                    (str(e)[:500], candidate_id),
                )

    _track(candidate_id, asyncio.create_task(_run()))
    return {"candidate_id": candidate_id, "run_id": run_id, "message": "Pipeline started"}


@router.post("/{candidate_id}/cancel")
async def cancel_cv_pipeline(candidate_id: str, request: Request):
    """Cancel an in-flight pipeline run OR dequeue a queued candidate."""
    user = get_current_user(request)
    cid = _resolve_cid(candidate_id)
    if cid is None:
        raise HTTPException(404, "Candidate not found")
    cand = get_candidate(cid)
    if not cand:
        raise HTTPException(404, "Candidate not found")
    ensure_owner_or_min(user, cand.get("owner_id"), "admin", request)
    # Try queue removal first (cheap)
    before = len(_pipeline_queue)
    _pipeline_queue[:] = [j for j in _pipeline_queue if j["cid"] != cid]
    dequeued = before - len(_pipeline_queue)
    task = _active_tasks.get(cid)
    if task and not task.done():
        task.cancel()
        return {"candidate_id": cid, "cancelled": True, "dequeued": dequeued}
    if dequeued:
        return {"candidate_id": cid, "cancelled": True, "dequeued": dequeued, "message": "Removed from queue"}
    return {"candidate_id": cid, "cancelled": False, "message": "No active task"}


@router.post("/cancel-all")
async def cancel_all_pipelines(request: Request):
    """Cancel ALL in-flight pipelines AND clear the queue. Admin only."""
    user = get_current_user(request)
    if not has_min_role(user, "admin"):
        raise HTTPException(403, "Admin required")
    dequeued = len(_pipeline_queue)
    _pipeline_queue.clear()
    cancelled = 0
    for cid, task in list(_active_tasks.items()):
        if task and not task.done():
            task.cancel()
            cancelled += 1
    return {"cancelled": cancelled, "dequeued": dequeued, "active_remaining": len(_active_tasks)}


@router.get("/{candidate_id}/file/sign")
async def sign_candidate_file_url(candidate_id: str, request: Request):
    """Issue a short-lived HMAC-signed URL for /file. Bearer auth required."""
    get_current_user(request)
    candidate_id = _resolve_cid(candidate_id)
    if candidate_id is None:
        raise HTTPException(404, "Candidate not found")
    candidate = get_candidate(candidate_id)
    if not candidate:
        raise HTTPException(404, "Candidate not found")
    exp = int(_time_mod.time()) + FILE_SIGN_TTL
    sig = _compute_file_sig(candidate_id, exp)
    return {
        "url": f"/api/candidates/{candidate_id}/file?sig={sig}&exp={exp}",
        "exp": exp,
        "ttl_s": FILE_SIGN_TTL,
    }


@router.get("/{candidate_id}/file")
async def get_candidate_file(
    candidate_id: str,
    request: Request,
    token: str | None = None,
    sig: str | None = None,
    exp: int | None = None,
):
    """Stream the uploaded CV file (PDF/image/DOCX) for preview.

    Auth precedence:
      1. sig + exp HMAC (preferred, short-lived)
      2. ?token=  legacy Bearer token (deprecated, still accepted)
      3. Authorization header
    """
    resolved_cid = _resolve_cid(candidate_id)
    if resolved_cid is None:
        raise HTTPException(404, "Candidate not found")
    candidate_id = resolved_cid
    if sig and exp is not None:
        now = int(_time_mod.time())
        if exp < now:
            raise HTTPException(401, "Signed URL expired")
        expected = _compute_file_sig(candidate_id, int(exp))
        if not hmac.compare_digest(expected, sig):
            raise HTTPException(401, "Invalid signature")
    elif token:
        logger.warning(
            "Deprecated ?token= used on /candidates/%s/file — switch to /file/sign HMAC URLs",
            candidate_id,
        )
        user = validate_token(token)
        if not user:
            raise HTTPException(401, "Invalid token")
    else:
        get_current_user(request)
    candidate = get_candidate(candidate_id)
    if not candidate:
        raise HTTPException(404, "Candidate not found")
    path_s = candidate.get("pdf_path")
    if not path_s:
        raise HTTPException(404, "No file")
    p = Path(path_s)
    if not p.exists():
        return JSONResponse(
            status_code=404,
            content={
                "error": "file_missing",
                "message": "Original file no longer on disk. Re-upload required.",
            },
        )
    from fastapi.responses import FileResponse
    ext = p.suffix.lower()
    mime = {
        ".pdf": "application/pdf",
        ".png": "image/png", ".jpg": "image/jpeg", ".jpeg": "image/jpeg",
        ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        ".doc": "application/msword",
    }.get(ext, "application/octet-stream")
    name = candidate.get("file_name") or p.name
    disposition = "attachment" if request.query_params.get("download") == "1" else "inline"
    return FileResponse(
        str(p),
        media_type=mime,
        headers={"Content-Disposition": f'{disposition}; filename="{name}"'},
    )


@router.get("/{candidate_id}/export/report.docx")
@limiter.limit("30/minute")
async def export_candidate_report(candidate_id: str, request: Request, token: str | None = None):
    """Generate comprehensive Word DOCX report with all candidate insights."""
    if token:
        user = validate_token(token)
        if not user:
            raise HTTPException(401, "Invalid token")
    else:
        get_current_user(request)

    candidate_id = _resolve_cid(candidate_id)
    if candidate_id is None:
        raise HTTPException(404, "Candidate not found")
    candidate = get_candidate(candidate_id)
    if not candidate:
        raise HTTPException(404, "Candidate not found")

    for fld in ("experience", "education", "certifications", "projects",
                "skills_technical", "skills_soft", "tools", "languages"):
        v = candidate.get(fld)
        if isinstance(v, str):
            try:
                candidate[fld] = json.loads(v)
            except Exception:
                candidate[fld] = []

    matches, scorecards, notes, flags, comp_scores, pipeline = [], [], [], [], [], []
    try:
        with get_cursor() as cur:
            cur.execute("""SELECT p.title, pc.match_score_composite, pc.stage,
                                  pc.skills_matched, pc.skills_missing
                           FROM position_candidates pc JOIN positions p ON p.id=pc.position_id
                           WHERE pc.candidate_id=%s ORDER BY pc.match_score_composite DESC NULLS LAST LIMIT 10""",
                        (candidate_id,))
            cols = [d[0] for d in cur.description]
            matches = [dict(zip(cols, r)) for r in cur.fetchall()]
            cur.execute("""SELECT sc.overall_score, sc.recommendation, sc.strengths, sc.concerns,
                                  i.interview_type, p.title AS position_title
                           FROM interview_scorecards sc JOIN interviews i ON i.id=sc.interview_id
                           LEFT JOIN positions p ON p.id=i.position_id
                           WHERE i.candidate_id=%s ORDER BY sc.created_at DESC LIMIT 10""", (candidate_id,))
            cols = [d[0] for d in cur.description]
            scorecards = [dict(zip(cols, r)) for r in cur.fetchall()]
            cur.execute("""SELECT content, note_type, created_at FROM candidate_notes
                           WHERE candidate_id=%s ORDER BY created_at DESC LIMIT 20""", (candidate_id,))
            cols = [d[0] for d in cur.description]
            notes = [dict(zip(cols, r)) for r in cur.fetchall()]
            try:
                cur.execute("""SELECT flag_type, title, description FROM candidate_flags
                               WHERE candidate_id=%s ORDER BY created_at DESC LIMIT 30""", (candidate_id,))
                cols = [d[0] for d in cur.description]
                flags = [dict(zip(cols, r)) for r in cur.fetchall()]
            except Exception: pass
            try:
                cur.execute("""SELECT c.code, c.name, ccs.score, ccs.source
                               FROM candidate_competency_scores ccs
                               JOIN competencies c ON c.id=ccs.competency_id
                               WHERE ccs.candidate_id=%s ORDER BY ccs.score DESC""", (candidate_id,))
                cols = [d[0] for d in cur.description]
                comp_scores = [dict(zip(cols, r)) for r in cur.fetchall()]
            except Exception: pass
            try:
                cur.execute("""SELECT step_name, status, latency_ms, cost_usd, model
                               FROM pipeline_trace WHERE candidate_id=%s
                               ORDER BY started_at DESC LIMIT 30""", (candidate_id,))
                cols = [d[0] for d in cur.description]
                pipeline = [dict(zip(cols, r)) for r in cur.fetchall()]
            except Exception: pass
    except Exception as e:
        logging.getLogger(__name__).warning(f"Report aux fetch failed: {e}")

    # --- Truncate large free-text fields to bound DOCX size / build time ---
    _MAX_TEXT = 50_000
    def _truncate(v):
        if isinstance(v, str) and len(v) > _MAX_TEXT:
            return v[:_MAX_TEXT] + "...[truncated]"
        return v
    for _fld in ("ai_summary", "summary_short", "raw_text", "cv_text", "notes"):
        if candidate.get(_fld):
            candidate[_fld] = _truncate(candidate[_fld])
    for _n in notes:
        if isinstance(_n, dict) and _n.get("content"):
            _n["content"] = _truncate(_n["content"])

    def _build_docx():
        from io import BytesIO
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        INK = RGBColor(0x38, 0x38, 0x32)
        PRIMARY = RGBColor(0x00, 0x75, 0x18)
        ACCENT = RGBColor(0x00, 0xa0, 0x30)
        MUTED = RGBColor(0x66, 0x66, 0x60)

        doc = Document()
        for section in doc.sections:
            section.top_margin = Cm(1.6)
            section.bottom_margin = Cm(1.6)
            section.left_margin = Cm(1.8)
            section.right_margin = Cm(1.8)

        def style_run(run, *, size=11, bold=False, color=INK, name="Calibri"):
            run.font.size = Pt(size)
            run.font.bold = bold
            run.font.color.rgb = color
            run.font.name = name

        def add_h(text, level=1, color=None):
            p = doc.add_paragraph()
            r = p.add_run(text.upper() if level == 1 else text)
            sizes = {1: 22, 2: 14, 3: 12}
            style_run(r, size=sizes.get(level, 11), bold=True,
                      color=color or (PRIMARY if level == 1 else INK))
            if level == 1:
                p.paragraph_format.space_after = Pt(8)
            return p

        def add_kv(label, value):
            p = doc.add_paragraph()
            r1 = p.add_run(f"{label}: ")
            style_run(r1, size=10, bold=True, color=MUTED)
            r2 = p.add_run(str(value) if value not in (None, "", []) else "—")
            style_run(r2, size=10)
            p.paragraph_format.space_after = Pt(2)

        def add_para(text, *, size=10, color=INK, italic=False):
            p = doc.add_paragraph()
            r = p.add_run(text)
            style_run(r, size=size, color=color)
            r.font.italic = italic
            return p

        def add_divider():
            p = doc.add_paragraph()
            pPr = p._element.get_or_add_pPr()
            bdr = OxmlElement('w:pBdr')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '6')
            bottom.set(qn('w:color'), '383832')
            bdr.append(bottom)
            pPr.append(bdr)

        # --- HEADER ---
        title = doc.add_paragraph()
        tr = title.add_run("CANDIDATE REPORT")
        style_run(tr, size=10, bold=True, color=MUTED)
        title.paragraph_format.space_after = Pt(2)

        name_p = doc.add_paragraph()
        nr = name_p.add_run((candidate.get("name") or "Unknown").upper())
        style_run(nr, size=28, bold=True, color=INK)
        name_p.paragraph_format.space_after = Pt(2)

        sub_p = doc.add_paragraph()
        role_text = candidate.get("current_role") or ""
        if candidate.get("current_company"):
            role_text += f" at {candidate['current_company']}"
        if candidate.get("total_experience_years"):
            role_text += f"  |  {candidate['total_experience_years']} yr exp"
        if candidate.get("seniority_level"):
            role_text += f"  |  {candidate['seniority_level']}"
        sr = sub_p.add_run(role_text)
        style_run(sr, size=12, color=ACCENT)
        sub_p.paragraph_format.space_after = Pt(8)

        contact_p = doc.add_paragraph()
        cl = []
        if candidate.get("location"): cl.append(f"📍 {candidate['location']}")
        if candidate.get("email"): cl.append(f"✉ {candidate['email']}")
        if candidate.get("phone"): cl.append(f"☎ {candidate['phone']}")
        if candidate.get("linkedin") or candidate.get("linkedin_url"):
            cl.append(f"in/ {candidate.get('linkedin') or candidate.get('linkedin_url')}")
        cr = contact_p.add_run("  ·  ".join(cl) or "—")
        style_run(cr, size=10, color=MUTED)
        add_divider()

        # --- AT-A-GLANCE ---
        add_h("AT A GLANCE", level=1)
        tbl = doc.add_table(rows=1, cols=5)
        tbl.autofit = True
        headers = ["Experience", "Quality Score", "Verified", "Positions", "Tags"]
        values = [
            f"{candidate.get('total_experience_years', 0)}y",
            f"{candidate.get('quality_score', 0)}/100",
            "✓" if candidate.get("opus_verified") else "—",
            str(len(matches)),
            str(len(candidate.get("tags") or [])),
        ]
        for i, h in enumerate(headers):
            cell = tbl.rows[0].cells[i]
            cell.text = ""
            p1 = cell.paragraphs[0]
            r1 = p1.add_run(values[i])
            style_run(r1, size=18, bold=True, color=PRIMARY)
            p2 = cell.add_paragraph()
            r2 = p2.add_run(h.upper())
            style_run(r2, size=8, bold=True, color=MUTED)
        doc.add_paragraph()

        # --- AI SUMMARY ---
        if candidate.get("ai_summary"):
            add_h("AI EXECUTIVE SUMMARY", level=1)
            for block in str(candidate["ai_summary"]).split("\n\n"):
                block = block.strip()
                if not block: continue
                if block.startswith("**") and block.endswith("**"):
                    add_h(block.strip("*").strip(), level=2)
                elif block.startswith("## "):
                    add_h(block[3:].strip(), level=2)
                else:
                    p = doc.add_paragraph()
                    rest = block
                    while "**" in rest:
                        pre, _, after = rest.partition("**")
                        bold_part, _, after2 = after.partition("**")
                        if pre:
                            r = p.add_run(pre); style_run(r, size=10)
                        if bold_part:
                            r = p.add_run(bold_part); style_run(r, size=10, bold=True)
                        rest = after2
                    if rest:
                        r = p.add_run(rest); style_run(r, size=10)
                    p.paragraph_format.space_after = Pt(4)
            add_divider()

        # --- CV SHORT SUMMARY ---
        if candidate.get("summary_short"):
            add_h("CV SUMMARY", level=1)
            add_para(candidate["summary_short"], size=10)
            add_divider()

        # --- EXPERIENCE ---
        if candidate.get("experience"):
            add_h("EXPERIENCE", level=1)
            for exp in candidate["experience"][:20]:
                if not isinstance(exp, dict): continue
                line = doc.add_paragraph()
                r = line.add_run((exp.get("role") or "").upper())
                style_run(r, size=11, bold=True, color=INK)
                if exp.get("company"):
                    r2 = line.add_run(f"  ·  {exp['company']}")
                    style_run(r2, size=11, color=PRIMARY)
                dates = " — ".join(filter(None, [exp.get("start_date"), exp.get("end_date") or "Present"]))
                if dates:
                    r3 = line.add_run(f"  ({dates})")
                    style_run(r3, size=9, color=MUTED, name="Calibri")
                    r3.font.italic = True
                if exp.get("description"):
                    add_para(exp["description"], size=10)
                line.paragraph_format.space_after = Pt(4)
            add_divider()

        # --- SKILLS ---
        sk_tech = candidate.get("skills_technical") or []
        sk_soft = candidate.get("skills_soft") or []
        sk_tools = candidate.get("tools") or []
        if sk_tech or sk_soft or sk_tools:
            add_h("SKILLS", level=1)
            if sk_tech: add_kv("Technical", ", ".join(map(str, sk_tech[:30])))
            if sk_soft: add_kv("Soft", ", ".join(map(str, sk_soft[:20])))
            if sk_tools: add_kv("Tools", ", ".join(map(str, sk_tools[:30])))
            if candidate.get("languages"):
                add_kv("Languages", ", ".join(map(str, candidate["languages"][:10])))
            add_divider()

        # --- EDUCATION ---
        if candidate.get("education"):
            add_h("EDUCATION", level=1)
            for ed in candidate["education"][:10]:
                if not isinstance(ed, dict): continue
                p = doc.add_paragraph()
                r = p.add_run(ed.get("degree") or "—")
                style_run(r, size=11, bold=True)
                if ed.get("institution"):
                    r2 = p.add_run(f"  ·  {ed['institution']}")
                    style_run(r2, size=11, color=PRIMARY)
                if ed.get("graduation_year") or ed.get("year"):
                    r3 = p.add_run(f"  ({ed.get('graduation_year') or ed.get('year')})")
                    style_run(r3, size=9, color=MUTED)
            add_divider()

        # --- CERTIFICATIONS ---
        add_h("CERTIFICATIONS", level=1)
        if candidate.get("certifications"):
            for cert in candidate["certifications"][:15]:
                if isinstance(cert, dict):
                    add_kv(cert.get("name", "—"), cert.get("issuer", ""))
                else:
                    add_para(f"• {cert}", size=10)
        else:
            add_para("No certifications listed.", size=10, italic=True, color=MUTED)
        add_divider()

        # --- COMPETENCIES ---
        add_h("COMPETENCIES (KF4D)", level=1)
        if comp_scores:
            for c in comp_scores:
                add_kv(c.get("name") or c.get("code"), f"{c.get('score', 0)}/5  ({c.get('source','')})")
        else:
            add_para("No competency scores recorded yet.", size=10, italic=True, color=MUTED)
        add_divider()

        # --- ASSIGNMENTS / POSITION MATCHES ---
        add_h("ASSIGNMENTS & POSITION MATCHES", level=1)
        if matches:
            for m in matches:
                p = doc.add_paragraph()
                r = p.add_run(m.get("title") or "—")
                style_run(r, size=11, bold=True)
                r2 = p.add_run(f"  ·  {m.get('match_score_composite') or 0}%  ·  Stage: {m.get('stage') or '—'}")
                style_run(r2, size=10, color=MUTED)
                if m.get("skills_missing"):
                    add_para(f"Missing skills: {m['skills_missing']}", size=9, italic=True, color=MUTED)
        else:
            add_para("No positions assigned yet.", size=10, italic=True, color=MUTED)
        add_divider()

        # --- FLAGS ---
        add_h("AI FLAGS", level=1)
        if flags:
            for f in flags:
                color = {"red": RGBColor(0xc0, 0x10, 0x10), "amber": RGBColor(0xc0, 0x80, 0x00),
                         "green": ACCENT}.get(f.get("flag_type"), MUTED)
                p = doc.add_paragraph()
                r = p.add_run(f"[{(f.get('flag_type') or '').upper()}] {f.get('title') or ''}")
                style_run(r, size=10, bold=True, color=color)
                if f.get("description"):
                    add_para(f["description"], size=9, color=MUTED)
        else:
            add_para("No flags raised.", size=10, italic=True, color=MUTED)
        add_divider()

        # --- INTERVIEWS / SCORECARDS ---
        add_h("INTERVIEWS & SCORECARDS", level=1)
        if scorecards:
            for s in scorecards:
                p = doc.add_paragraph()
                r = p.add_run(f"{s.get('position_title') or 'N/A'}  ·  {s.get('overall_score','—')}/5  ·  {s.get('recommendation','')}")
                style_run(r, size=10, bold=True)
                if s.get("strengths"): add_para(f"Strengths: {s['strengths']}", size=9)
                if s.get("concerns"): add_para(f"Concerns: {s['concerns']}", size=9, color=MUTED)
        else:
            add_para("No interviews scheduled or scorecards submitted yet.", size=10, italic=True, color=MUTED)
        add_divider()

        # --- NOTES ---
        add_h("NOTES", level=1)
        if notes:
            for n in notes[:10]:
                ts = (n.get("created_at") or "")
                add_para(f"[{n.get('note_type','general')}] {ts}", size=9, color=MUTED)
                add_para(n.get("content") or "", size=10)
        else:
            add_para("No notes recorded yet.", size=10, italic=True, color=MUTED)
        add_divider()

        # --- DEMOGRAPHICS ---
        add_h("DEMOGRAPHICS", level=1)
        demo_keys = ["dob", "national_id", "gender", "marital_status", "nationality",
                     "religion", "height", "weight", "father_name"]
        if any(candidate.get(k) for k in demo_keys):
            for k in demo_keys:
                if candidate.get(k):
                    add_kv(k.replace("_", " ").title(), candidate[k])
        else:
            add_para("No demographic data extracted.", size=10, italic=True, color=MUTED)
        add_divider()

        # --- ACTIVITY ---
        activity = []
        try:
            with get_cursor() as cur:
                cur.execute("""SELECT action, details, created_at FROM candidate_activity
                               WHERE candidate_id=%s ORDER BY created_at DESC LIMIT 30""", (candidate_id,))
                cols = [d[0] for d in cur.description]
                activity = [dict(zip(cols, r)) for r in cur.fetchall()]
        except Exception: pass
        add_h("ACTIVITY LOG", level=1)
        if activity:
            for a in activity[:20]:
                add_para(f"{a.get('created_at','')}  ·  {a.get('action','')}  ·  {a.get('details','')}",
                         size=9, color=MUTED)
        else:
            add_para("No activity recorded yet.", size=10, italic=True, color=MUTED)
        add_divider()

        # --- PIPELINE TRACE ---
        add_h("PIPELINE TRACE", level=1)
        if pipeline:
            total_cost = sum(float(p.get("cost_usd") or 0) for p in pipeline)
            total_lat = sum(int(p.get("latency_ms") or 0) for p in pipeline)
            add_kv("Total cost", f"${total_cost:.4f}")
            add_kv("Total latency", f"{total_lat/1000:.2f}s")
            for s in pipeline[:15]:
                add_kv(s.get("step_name") or "?",
                       f"{s.get('status','—')}  ·  {s.get('model') or ''}  ·  {(s.get('latency_ms') or 0)}ms")
        else:
            add_para("No pipeline runs recorded.", size=10, italic=True, color=MUTED)

        # Footer
        foot = doc.add_paragraph()
        fr = foot.add_run(f"\nGenerated by HIRE Talent Intelligence  ·  Candidate ID #{candidate_id}")
        style_run(fr, size=8, color=MUTED)
        foot.alignment = WD_ALIGN_PARAGRAPH.CENTER

        buf = BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf

    import asyncio
    from fastapi.responses import StreamingResponse
    loop = asyncio.get_event_loop()
    try:
        buf = await asyncio.wait_for(loop.run_in_executor(None, _build_docx), timeout=30)
    except asyncio.TimeoutError:
        return JSONResponse(status_code=504, content={"detail": "Report generation timed out after 30s"})
    safe_name = re.sub(r"[^A-Za-z0-9_-]", "_", (candidate.get("name") or f"candidate_{candidate_id}"))
    filename = f"{safe_name}_HIRE_report.docx"
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@router.delete("/{candidate_id}")
async def archive_candidate(candidate_id: str, request: Request, hard: bool = Query(False)):
    """Soft-delete (archive) a candidate, or hard-delete with ?hard=true."""
    user = get_current_user(request)
    from backend.core.database import get_cursor

    candidate_id = _resolve_cid(candidate_id)
    if candidate_id is None:
        raise HTTPException(404, "Candidate not found")

    if hard:
        # Capture file path + creator before delete; enforce creator-OR-superadmin
        file_path = None
        with get_cursor() as cur:
            cur.execute("SELECT pdf_path, owner_id FROM candidates WHERE id = %s", (candidate_id,))
            row = cur.fetchone()
            if not row:
                raise HTTPException(404, "Candidate not found")
            file_path, owner_id = row
            is_creator = owner_id == (user.get("user_id") or 0)
            is_super = user.get("role") == "superadmin"
            if not (is_creator or is_super):
                raise HTTPException(403, "Only the uploader or superadmin can hard-delete this CV")
            cur.execute("DELETE FROM candidates WHERE id = %s", (candidate_id,))
        # Remove file from disk (best-effort)
        if file_path:
            try:
                p = Path(file_path)
                if p.exists() and p.is_file():
                    p.unlink()
            except Exception:
                pass
        return {"message": f"Candidate {candidate_id} hard-deleted", "hard": True}

    # Soft archive — owner or admin+
    with get_cursor() as cur:
        cur.execute("SELECT owner_id FROM candidates WHERE id = %s", (candidate_id,))
        _r = cur.fetchone()
        if not _r:
            raise HTTPException(404, "Candidate not found")
        ensure_owner_or_min(user, _r[0], "admin", request)
        cur.execute(
            "UPDATE candidates SET status = 'archived', updated_at = NOW() WHERE id = %s",
            (candidate_id,),
        )
    return {"message": f"Candidate {candidate_id} archived"}


@router.patch("/{candidate_id}")
async def update_candidate(candidate_id: str, request: Request):
    """Update candidate fields."""
    user = get_current_user(request)
    body = await request.json()
    candidate_id = _resolve_cid(candidate_id)
    if candidate_id is None:
        raise HTTPException(404, "Candidate not found")

    with get_cursor() as _c:
        _c.execute("SELECT owner_id FROM candidates WHERE id = %s", (candidate_id,))
        _r = _c.fetchone()
    if not _r:
        raise HTTPException(404, "Candidate not found")
    ensure_owner_or_min(user, _r[0], "admin", request)

    allowed_fields = {"name", "email", "phone", "location", "tags", "seniority_level", "status"}
    updates = {k: v for k, v in body.items() if k in allowed_fields}

    if not updates:
        raise HTTPException(status_code=400, detail="No valid fields to update")

    # Safe: field names come from allowed_fields whitelist, values are parameterized
    set_clause = ", ".join(f"{k} = %s" for k in updates)
    values = list(updates.values()) + [candidate_id]

    with get_cursor() as cur:
        cur.execute(
            f"UPDATE candidates SET {set_clause}, updated_at = NOW(), "
            f"updated_by = %s, updated_by_id = %s WHERE id = %s",
            values[:-1] + [user.get("user_id"), user.get("user_id"), candidate_id],
        )
    return {"message": "Updated", "fields": list(updates.keys())}


# ---------------------------------------------------------------------------
# LLM extraction helper
# ---------------------------------------------------------------------------
def _extract_candidate_from_text(text: str) -> dict:
    """Use LLM to extract structured candidate data from raw text."""
    prompt = f"""Extract structured candidate data from this LinkedIn profile / resume text. Return ONLY valid JSON with:
{{
  "name": "string",
  "email": "string or null",
  "current_role": "string or null",
  "current_company": "string or null",
  "location": "string or null",
  "total_experience_years": "number or null",
  "seniority_level": "junior|mid|senior|staff|principal|lead|manager|director or null",
  "skills_technical": ["string"],
  "experience": [{{"company": "string", "role": "string", "start_date": "string", "end_date": "string or null", "description": "string"}}],
  "education": [{{"institution": "string", "degree": "string", "field": "string", "year": "string or null"}}],
  "certifications": ["string"],
  "summary_short": "1-2 sentence summary"
}}

Text:
{text[:8000]}"""

    raw = llm_call(prompt, model=CHAT_MODEL, temperature=0.1, max_tokens=4000)
    if not raw:
        raise HTTPException(status_code=500, detail="LLM extraction failed")

    # Parse JSON from response (handle markdown code blocks)
    cleaned = raw.strip()
    if cleaned.startswith("```"):
        cleaned = cleaned.split("\n", 1)[1] if "\n" in cleaned else cleaned[3:]
        if cleaned.endswith("```"):
            cleaned = cleaned[:-3]
    try:
        return json.loads(cleaned)
    except json.JSONDecodeError:
        logger.error(f"LLM returned invalid JSON: {raw[:500]}")
        raise HTTPException(status_code=500, detail="LLM returned invalid JSON")


def _create_candidate_from_extracted(data: dict, source: str) -> int:
    """Create a candidate record from LLM-extracted data."""
    experience = data.get("experience", [])
    education = data.get("education", [])
    certifications = data.get("certifications", [])

    candidate_data = {
        "name": data.get("name") or "Unknown",
        "email": data.get("email"),
        "phone": data.get("phone"),
        "location": data.get("location"),
        "linkedin_url": data.get("linkedin_url"),
        "current_role": data.get("current_role"),
        "current_company": data.get("current_company"),
        "total_experience_years": data.get("total_experience_years"),
        "seniority_level": data.get("seniority_level"),
        "skills_technical": data.get("skills_technical", []),
        "skills_soft": data.get("skills_soft", []),
        "tools": data.get("tools", []),
        "languages": data.get("languages", []),
        "experience": json.dumps(experience) if isinstance(experience, list) else "[]",
        "education": json.dumps(education) if isinstance(education, list) else "[]",
        "certifications": json.dumps(certifications) if isinstance(certifications, list) else "[]",
        "projects": "[]",
        "summary_short": data.get("summary_short"),
        "summary_detailed": data.get("summary_detailed"),
        "raw_text": data.get("_raw_text"),
        "pdf_path": None,
        "file_type": None,
        "file_name": None,
        "page_count": 0,
        "quality_score": 0,
        "tags": [],
        "source": source,
        "status": "active",
        "is_processed": True,
    }
    return insert_candidate(candidate_data)


def _link_to_position(candidate_id: int, position_slug: str):
    """Link candidate to a position if slug provided."""
    try:
        with get_cursor() as cur:
            cur.execute("SELECT id FROM positions WHERE slug = %s", (position_slug,))
            row = cur.fetchone()
            if row:
                position_id = row[0]
                cur.execute("""
                    INSERT INTO position_candidates (position_id, candidate_id, stage)
                    VALUES (%s, %s, 'uploaded')
                    ON CONFLICT (position_id, candidate_id) DO NOTHING
                """, (position_id, candidate_id))
    except Exception as e:
        logger.warning(f"Failed to link candidate {candidate_id} to position {position_slug}: {e}")


# ---------------------------------------------------------------------------
# LinkedIn Import
# ---------------------------------------------------------------------------
@router.post("/import-linkedin", dependencies=[Depends(require_role("recruiter"))])
async def import_linkedin(request: Request):
    """Import candidate from LinkedIn URL or pasted profile text."""
    user = get_current_user(request)
    body = await request.json()

    linkedin_url = body.get("linkedin_url")
    linkedin_text = body.get("linkedin_text")
    position_slug = body.get("position_slug")

    profile_text = None

    # Try fetching from URL first
    if linkedin_url:
        try:
            async with httpx.AsyncClient(timeout=15, follow_redirects=True) as client:
                resp = await client.get(linkedin_url, headers={
                    "User-Agent": "Mozilla/5.0 (compatible; HireBot/1.0)",
                })
                if resp.status_code == 200 and len(resp.text) > 500:
                    profile_text = resp.text
                else:
                    logger.info(f"LinkedIn URL returned {resp.status_code}, length={len(resp.text)}")
        except Exception as e:
            logger.warning(f"Failed to fetch LinkedIn URL: {e}")

    # Fall back to pasted text
    if not profile_text:
        if linkedin_text:
            profile_text = linkedin_text
        elif linkedin_url:
            # Use LLM to work with whatever we got, or note the URL
            profile_text = f"LinkedIn profile URL: {linkedin_url}\n(Unable to fetch page content)"
        else:
            raise HTTPException(status_code=400, detail="Provide linkedin_url or linkedin_text")

    # Extract structured data via LLM
    extracted = _extract_candidate_from_text(profile_text)
    extracted["linkedin_url"] = linkedin_url
    extracted["_raw_text"] = profile_text[:10000]

    # Create candidate
    candidate_id = _create_candidate_from_extracted(extracted, source="linkedin")

    # Link to position if provided
    if position_slug:
        _link_to_position(candidate_id, position_slug)

    logger.info(f"LinkedIn import: candidate_id={candidate_id}, name={extracted.get('name')}")
    return {
        "candidate_id": candidate_id,
        "name": extracted.get("name"),
        "source": "linkedin",
        "message": "Candidate imported from LinkedIn",
    }


# ---------------------------------------------------------------------------
# GitHub Profile Analysis
# ---------------------------------------------------------------------------
@router.post("/{candidate_id}/analyze-github")
async def analyze_github(candidate_id: str, request: Request):
    """Analyze a GitHub profile and attach analysis to candidate."""
    user = get_current_user(request)
    _cid_int = _resolve_cid(candidate_id)
    if _cid_int is not None:
        with get_cursor() as _c:
            _c.execute("SELECT owner_id FROM candidates WHERE id = %s", (_cid_int,))
            _r = _c.fetchone()
        if _r is not None:
            ensure_owner_or_min(user, _r[0], "admin", request)
    body = await request.json()
    github_url = body.get("github_url")
    if not github_url:
        raise HTTPException(status_code=400, detail="github_url is required")

    candidate_id = _resolve_cid(candidate_id)
    if candidate_id is None:
        raise HTTPException(status_code=404, detail="Candidate not found")
    candidate = get_candidate(candidate_id)
    if not candidate:
        raise HTTPException(status_code=404, detail="Candidate not found")

    # Extract username from URL
    username = github_url.rstrip("/").split("/")[-1]
    if not username:
        raise HTTPException(status_code=400, detail="Invalid GitHub URL")

    # Fetch GitHub data
    async with httpx.AsyncClient(timeout=15) as client:
        try:
            user_resp = await client.get(
                f"https://api.github.com/users/{username}",
                headers={"Accept": "application/vnd.github.v3+json"},
            )
            if user_resp.status_code != 200:
                raise HTTPException(status_code=400, detail=f"GitHub user not found: {username}")
            gh_user = user_resp.json()

            repos_resp = await client.get(
                f"https://api.github.com/users/{username}/repos?sort=updated&per_page=30",
                headers={"Accept": "application/vnd.github.v3+json"},
            )
            gh_repos = repos_resp.json() if repos_resp.status_code == 200 else []
        except httpx.HTTPError as e:
            raise HTTPException(status_code=502, detail=f"GitHub API error: {e}")

    # Analyze languages
    languages = {}
    top_repos = []
    for repo in gh_repos:
        if repo.get("fork"):
            continue
        lang = repo.get("language")
        if lang:
            languages[lang] = languages.get(lang, 0) + 1
        top_repos.append({
            "name": repo.get("name"),
            "stars": repo.get("stargazers_count", 0),
            "description": repo.get("description") or "",
            "language": lang,
        })

    # Sort top repos by stars
    top_repos.sort(key=lambda r: r["stars"], reverse=True)
    top_repos = top_repos[:10]

    # Sort languages by count
    total_lang = sum(languages.values()) or 1
    language_pcts = {k: round(v / total_lang * 100) for k, v in
                     sorted(languages.items(), key=lambda x: -x[1])}

    # Build analysis summary
    analysis_data = {
        "username": username,
        "public_repos": gh_user.get("public_repos", 0),
        "followers": gh_user.get("followers", 0),
        "following": gh_user.get("following", 0),
        "languages": language_pcts,
        "top_repos": top_repos,
        "bio": gh_user.get("bio"),
        "company": gh_user.get("company"),
        "location": gh_user.get("location"),
        "created_at": gh_user.get("created_at"),
    }

    # Use LLM for technical assessment
    assessment_prompt = f"""Analyze this GitHub profile and provide a brief technical assessment (2-3 sentences).

Username: {username}
Bio: {gh_user.get('bio', 'N/A')}
Public repos: {gh_user.get('public_repos', 0)}
Followers: {gh_user.get('followers', 0)}
Languages: {json.dumps(language_pcts)}
Top repos: {json.dumps(top_repos[:5])}
Account created: {gh_user.get('created_at', 'N/A')}

Focus on: technical breadth, open-source activity, language expertise, project quality."""

    assessment = llm_call(assessment_prompt, model=CHAT_MODEL, temperature=0.3, max_tokens=500)
    analysis_data["assessment"] = assessment or "Assessment unavailable"

    # Save as note on candidate
    with get_cursor() as cur:
        cur.execute("""
            INSERT INTO candidate_notes (candidate_id, content, note_type, created_by)
            VALUES (%s, %s, %s, %s)
            RETURNING id
        """, (
            candidate_id,
            json.dumps(analysis_data, indent=2),
            "github_analysis",
            user["user_id"],
        ))
        note_id = cur.fetchone()[0]

        # Update candidate skills with detected languages
        existing_skills = candidate.get("skills_technical") or []
        new_skills = list(set(existing_skills + list(languages.keys())))
        cur.execute("""
            UPDATE candidates SET skills_technical = %s, updated_at = NOW() WHERE id = %s
        """, (new_skills, candidate_id))

    logger.info(f"GitHub analysis: candidate_id={candidate_id}, username={username}")
    return {
        "note_id": note_id,
        "analysis": analysis_data,
        "skills_added": [l for l in languages.keys() if l not in existing_skills],
    }


# ---------------------------------------------------------------------------
# Import from Text (Universal)
# ---------------------------------------------------------------------------
@router.post("/import-text", dependencies=[Depends(require_role("recruiter"))])
async def import_text(request: Request):
    """Import candidate from raw pasted text (LinkedIn, email, website, etc.)."""
    user = get_current_user(request)
    body = await request.json()

    text = body.get("text")
    if not text or len(text.strip()) < 20:
        raise HTTPException(status_code=400, detail="text is required (minimum 20 characters)")

    source = body.get("source", "text_import")
    position_slug = body.get("position_slug")

    # Extract structured data via LLM
    extracted = _extract_candidate_from_text(text)
    extracted["_raw_text"] = text[:10000]

    # Create candidate
    candidate_id = _create_candidate_from_extracted(extracted, source=source)

    # Link to position if provided
    if position_slug:
        _link_to_position(candidate_id, position_slug)

    logger.info(f"Text import: candidate_id={candidate_id}, name={extracted.get('name')}, source={source}")
    return {
        "candidate_id": candidate_id,
        "name": extracted.get("name"),
        "source": source,
        "message": "Candidate imported from text",
    }


# ---------------------------------------------------------------------------
# Share / Visibility
# ---------------------------------------------------------------------------
from pydantic import BaseModel as _BaseModel
import json as _json

class CVShareRequest(_BaseModel):
    visibility: Optional[str] = None  # legacy
    shared_sector: Optional[bool] = None
    shared_global: Optional[bool] = None


@router.post("/{candidate_id}/share")
async def share_candidate(candidate_id: str, body: CVShareRequest, request: Request):
    user = _user_with_sector(get_current_user(request))
    candidate_id = _resolve_cid(candidate_id)
    if candidate_id is None:
        raise HTTPException(404, "Candidate not found")

    if body.shared_sector is None and body.shared_global is None:
        v = (body.visibility or "private").lower()
        if v not in ("private", "sector", "global"):
            raise HTTPException(400, "visibility must be private | sector | global")
        shared_sector = (v == "sector")
        shared_global = (v == "global")
    else:
        shared_sector = bool(body.shared_sector)
        shared_global = bool(body.shared_global)

    from backend.core.database import get_cursor as _gc
    with _gc() as cur:
        cur.execute("SELECT owner_id, sector_id FROM candidates WHERE id = %s", (candidate_id,))
        row = cur.fetchone()
        if not row:
            raise HTTPException(404, "Candidate not found")
        owner_id, cand_sector = row
        is_owner = owner_id == (user.get("user_id") or 0)
        is_admin = user.get("role") in ("admin", "group_hr")

        if shared_global and not is_admin:
            raise HTTPException(403, "Only admin / group_hr can publish to talent pool")
        if not (is_owner or is_admin):
            raise HTTPException(403, "Only the owner or admin can change visibility")

        legacy = "global" if shared_global else ("sector" if shared_sector else "private")
        new_sector = cand_sector or user.get("sector_id")
        cur.execute("""
            UPDATE candidates
               SET visibility = %s, shared_sector = %s, shared_global = %s,
                   sector_id = %s, updated_at = NOW()
             WHERE id = %s
        """, (legacy, shared_sector, shared_global, new_sector, candidate_id))
        cur.execute(
            "INSERT INTO audit_log (user_id, action, resource_type, resource_id, details) VALUES (%s,%s,%s,%s,%s)",
            (user.get("user_id"), "CV_SHARE", "candidates", candidate_id,
             _json.dumps({"shared_sector": shared_sector, "shared_global": shared_global})),
        )
    return {"id": candidate_id, "shared_sector": shared_sector, "shared_global": shared_global, "visibility": legacy}


# ---------------------------------------------------------------------------
# CSV export
# ---------------------------------------------------------------------------
@router.get("/export.csv")
async def export_candidates_csv(
    request: Request,
    search: Optional[str] = Query(None),
    stage: Optional[str] = Query(None),
):
    """Export candidates as CSV. Honors search + stage filters."""
    get_current_user(request)
    from backend.core.database import get_cursor as _gc

    conds = ["c.status = 'active'"]
    params: list = []
    if search:
        conds.append("(c.name ILIKE %s OR c.email ILIKE %s OR c.current_role ILIKE %s)")
        params.extend([f"%{search}%", f"%{search}%", f"%{search}%"])
    join_sql = ""
    if stage:
        join_sql = "JOIN position_candidates pc ON pc.candidate_id = c.id"
        conds.append("pc.stage = %s")
        params.append(stage)
    where = "WHERE " + " AND ".join(conds)

    buf = _io.StringIO()
    w = _csv.writer(buf)
    w.writerow(["public_id", "name", "email", "phone", "current_title",
                "current_company", "location", "years_experience", "source", "created_at"])
    with _gc() as cur:
        cur.execute(f"""
            SELECT DISTINCT c.public_id, c.name, c.email, c.phone, c.current_role,
                   c.current_company, c.location, c.total_experience_years,
                   c.source, c.created_at
            FROM candidates c {join_sql}
            {where}
            ORDER BY c.created_at DESC
        """, params)
        for r in cur.fetchall():
            w.writerow([("" if v is None else v) for v in r])
    return Response(
        content=buf.getvalue(),
        media_type="text/csv",
        headers={"Content-Disposition": 'attachment; filename="candidates.csv"'},
    )
