"""JD Repository — Create, store, manage, and AI-generate job descriptions."""
from __future__ import annotations

import json
import re
import csv as _csv
import io as _io
import logging
from typing import Optional

from fastapi import APIRouter, Depends, Request, HTTPException, Query
from fastapi.responses import Response
from pydantic import BaseModel

from backend.core.auth import get_current_user
from backend.core.database import get_cursor
from backend.core.permissions import (
    has_min_role,
    require_role,
    ensure_owner_or_min,
)
from backend.core.config import llm_call, CHAT_MODEL, LITE_MODEL
from backend.core.ids import gen_public_id, resolve_id_sync

logger = logging.getLogger(__name__)
router = APIRouter()


def _resolve_jd_id(ref) -> int:
    """Resolve int or public_id JD reference to internal int id, or raise 404."""
    rid = resolve_id_sync("jd_repository", ref)
    if rid is None:
        raise HTTPException(status_code=404, detail="JD not found")
    return rid


# ---------------------------------------------------------------------------
# Models
# ---------------------------------------------------------------------------
class JDCreateRequest(BaseModel):
    title: str
    department: str = ""
    seniority_level: str = ""
    employment_type: str = "full-time"
    jd_text: str = ""
    # corporate template fields
    job_code: Optional[str] = None
    business_sector: Optional[str] = None
    grading: Optional[str] = None
    reporting_to: Optional[str] = None
    location: Optional[str] = None
    work_mode: Optional[str] = None
    job_purpose: Optional[str] = None
    preferred_education: Optional[str] = None
    travel_requirement: Optional[str] = None
    physical_conditions: Optional[str] = None
    doc_owner: Optional[str] = None


class JDGenerateRequest(BaseModel):
    title: str
    department: str = ""
    seniority_level: str = ""
    employment_type: str = "full-time"
    bullets: list[str] = []
    tone: str = "professional"
    # corporate template fields
    job_code: Optional[str] = None
    business_sector: Optional[str] = None
    grading: Optional[str] = None
    reporting_to: Optional[str] = None
    location: Optional[str] = None
    work_mode: Optional[str] = None
    travel_requirement: Optional[str] = None
    physical_conditions: Optional[str] = None
    doc_owner: Optional[str] = None


class JDEnhanceRequest(BaseModel):
    jd_id: int


# ---------------------------------------------------------------------------
# CRUD
# ---------------------------------------------------------------------------
def _scope_where(scope: str, user: dict):
    """Return (sql_fragment, params) for visibility scope filter.
    A JD can be shared to sector AND global simultaneously (shared_sector + shared_global)."""
    uid = user.get("user_id") or 1
    sector_id = user.get("sector_id")
    if scope == "mine":
        return "created_by = %s", [uid]
    if scope == "sector":
        if sector_id is None:
            return "(visibility = 'sector' OR shared_sector = TRUE)", []
        return "(visibility = 'sector' OR shared_sector = TRUE) AND sector_id = %s", [sector_id]
    if scope == "global":
        return "(visibility = 'global' OR shared_global = TRUE)", []
    if sector_id is None:
        return "(created_by = %s OR visibility = 'global' OR shared_global = TRUE)", [uid]
    return ("(created_by = %s "
            " OR ((visibility = 'sector' OR shared_sector = TRUE) AND sector_id = %s) "
            " OR visibility = 'global' OR shared_global = TRUE)"), [uid, sector_id]


def _user_with_sector(user: dict) -> dict:
    """Augment current_user with sector_id from DB if not present."""
    if user.get("sector_id") is not None:
        return user
    with get_cursor() as cur:
        cur.execute("SELECT sector_id FROM users WHERE id = %s", (user.get("user_id") or 1,))
        row = cur.fetchone()
        if row:
            user["sector_id"] = row[0]
    return user


SORT_COLS_JD = {
    "created_at": "j.created_at",
    "updated_at": "j.updated_at",
    "title": "j.title",
    "department": "j.department",
    "seniority_level": "j.seniority_level",
}


@router.get("/")
async def list_jds(
    request: Request,
    department: Optional[str] = Query(None),
    seniority: Optional[str] = Query(None),
    status: Optional[str] = Query(None),
    search: Optional[str] = Query(None),
    scope: Optional[str] = Query(None, description="mine | sector | global"),
    created_by: Optional[int] = Query(None),
    modified_by: Optional[int] = Query(None),
    created_after: Optional[str] = Query(None),
    created_before: Optional[str] = Query(None),
    modified_after: Optional[str] = Query(None),
    modified_before: Optional[str] = Query(None),
    sort: str = Query("updated_at"),
    dir: str = Query("desc"),
    skills: Optional[str] = Query(None, description="CSV of skill canonical values (any-of)"),
    depts: Optional[str] = Query(None, description="CSV of department values"),
    locations: Optional[str] = Query(None, description="CSV of location values"),
    employment_types: Optional[str] = Query(None, description="CSV of employment_type values"),
    seniorities: Optional[str] = Query(None, description="CSV of seniority_level values"),
):
    """List JDs with scope, filters, sort."""
    user = _user_with_sector(get_current_user(request))
    conditions, params = [], []

    def add(sql, *p):
        conditions.append(sql); params.extend(p)

    def _csv(v):
        if not v:
            return []
        return [s.strip() for s in v.split(",") if s.strip()]

    if department:    add("j.department = %s", department)
    if seniority:     add("j.seniority_level = %s", seniority)
    if status:        add("j.status = %s", status)
    else:             add("j.status = 'active'")
    if search:
        add("(j.title ILIKE %s OR j.jd_text ILIKE %s)", f"%{search}%", f"%{search}%")
    if created_by:    add("j.created_by = %s", created_by)
    if modified_by:   add("j.updated_by = %s", modified_by)
    if created_after: add("j.created_at >= %s::timestamptz", created_after)
    if created_before:add("j.created_at <  %s::timestamptz", created_before)
    if modified_after:add("j.updated_at >= %s::timestamptz", modified_after)
    if modified_before:add("j.updated_at < %s::timestamptz", modified_before)

    # ── Facet-rail CSV filters (any-of within group) ──
    sk_list = _csv(skills)
    if sk_list:
        # required_skills + nice_to_have_skills are text[] arrays — case-insensitive overlap
        # Build a lowercased list and compare via lower(unnest) ANY pattern.
        # Simplest correct form: ILIKE any of values across both arrays via array_to_string.
        clauses = []
        for v in sk_list:
            clauses.append(
                "(array_to_string(COALESCE(j.required_skills,'{}'::text[]) || "
                "COALESCE(j.nice_to_have_skills,'{}'::text[]),'|') ILIKE %s)"
            )
            params.append(f"%{v}%")
        conditions.append("(" + " OR ".join(clauses) + ")")

    dp_list = _csv(depts)
    if dp_list:
        ph = ",".join(["%s"] * len(dp_list))
        conditions.append(f"LOWER(j.department) IN ({ph})")
        params.extend([v.lower() for v in dp_list])

    loc_list = _csv(locations)
    if loc_list:
        clauses = []
        for v in loc_list:
            clauses.append("j.location ILIKE %s")
            params.append(f"%{v}%")
        conditions.append("(" + " OR ".join(clauses) + ")")

    et_list = _csv(employment_types)
    if et_list:
        ph = ",".join(["%s"] * len(et_list))
        conditions.append(f"LOWER(j.employment_type) IN ({ph})")
        params.extend([v.lower() for v in et_list])

    sn_list = _csv(seniorities)
    if sn_list:
        ph = ",".join(["%s"] * len(sn_list))
        conditions.append(f"LOWER(j.seniority_level) IN ({ph})")
        params.extend([v.lower() for v in sn_list])

    scope_sql, scope_params = _scope_where(scope or "all", user)
    if scope_sql:
        # rewrite scope_sql column refs (already on jd_repository, prefix with j.)
        conditions.append(scope_sql.replace("created_by", "j.created_by")
                                   .replace("visibility", "j.visibility")
                                   .replace("sector_id", "j.sector_id"))
        params.extend(scope_params)

    where = ("WHERE " + " AND ".join(conditions)) if conditions else ""
    sort_col = SORT_COLS_JD.get(sort, "j.updated_at")
    sort_dir = "ASC" if dir.lower() == "asc" else "DESC"

    with get_cursor() as cur:
        cur.execute(f"""
            SELECT j.*,
                   uc.display_name AS created_by_name,
                   um.display_name AS updated_by_name,
                   (j.expires_at IS NOT NULL AND j.expires_at < NOW()) AS is_expired,
                   CASE WHEN j.expires_at IS NULL THEN NULL
                        ELSE EXTRACT(DAY FROM (j.expires_at - NOW()))::int
                   END AS days_until_expiry
            FROM jd_repository j
            LEFT JOIN users uc ON uc.id = j.created_by
            LEFT JOIN users um ON um.id = j.updated_by
            {where}
            ORDER BY {sort_col} {sort_dir} NULLS LAST
        """, params)
        cols = [desc[0] for desc in cur.description]
        jds = [dict(zip(cols, row)) for row in cur.fetchall()]

        # Compute scope counts (active only)
        uid = user.get("user_id") or 1
        sector_id = user.get("sector_id")
        cur.execute("SELECT COUNT(*) FROM jd_repository WHERE status='active' AND created_by = %s", (uid,))
        n_mine = cur.fetchone()[0]
        if sector_id is not None:
            cur.execute(
                "SELECT COUNT(*) FROM jd_repository WHERE status='active' "
                "AND (visibility='sector' OR shared_sector=TRUE) AND sector_id = %s",
                (sector_id,))
        else:
            cur.execute("SELECT COUNT(*) FROM jd_repository WHERE status='active' "
                        "AND (visibility='sector' OR shared_sector=TRUE)")
        n_sector = cur.fetchone()[0]
        cur.execute("SELECT COUNT(*) FROM jd_repository WHERE status='active' "
                    "AND (visibility='global' OR shared_global=TRUE)")
        n_global = cur.fetchone()[0]

    return {
        "jds": jds,
        "total": len(jds),
        "counts": {"mine": n_mine, "sector": n_sector, "global": n_global},
    }


@router.get("/{jd_id}")
async def get_jd(jd_id: str, request: Request):
    """Get a single JD."""
    get_current_user(request)
    jd_id = _resolve_jd_id(jd_id)

    with get_cursor() as cur:
        cur.execute("SELECT * FROM jd_repository WHERE id = %s", (jd_id,))
        row = cur.fetchone()
        if not row:
            raise HTTPException(status_code=404, detail="JD not found")
        cols = [desc[0] for desc in cur.description]
        return dict(zip(cols, row))


@router.post("/", dependencies=[Depends(require_role("recruiter"))])
async def create_jd(body: JDCreateRequest, request: Request):
    """Create a JD manually (paste or write)."""
    user = get_current_user(request)

    if not body.jd_text.strip():
        raise HTTPException(status_code=400, detail="JD text is required")

    # Auto-reformat layout (preserves wording) — paste path
    raw_original = body.jd_text
    final_text = raw_original
    text_source = "raw"
    quality_score = 1.0
    try:
        from backend.agents.jd_reformat import safe_reformat
        import asyncio as _aio
        final_text, text_source, quality_score = await _aio.to_thread(safe_reformat, raw_original)
    except Exception as _e:
        logger.warning("[jd_reformat] paste path failed: %s", _e)

    # Auto-extract requirements (on reformatted text — better signal)
    extracted = await _extract_requirements(final_text)

    # mig 055/056: stamp expires_at + audit
    try:
        from backend.core.expiry import get_default_expiry_days as _gdays
        _exp_days = _gdays("jd")
    except Exception:
        _exp_days = 90
    _uid = user.get("user_id")
    with get_cursor() as cur:
        cur.execute("""
            INSERT INTO jd_repository (
                public_id,
                title, department, seniority_level, employment_type, jd_text,
                jd_text_raw, jd_text_source, reformat_quality_score,
                required_skills, nice_to_have_skills, min_experience_years,
                education_level, industry_keywords, required_certifications,
                tags, created_by,
                job_code, business_sector, grading, reporting_to, location, work_mode,
                job_purpose, preferred_education, travel_requirement, physical_conditions,
                doc_owner, doc_last_review,
                created_by_id, updated_by_id, expires_at
            ) VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s,
                      %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, CURRENT_DATE,
                      %s, %s, NOW() + (%s || ' days')::interval)
            RETURNING id
        """, (
            gen_public_id("jd_repository"),
            body.title, body.department, body.seniority_level, body.employment_type,
            final_text, raw_original, text_source, quality_score,
            extracted.get("required_skills", []),
            extracted.get("nice_to_have", []),
            extracted.get("min_experience_years", 0),
            extracted.get("education_level", ""),
            extracted.get("industry_keywords", []),
            extracted.get("certifications", []),
            [body.department.lower()] if body.department else [],
            _uid,
            body.job_code, body.business_sector, body.grading, body.reporting_to,
            body.location, body.work_mode, body.job_purpose, body.preferred_education,
            body.travel_requirement, body.physical_conditions, body.doc_owner,
            _uid, _uid, _exp_days,
        ))
        jd_id = cur.fetchone()[0]

    try:
        from backend.agents.dedup import propose_jd_dups
        propose_jd_dups(jd_id)
    except Exception as e:
        logger.warning(f"dedup propose_jd_dups failed: {e}")

    try:
        from backend.agents.facet_miner import mine_jd
        fm = mine_jd(jd_id)
        logger.info(f"jd facets {jd_id}: +{fm.get('added',0)} ↑{fm.get('bumped',0)}")
    except Exception as e:
        logger.warning(f"jd facet mining failed: {e}")

    # Auto-extract KF4D competencies from JD text
    try:
        comp_res = _internal_extract_jd_competencies(jd_id)
        logger.info(f"jd competencies {jd_id}: +{comp_res.get('created',0)} ↑{comp_res.get('updated',0)}")
    except Exception as e:
        logger.warning(f"jd competency extraction failed for {jd_id}: {e}")

    return {"jd_id": jd_id, "message": "JD created", "extracted": extracted}


@router.post("/generate", dependencies=[Depends(require_role("recruiter"))])
async def generate_jd(body: JDGenerateRequest, request: Request, preview: bool = Query(False)):
    """AI-generate a JD from title + bullet points. If preview=true, returns generated text WITHOUT persisting."""
    user = get_current_user(request)

    from backend.core.config import is_ai_available
    if not is_ai_available():
        raise HTTPException(
            status_code=503,
            detail="AI features not available. Set OPENROUTER_API_KEY environment variable to enable JD generation. Get a key from https://openrouter.ai/keys"
        )

    prompt = f"""Generate a corporate Job Description matching this exact 6-section template. Output Markdown only. At least 500 words.

INPUT
- Position Title: {body.title}
- Job Code/ID: {body.job_code or '(generate one like ENG-SE-001)'}
- Business Sector/Entity: {body.business_sector or 'Not specified'}
- Department/Function: {body.department or 'Not specified'}
- Grading/Ranking: {body.grading or body.seniority_level or 'Not specified'}
- Reporting To: {body.reporting_to or 'Not specified'}
- Location: {body.location or 'Not specified'}
- Work Mode: {body.work_mode or 'Onsite'}
- Employment Type: {body.employment_type}
- Travel Requirement: {body.travel_requirement or 'Minimal (<10%)'}
- Physical/Specific Conditions: {body.physical_conditions or 'Standard office'}
- Tone: {body.tone}

Hiring Manager Notes:
{chr(10).join('- ' + b for b in body.bullets) if body.bullets else '(none — infer from role)'}

OUTPUT — produce exactly these sections in this order, using these headings verbatim:

## I. Role Information
| Field | Value |
|---|---|
| Position Title | ... |
| Job Code/ID | ... |
| Business Sector/Entity | ... |
| Department/Function | ... |
| Grading/Ranking | ... |
| Reporting To | ... |
| Location | ... |
| Work Mode | ... |
| Employment Type | ... |

## II. Role Scope
**Job Purpose:** (2-4 sentences — why this role exists, what business outcome it owns)

## III. Key Responsibilities
(8-12 measurable bullets starting with strong verbs: Lead, Build, Design, Drive, Analyze. Include scope and impact metrics where possible.)

## IV. Job Requirements
**Minimum Education:** (specific degree)
**Preferred Education:** (advanced/specialty)
**Total Experience:** (e.g. 5+ years; specify domain)
**Technical Skills:**
- (specific tech, frameworks, tools — be precise)

## V. Working Conditions
**Travel Requirement:** ...
**Physical/Specific Conditions:** ...

## VI. Governance
**Equal Opportunity:** We are an equal opportunity employer; all qualified applicants will receive consideration without regard to race, color, religion, sex, sexual orientation, gender identity, national origin, disability, or veteran status.
**Document Control:** Version 1.0 · Owner: {body.doc_owner or '(HR Department)'} · Last Reviewed: (today)

RULES
- Inclusive, gender-neutral language (no "rockstar", "ninja", "guru", "manpower")
- Specific technologies and years, not vague phrasing
- Action-oriented responsibilities
- Output ONLY the Markdown — no preamble or postscript
"""

    jd_text = llm_call(prompt, model=CHAT_MODEL, temperature=0.4, max_tokens=4000)
    if not jd_text:
        raise HTTPException(status_code=500, detail="JD generation failed")

    # Extract requirements
    extracted = await _extract_requirements(jd_text)

    # Auto-pick scoring profile based on department / title
    profile = "engineering"
    title_lower = (body.title or "").lower()
    dept_lower = (body.department or "").lower()
    if "sales" in title_lower or "sales" in dept_lower or "account exec" in title_lower:
        profile = "sales"
    elif "design" in title_lower or "ux" in title_lower or "ui" in title_lower:
        profile = "design"
    elif "product manager" in title_lower or "product owner" in title_lower or "product" == dept_lower:
        profile = "product"
    elif "marketing" in title_lower or "marketing" in dept_lower or "growth" in title_lower:
        profile = "marketing"
    elif "compliance" in title_lower or "risk" in title_lower or "audit" in title_lower or "legal" in dept_lower:
        profile = "compliance"
    elif "doctor" in title_lower or "nurse" in title_lower or "clinical" in title_lower or "healthcare" in dept_lower:
        profile = "healthcare"
    elif "finance" in title_lower or "accountant" in title_lower or "treasur" in title_lower or "finance" in dept_lower:
        profile = "finance"
    p = SCORING_PRESETS[profile]

    if preview:
        return {
            "preview": True,
            "jd_text": jd_text,
            "extracted": extracted,
            "scoring_profile": profile,
            "weights": {
                "weight_skills": p["weight_skills"],
                "weight_experience": p["weight_experience"],
                "weight_industry": p["weight_industry"],
                "weight_education": p["weight_education"],
                "weight_certifications": p["weight_certifications"],
                "weight_culture": p["weight_culture"],
            },
            "message": "JD generated (preview only — not saved)",
        }

    # Save to repository (mig 055/056: stamp expires_at + audit cols)
    try:
        from backend.core.expiry import get_default_expiry_days as _gdays
        _exp_days = _gdays("jd")
    except Exception:
        _exp_days = 90
    _uid = user.get("user_id")
    with get_cursor() as cur:
        cur.execute("""
            INSERT INTO jd_repository (
                public_id,
                title, department, seniority_level, employment_type, jd_text,
                required_skills, nice_to_have_skills, min_experience_years,
                education_level, industry_keywords, required_certifications,
                tags, created_by,
                job_code, business_sector, grading, reporting_to, location, work_mode,
                travel_requirement, physical_conditions, doc_owner, doc_last_review,
                weight_skills, weight_experience, weight_industry, weight_education,
                weight_certifications, weight_culture, scoring_profile,
                created_by_id, updated_by_id, expires_at
            ) VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s,
                      %s, %s, %s, %s, %s, %s, %s, %s, %s, CURRENT_DATE,
                      %s, %s, %s, %s, %s, %s, %s,
                      %s, %s, NOW() + (%s || ' days')::interval)
            RETURNING id
        """, (
            gen_public_id("jd_repository"),
            body.title, body.department, body.seniority_level, body.employment_type,
            jd_text,
            extracted.get("required_skills", []),
            extracted.get("nice_to_have", []),
            extracted.get("min_experience_years", 0),
            extracted.get("education_level", ""),
            extracted.get("industry_keywords", []),
            extracted.get("certifications", []),
            [body.department.lower()] if body.department else [],
            _uid,
            body.job_code, body.business_sector, body.grading, body.reporting_to,
            body.location, body.work_mode, body.travel_requirement,
            body.physical_conditions, body.doc_owner,
            p["weight_skills"], p["weight_experience"], p["weight_industry"],
            p["weight_education"], p["weight_certifications"], p["weight_culture"],
            profile,
            _uid, _uid, _exp_days,
        ))
        jd_id = cur.fetchone()[0]

    try:
        from backend.agents.facet_miner import mine_jd
        mine_jd(jd_id)
    except Exception as e:
        logger.warning(f"jd facet mining (generate) failed: {e}")

    try:
        comp_res = _internal_extract_jd_competencies(jd_id)
        logger.info(f"jd competencies (generate) {jd_id}: +{comp_res.get('created',0)}")
    except Exception as e:
        logger.warning(f"jd competency extraction (generate) failed for {jd_id}: {e}")

    return {
        "jd_id": jd_id,
        "jd_text": jd_text,
        "extracted": extracted,
        "message": "JD generated and saved",
    }


@router.post("/{jd_id}/enhance")
async def enhance_jd(jd_id: str, request: Request, preview: bool = Query(False)):
    """AI-enhance an existing JD. If preview=true, returns proposed changes WITHOUT writing to DB."""
    user = get_current_user(request)
    jd_id = _resolve_jd_id(jd_id)
    if not preview:
        with get_cursor() as _c:
            _c.execute("SELECT created_by FROM jd_repository WHERE id = %s", (jd_id,))
            _r = _c.fetchone()
        if _r is not None:
            ensure_owner_or_min(user, _r[0], "admin", request)

    from backend.core.config import is_ai_available
    if not is_ai_available():
        raise HTTPException(
            status_code=503,
            detail="AI features not available. Set OPENROUTER_API_KEY environment variable to enable JD enhancement. Get a key from https://openrouter.ai/keys"
        )

    with get_cursor() as cur:
        cur.execute("SELECT jd_text FROM jd_repository WHERE id = %s", (jd_id,))
        row = cur.fetchone()
        if not row:
            raise HTTPException(status_code=404, detail="JD not found")
        jd_text = row[0]

    prompt = f"""You are an expert HR consultant. Review and significantly enhance this job description.

TASKS:
1. **DEI Language Audit**: Find and replace all non-inclusive terms (gendered language, ableist terms, age-biased requirements, exclusionary phrases). Score inclusivity 0-100.
2. **Legal Compliance**: Check for discriminatory requirements (age limits, unnecessary physical requirements, nationality restrictions). Pass/Fail.
3. **Completeness Check**: Ensure these sections exist: About the Role, Responsibilities (8+), Requirements (6+), Nice to Have, What We Offer, Technical Skills. Score 0-100.
4. **Quality Enhancement**:
   - Make vague requirements specific ("programming experience" → "5+ years Python/Django")
   - Add missing context about team and impact
   - Ensure responsibilities are action-oriented (verb-first)
   - Add measurable outcomes where possible
   - Improve the overall appeal and clarity
5. **Add missing sections** if not present — do NOT just return the same text

Current JD (enhance this — keep all existing content but improve it significantly):
{jd_text[:12000]}

Return the COMPLETE enhanced JD in Markdown (not truncated — include ALL sections).

Then at the very end, add:
---
## Enhancement Report
- DEI Score: X/100
- Legal: PASS/FAIL
- Completeness: X/100
- Changes Made:
  - [specific change 1]
  - [specific change 2]
  - ..."""

    enhanced = llm_call(prompt, model=CHAT_MODEL, temperature=0.3, max_tokens=5000)
    if not enhanced:
        raise HTTPException(status_code=500, detail="Enhancement failed")

    # Parse scores from the enhancement report
    compliance = _parse_compliance(enhanced)

    # Re-extract requirements
    extracted = await _extract_requirements(enhanced)

    from backend.agents.ingest import _jd_content_hash, _clean_jd_title
    new_title = (extracted.get("title") or "").strip() or _clean_jd_title(enhanced, None)
    new_hash = _jd_content_hash(enhanced)

    if preview:
        # Don't write to DB — return draft so frontend can show + let user SAVE manually
        return {
            "jd_id": jd_id,
            "preview": True,
            "draft": {
                "jd_text": enhanced,
                "title": new_title,
                "department": extracted.get("department", "") or "",
                "seniority_level": extracted.get("seniority_level", "") or "",
                "required_skills": extracted.get("required_skills", []),
                "nice_to_have_skills": extracted.get("nice_to_have", []),
                "min_experience_years": int(extracted.get("min_experience_years") or 0),
                "education_level": extracted.get("education_level", "") or "",
                "industry_keywords": extracted.get("industry_keywords", []),
                "required_certifications": extracted.get("certifications", []),
                "dei_score": compliance.get("dei_score"),
                "completeness_score": compliance.get("completeness"),
            },
            "compliance": compliance,
            "extracted": extracted,
            "message": "Preview ready — click SAVE to commit",
        }

    with get_cursor() as cur:
        cur.execute("""
            UPDATE jd_repository SET
                title = COALESCE(NULLIF(%s,''), title),
                department = COALESCE(NULLIF(%s,''), department),
                seniority_level = COALESCE(NULLIF(%s,''), seniority_level),
                jd_text = %s, jd_enhanced = TRUE,
                required_skills = %s, nice_to_have_skills = %s,
                min_experience_years = %s, education_level = %s,
                industry_keywords = %s, required_certifications = %s,
                dei_score = %s, legal_check = %s, completeness_score = %s,
                jd_content_hash = %s,
                updated_at = NOW()
            WHERE id = %s
        """, (
            new_title,
            extracted.get("department", ""),
            extracted.get("seniority_level", ""),
            enhanced,
            extracted.get("required_skills", []),
            extracted.get("nice_to_have", []),
            extracted.get("min_experience_years", 0),
            extracted.get("education_level", ""),
            extracted.get("industry_keywords", []),
            extracted.get("certifications", []),
            compliance.get("dei_score"),
            json.dumps({"pass": compliance.get("legal_pass"), "issues": compliance.get("legal_issues", [])}),
            compliance.get("completeness"),
            new_hash,
            jd_id,
        ))

    try:
        from backend.agents.facet_miner import mine_jd
        mine_jd(jd_id)
    except Exception as e:
        logger.warning(f"jd facet mining (enhance) failed: {e}")

    try:
        comp_res = _internal_extract_jd_competencies(jd_id)
        logger.info(f"jd competencies (enhance) {jd_id}: +{comp_res.get('created',0)} ↑{comp_res.get('updated',0)}")
    except Exception as e:
        logger.warning(f"jd competency extraction (enhance) failed for {jd_id}: {e}")

    return {
        "jd_id": jd_id,
        "jd_text": enhanced,
        "compliance": compliance,
        "extracted": extracted,
        "message": "JD enhanced",
    }


@router.post("/{jd_id}/duplicate", dependencies=[Depends(require_role("recruiter"))])
async def duplicate_jd(jd_id: str, request: Request):
    """Duplicate a JD as a starting point for a new one."""
    user = get_current_user(request)
    jd_id = _resolve_jd_id(jd_id)

    with get_cursor() as cur:
        cur.execute("SELECT * FROM jd_repository WHERE id = %s", (jd_id,))
        row = cur.fetchone()
        if not row:
            raise HTTPException(status_code=404, detail="JD not found")
        cols = [desc[0] for desc in cur.description]
        original = dict(zip(cols, row))

    try:
        from backend.core.expiry import get_default_expiry_days as _gdays
        _exp_days = _gdays("jd")
    except Exception:
        _exp_days = 90
    _uid = user.get("user_id")
    with get_cursor() as cur:
        cur.execute("""
            INSERT INTO jd_repository (
                public_id,
                title, department, seniority_level, employment_type, jd_text,
                required_skills, nice_to_have_skills, min_experience_years,
                education_level, industry_keywords, required_certifications,
                tags, created_by,
                created_by_id, updated_by_id, expires_at
            ) VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s,
                      %s, %s, NOW() + (%s || ' days')::interval)
            RETURNING id
        """, (
            gen_public_id("jd_repository"),
            f"{original['title']} (Copy)",
            original.get("department"),
            original.get("seniority_level"),
            original.get("employment_type"),
            original.get("jd_text"),
            original.get("required_skills", []),
            original.get("nice_to_have_skills", []),
            original.get("min_experience_years", 0),
            original.get("education_level"),
            original.get("industry_keywords", []),
            original.get("required_certifications", []),
            original.get("tags", []),
            _uid,
            _uid, _uid, _exp_days,
        ))
        new_id = cur.fetchone()[0]

    return {"jd_id": new_id, "message": "JD duplicated"}


@router.post("/{jd_id}/use", dependencies=[Depends(require_role("recruiter"))])
async def use_jd_for_position(jd_id: str, request: Request):
    """Link a JD from the repo to a position (creates position if slug provided)."""
    user = get_current_user(request)
    jd_id = _resolve_jd_id(jd_id)
    body = await request.json()
    position_slug = body.get("position_slug")

    with get_cursor() as cur:
        cur.execute("SELECT * FROM jd_repository WHERE id = %s", (jd_id,))
        row = cur.fetchone()
        if not row:
            raise HTTPException(status_code=404, detail="JD not found")
        cols = [desc[0] for desc in cur.description]
        jd = dict(zip(cols, row))

    if not position_slug:
        raise HTTPException(status_code=400, detail="position_slug required")

    # Update position with JD (incl. weights, unless position has overridden weights)
    with get_cursor() as cur:
        cur.execute("""
            UPDATE positions SET
                jd_text = %s,
                required_skills = %s, nice_to_have_skills = %s,
                min_experience_years = %s, education_level = %s,
                industry_keywords = %s, required_certifications = %s,
                weight_skills = CASE WHEN COALESCE(weights_overridden, FALSE) THEN weight_skills ELSE %s END,
                weight_experience = CASE WHEN COALESCE(weights_overridden, FALSE) THEN weight_experience ELSE %s END,
                weight_industry = CASE WHEN COALESCE(weights_overridden, FALSE) THEN weight_industry ELSE %s END,
                weight_education = CASE WHEN COALESCE(weights_overridden, FALSE) THEN weight_education ELSE %s END,
                weight_certifications = CASE WHEN COALESCE(weights_overridden, FALSE) THEN weight_certifications ELSE %s END,
                weight_culture = CASE WHEN COALESCE(weights_overridden, FALSE) THEN weight_culture ELSE %s END,
                knockout_threshold = CASE WHEN COALESCE(weights_overridden, FALSE) THEN knockout_threshold ELSE %s END,
                scoring_profile = COALESCE(scoring_profile, %s),
                weights_source_jd_id = %s,
                updated_at = NOW()
            WHERE slug = %s
        """, (
            jd["jd_text"],
            jd.get("required_skills", []),
            jd.get("nice_to_have_skills", []),
            jd.get("min_experience_years", 0),
            jd.get("education_level"),
            jd.get("industry_keywords", []),
            jd.get("required_certifications", []),
            jd.get("weight_skills", 40),
            jd.get("weight_experience", 25),
            jd.get("weight_industry", 15),
            jd.get("weight_education", 10),
            jd.get("weight_certifications", 10),
            jd.get("weight_culture", 0),
            jd.get("knockout_threshold", 0),
            jd.get("scoring_profile", "engineering"),
            jd_id,
            position_slug,
        ))

    # Increment used_count
    with get_cursor() as cur:
        cur.execute("UPDATE jd_repository SET used_count = used_count + 1 WHERE id = %s", (jd_id,))

    # Sync competencies from JD → position_competencies
    try:
        with get_cursor() as cur:
            cur.execute("SELECT id FROM positions WHERE slug = %s", (position_slug,))
            prow = cur.fetchone()
        if prow:
            _sync_position_competencies_from_jd(jd_id, prow[0])
    except Exception as e:
        logger.warning(f"sync position competencies from jd failed: {e}")

    try:
        from backend.agents.dedup import propose_jd_dups
        propose_jd_dups(jd_id)
    except Exception as e:
        logger.warning(f"dedup propose_jd_dups failed: {e}")

    # Trigger embedding + auto-scan so AI tab populates after JD attach.
    # Without this, position has jd_text but never embedded/scanned, so the
    # AI tab shows "NO JD ON POSITION" / "no scan ever ran".
    try:
        import asyncio as _asyncio
        from backend.routes.positions import (
            _extract_and_embed_jd,
            auto_scan_for_position,
            _create_ai_scan_row,
            _safe_background,
        )
        with get_cursor() as cur:
            cur.execute("SELECT id FROM positions WHERE slug = %s", (position_slug,))
            prow = cur.fetchone()
        if prow and jd.get("jd_text"):
            pid = prow[0]
            try:
                await _extract_and_embed_jd(position_slug, jd["jd_text"])
            except Exception as e:
                logger.warning(f"[use_jd] extract+embed failed: {e}")
            try:
                _scan_id = _create_ai_scan_row(pid)
                _asyncio.create_task(_safe_background(
                    auto_scan_for_position(pid, position_slug, scan_id=_scan_id,
                                           match_source="auto_scan_on_jd_update"),
                    "auto_scan_on_jd_attach",
                ))
            except Exception as e:
                logger.warning(f"[use_jd] auto_scan trigger failed: {e}")
    except Exception as e:
        logger.warning(f"[use_jd] post-attach hooks failed: {e}")

    return {"message": f"JD applied to position {position_slug}"}


@router.post("/{jd_id}/upload-body")
async def upload_jd_body(jd_id: str, request: Request):
    """Replace JD body text from an uploaded file (no AI). Accepts pdf/docx/txt."""
    from fastapi import UploadFile, File
    from pathlib import Path as _P
    from backend.core.cv_parser import extract_pdf, extract_docx
    user = get_current_user(request)
    jd_id = _resolve_jd_id(jd_id)
    with get_cursor() as _c:
        _c.execute("SELECT created_by FROM jd_repository WHERE id = %s", (jd_id,))
        _r = _c.fetchone()
    if _r is not None:
        ensure_owner_or_min(user, _r[0], "admin", request)

    form = await request.form()
    f = form.get("file")
    if not f:
        raise HTTPException(status_code=400, detail="file required")
    ext = _P(f.filename or "").suffix.lower()
    content = await f.read()
    if len(content) > 20 * 1024 * 1024:
        raise HTTPException(status_code=413, detail="file too large (max 20MB)")

    import os, uuid as _u
    DATA = _P(os.getenv("DATA_DIR", "/data"))
    JD_DIR = DATA / "jds"
    JD_DIR.mkdir(parents=True, exist_ok=True)
    save_path = JD_DIR / f"{_u.uuid4()}{ext}"
    save_path.write_bytes(content)

    if ext == ".docx":
        text = extract_docx(str(save_path)).get("text", "")
    elif ext == ".pdf":
        text = extract_pdf(str(save_path)).get("text", "")
    elif ext in {".txt", ".md"}:
        text = save_path.read_text(encoding="utf-8", errors="ignore")
    else:
        raise HTTPException(status_code=400, detail=f"unsupported type {ext}")

    if not text.strip():
        raise HTTPException(status_code=400, detail="extracted text empty")

    # Auto-reformat layout (preserves wording)
    raw_original = text
    final_text = text
    text_source = "extracted"
    quality_score = 1.0
    try:
        from backend.agents.jd_reformat import safe_reformat
        import asyncio as _aio
        final_text, text_source, quality_score = await _aio.to_thread(safe_reformat, raw_original)
    except Exception as _e:
        logger.warning("[jd_reformat] upload-body path failed: %s", _e)

    from backend.agents.ingest import _jd_content_hash
    chash = _jd_content_hash(final_text)
    with get_cursor() as cur:
        cur.execute(
            "UPDATE jd_repository SET jd_text=%s, jd_text_raw=%s, jd_text_source=%s, reformat_quality_score=%s, source_file_path=%s, source_file_name=%s, source_file_type=%s, jd_content_hash=%s, updated_at=NOW() WHERE id=%s RETURNING id",
            (final_text, raw_original, text_source, quality_score, str(save_path), f.filename, ext.lstrip("."), chash, jd_id),
        )
        if not cur.fetchone():
            raise HTTPException(status_code=404, detail="JD not found")

    return {"jd_id": jd_id, "chars": len(final_text), "filename": f.filename, "source": text_source, "score": quality_score, "message": "JD body replaced"}


@router.post("/{jd_id}/reformat")
async def manual_reformat_jd(jd_id: str, request: Request, force: bool = Query(False)):
    """Manually trigger AI reformat (layout-only, preserves wording).

    Idempotent: skips if score already above threshold unless force=true.
    Stores original in jd_text_raw, sets jd_text_source='ai_reformatted'.
    """
    user = get_current_user(request)
    jd_id = _resolve_jd_id(jd_id)
    with get_cursor() as cur:
        cur.execute("SELECT created_by, jd_text, jd_text_raw FROM jd_repository WHERE id=%s", (jd_id,))
        row = cur.fetchone()
    if not row:
        raise HTTPException(status_code=404, detail="JD not found")
    ensure_owner_or_min(user, row[0], "admin", request)
    current_text = row[1] or ""
    raw_backup = row[2] or current_text  # keep first raw if already set
    if not current_text.strip():
        raise HTTPException(status_code=400, detail="JD text is empty")

    from backend.agents.jd_reformat import safe_reformat, reformat_text, score_structure, MIN_RETENTION
    import asyncio as _aio

    if force:
        new_text = await _aio.to_thread(reformat_text, current_text)
        if not new_text:
            raise HTTPException(status_code=502, detail="LLM reformat failed")
        retention = len(new_text.split()) / max(len(current_text.split()), 1)
        if retention < MIN_RETENTION:
            raise HTTPException(status_code=422, detail=f"reformat dropped too much content (retention {retention:.2f} < {MIN_RETENTION})")
        new_score = score_structure(new_text)
        text_source = "ai_reformatted"
    else:
        new_text, text_source, new_score = await _aio.to_thread(safe_reformat, current_text)
        if text_source != "ai_reformatted":
            return {"jd_id": jd_id, "skipped": True, "reason": "structure score already acceptable", "score": new_score}

    from backend.agents.ingest import _jd_content_hash, _embed_jd, _safe
    chash = _jd_content_hash(new_text)
    with get_cursor() as cur:
        cur.execute(
            "UPDATE jd_repository SET jd_text=%s, jd_text_raw=%s, jd_text_source=%s, reformat_quality_score=%s, jd_content_hash=%s, updated_at=NOW() WHERE id=%s RETURNING title",
            (new_text, raw_backup, text_source, new_score, chash, jd_id),
        )
        title = cur.fetchone()[0]
    _aio.create_task(_safe(_embed_jd(jd_id, title or "", new_text), f"jd-reembed {jd_id}"))
    return {"jd_id": jd_id, "source": text_source, "score": new_score, "chars": len(new_text), "message": "JD reformatted"}


@router.post("/{jd_id}/revert-raw")
async def revert_jd_raw(jd_id: str, request: Request):
    """Restore JD body to original pre-reformat text (jd_text_raw)."""
    user = get_current_user(request)
    jd_id = _resolve_jd_id(jd_id)
    with get_cursor() as cur:
        cur.execute("SELECT created_by, jd_text_raw FROM jd_repository WHERE id=%s", (jd_id,))
        row = cur.fetchone()
    if not row:
        raise HTTPException(status_code=404, detail="JD not found")
    ensure_owner_or_min(user, row[0], "admin", request)
    raw = row[1]
    if not raw:
        raise HTTPException(status_code=400, detail="no raw backup available — JD was never reformatted")
    from backend.agents.ingest import _jd_content_hash, _embed_jd, _safe
    import asyncio as _aio
    chash = _jd_content_hash(raw)
    with get_cursor() as cur:
        cur.execute(
            "UPDATE jd_repository SET jd_text=%s, jd_text_source='raw', reformat_quality_score=NULL, jd_content_hash=%s, updated_at=NOW() WHERE id=%s RETURNING title",
            (raw, chash, jd_id),
        )
        title = cur.fetchone()[0]
    _aio.create_task(_safe(_embed_jd(jd_id, title or "", raw), f"jd-reembed {jd_id}"))
    return {"jd_id": jd_id, "message": "JD reverted to raw original"}


class JDBodyUpdate(BaseModel):
    jd_text: str


class JDFullUpdate(BaseModel):
    jd_text: Optional[str] = None
    title: Optional[str] = None
    department: Optional[str] = None
    seniority_level: Optional[str] = None
    employment_type: Optional[str] = None
    required_skills: Optional[list[str]] = None
    nice_to_have_skills: Optional[list[str]] = None
    min_experience_years: Optional[int] = None
    education_level: Optional[str] = None
    industry_keywords: Optional[list[str]] = None
    required_certifications: Optional[list[str]] = None
    dei_score: Optional[float] = None
    completeness_score: Optional[float] = None


@router.patch("/{jd_id}")
async def update_jd(jd_id: str, body: JDFullUpdate, request: Request):
    """Save any subset of JD fields. Used by 'Save' button after AI Enhance preview."""
    user = get_current_user(request)
    jd_id = _resolve_jd_id(jd_id)
    with get_cursor() as _c:
        _c.execute("SELECT created_by FROM jd_repository WHERE id = %s", (jd_id,))
        _r = _c.fetchone()
    if _r is not None:
        ensure_owner_or_min(user, _r[0], "admin", request)
    fields = body.model_dump(exclude_none=True)
    if not fields:
        return {"jd_id": jd_id, "message": "no changes"}
    sets = []
    params = []
    for k, v in fields.items():
        sets.append(f"{k}=%s")
        params.append(v)
    if "jd_text" in fields:
        from backend.agents.ingest import _jd_content_hash
        sets.append("jd_content_hash=%s")
        params.append(_jd_content_hash(fields["jd_text"]))
    sets.append("updated_at=NOW()")
    sets.append("updated_by=%s")
    params.append(user.get("user_id"))
    sets.append("updated_by_id=%s")
    params.append(user.get("user_id"))
    params.append(jd_id)
    with get_cursor() as cur:
        cur.execute(f"UPDATE jd_repository SET {', '.join(sets)} WHERE id=%s RETURNING id, title, jd_text", params)
        row = cur.fetchone()
        if not row:
            raise HTTPException(status_code=404, detail="JD not found")
    if "jd_text" in fields or "title" in fields:
        import asyncio as _aio
        from backend.agents.ingest import _embed_jd, _safe
        _aio.create_task(_safe(_embed_jd(row[0], row[1] or "", row[2] or ""), f"jd-reembed {jd_id}"))
    return {"jd_id": jd_id, "saved_fields": list(fields.keys()), "message": "JD saved"}


@router.patch("/{jd_id}/body")
async def update_jd_body(jd_id: str, body: JDBodyUpdate, request: Request):
    """Save manual edits to JD body text. Recomputes content hash."""
    user = get_current_user(request)
    jd_id = _resolve_jd_id(jd_id)
    with get_cursor() as _c:
        _c.execute("SELECT created_by FROM jd_repository WHERE id = %s", (jd_id,))
        _r = _c.fetchone()
    if _r is not None:
        ensure_owner_or_min(user, _r[0], "admin", request)
    if not body.jd_text.strip():
        raise HTTPException(status_code=400, detail="jd_text empty")
    from backend.agents.ingest import _jd_content_hash
    chash = _jd_content_hash(body.jd_text)
    with get_cursor() as cur:
        cur.execute(
            "UPDATE jd_repository SET jd_text=%s, jd_content_hash=%s, updated_at=NOW(), "
            "updated_by=%s, updated_by_id=%s WHERE id=%s RETURNING id, title",
            (body.jd_text, chash, user.get("user_id"), user.get("user_id"), jd_id),
        )
        row = cur.fetchone()
        if not row:
            raise HTTPException(status_code=404, detail="JD not found")
    import asyncio as _aio
    from backend.agents.ingest import _embed_jd, _safe
    _aio.create_task(_safe(_embed_jd(row[0], row[1] or "", body.jd_text), f"jd-reembed {jd_id}"))
    return {"jd_id": jd_id, "chars": len(body.jd_text), "message": "JD body saved"}


@router.post(
    "/{jd_id}/enrich",
    summary="Re-extract structured fields from JD text",
    description="Runs the LLM extractor on jd_text and updates department, seniority_level, "
                "min_experience_years, required_skills, certifications. Does NOT rewrite jd_text. "
                "Use for stale rows where the background agent hasn't filled fields yet.",
)
async def enrich_jd_fields(jd_id: str, request: Request, force: bool = Query(False)):
    """Manual on-demand enrichment. Cheap (single LLM call). Owner or admin only."""
    user = get_current_user(request)
    jd_id = _resolve_jd_id(jd_id)
    with get_cursor() as cur:
        cur.execute(
            "SELECT created_by, jd_text, department, seniority_level, min_experience_years, "
            "required_skills, nice_to_have_skills FROM jd_repository WHERE id=%s",
            (jd_id,),
        )
        row = cur.fetchone()
    if not row:
        raise HTTPException(status_code=404, detail="JD not found")
    ensure_owner_or_min(user, row[0], "admin", request)
    jd_text = row[1] or ""
    if not jd_text.strip():
        raise HTTPException(status_code=400, detail="JD has no text to enrich")

    extracted = await _extract_requirements(jd_text)
    if not extracted:
        raise HTTPException(status_code=502, detail="LLM extraction returned empty")

    sets, vals, updated_cols = [], [], []
    def _maybe(col, new_val, current_val, is_array=False):
        if new_val is None:
            return
        if not force:
            empty = (current_val is None) or (current_val == "") or (is_array and not current_val) or (col == "min_experience_years" and (current_val == 0 or current_val is None))
            if not empty:
                return
        sets.append(f"{col}=%s")
        vals.append(new_val)
        updated_cols.append(col)
    _maybe("department", extracted.get("department"), row[2])
    _maybe("seniority_level", extracted.get("seniority_level"), row[3])
    _maybe("min_experience_years", extracted.get("min_experience_years"), row[4])
    _maybe("required_skills", extracted.get("required_skills"), row[5], is_array=True)
    _maybe("nice_to_have_skills", extracted.get("nice_to_have"), row[6], is_array=True)

    if not sets:
        return {"jd_id": jd_id, "updated": [], "extracted": extracted, "message": "All fields already filled"}

    sets.append("updated_at=NOW()")
    vals.append(jd_id)
    sql = f"UPDATE jd_repository SET {', '.join(sets)} WHERE id=%s RETURNING department, seniority_level, min_experience_years, required_skills, nice_to_have_skills"
    with get_cursor() as cur:
        cur.execute(sql, tuple(vals))
        updated_row = cur.fetchone()
    cols = ["department", "seniority_level", "min_experience_years", "required_skills", "nice_to_have_skills"]
    return {"jd_id": jd_id, "updated": updated_cols, "fields": dict(zip(cols, updated_row)), "extracted": extracted}


@router.get("/{jd_id}/file")
async def get_jd_file(jd_id: str, request: Request, t: Optional[str] = Query(None)):
    """Stream the original uploaded JD source file. Accepts JWT in header OR ?t= query (for <a> tags)."""
    from fastapi.responses import FileResponse
    from pathlib import Path as _P
    if t and "authorization" not in {k.lower() for k in request.headers}:
        # Inject token from query into request scope so get_current_user picks it up
        from backend.core.jwt_auth import verify_token as _vjwt
        try:
            _vjwt(t)
        except Exception:
            raise HTTPException(status_code=401, detail="Invalid token")
    else:
        get_current_user(request)
    jd_id = _resolve_jd_id(jd_id)
    with get_cursor() as cur:
        cur.execute("SELECT source_file_path, source_file_name, source_file_type FROM jd_repository WHERE id=%s", (jd_id,))
        row = cur.fetchone()
    if not row:
        raise HTTPException(status_code=404, detail="JD not found")
    path, name, ftype = row
    if not path or not _P(path).exists():
        raise HTTPException(status_code=404, detail="No source file stored for this JD")
    media = {
        "pdf": "application/pdf",
        "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "doc": "application/msword",
        "txt": "text/plain",
        "md": "text/markdown",
    }.get((ftype or "").lower(), "application/octet-stream")
    return FileResponse(path, media_type=media, filename=name or _P(path).name)


@router.delete("/{jd_id}")
async def archive_jd(jd_id: str, request: Request, hard: bool = Query(False)):
    """Archive a JD (default), or hard-delete with ?hard=true (creator or superadmin only)."""
    user = get_current_user(request)
    jd_id = _resolve_jd_id(jd_id)
    with get_cursor() as cur:
        if hard:
            cur.execute("SELECT created_by FROM jd_repository WHERE id = %s", (jd_id,))
            row = cur.fetchone()
            if not row:
                raise HTTPException(404, "JD not found")
            creator_id = row[0]
            is_creator = creator_id == (user.get("user_id") or 0)
            is_super = user.get("role") == "superadmin"
            if not (is_creator or is_super):
                raise HTTPException(403, "Only the creator or superadmin can hard-delete this JD")
            cur.execute("DELETE FROM jd_repository WHERE id = %s", (jd_id,))
            return {"message": f"JD {jd_id} hard-deleted", "hard": True}
        # Soft archive — owner or admin+
        cur.execute("SELECT created_by FROM jd_repository WHERE id = %s", (jd_id,))
        _r = cur.fetchone()
        if not _r:
            raise HTTPException(404, "JD not found")
        ensure_owner_or_min(user, _r[0], "admin", request)
        cur.execute("UPDATE jd_repository SET status = 'archived', updated_at = NOW() WHERE id = %s", (jd_id,))
    return {"message": "JD archived"}


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------
async def _extract_requirements(jd_text: str) -> dict:
    """LLM-extract structured requirements from JD text."""
    prompt = f"""Extract ALL structured requirements from this job description. Be thorough — extract every skill, tool, and technology mentioned.

Return ONLY valid JSON:
{{
    "title": "concise role title (e.g. Group HSE Manager)",
    "department": "department or function",
    "seniority_level": "junior | mid | senior | lead | director | vp | c-level",
    "required_skills": ["Python", "Django", "PostgreSQL"],
    "nice_to_have": ["Redis", "Docker"],
    "min_experience_years": 5,
    "education_level": "bachelor",
    "industry_keywords": ["fintech", "saas"],
    "certifications": ["AWS Solutions Architect"]
}}

RULES:
- Title: strip prefixes like "Job Title:" or table syntax; just the role
- Extract ALL technical skills, tools, frameworks, and technologies mentioned
- Separate required vs nice-to-have based on context ("must have" vs "preferred"/"bonus")
- Parse experience years from text ("5+ years" → 5, "3-5 years" → 3)
- Detect education level (bachelor/master/phd or none if not mentioned)
- Extract industry/domain keywords

JD:
{jd_text[:8000]}"""

    result = llm_call(prompt, model=LITE_MODEL, temperature=0.1, max_tokens=1200)
    if not result:
        return {}

    try:
        match = re.search(r'\{[\s\S]*\}', result)
        if match:
            return json.loads(match.group())
    except json.JSONDecodeError:
        pass
    return {}


def _parse_compliance(enhanced_text: str) -> dict:
    """Parse DEI/legal/completeness scores from enhancement report."""
    result = {"dei_score": None, "legal_pass": None, "completeness": None}

    text_lower = enhanced_text.lower()

    # Parse DEI score
    import re
    dei_match = re.search(r'dei\s*(?:score)?[:\s]*(\d+)', text_lower)
    if dei_match:
        result["dei_score"] = int(dei_match.group(1))

    # Parse legal
    if "legal: pass" in text_lower or "legal:**pass" in text_lower.replace(" ", ""):
        result["legal_pass"] = True
    elif "legal: fail" in text_lower:
        result["legal_pass"] = False

    # Parse completeness
    comp_match = re.search(r'completeness[:\s]*(\d+)', text_lower)
    if comp_match:
        result["completeness"] = int(comp_match.group(1))

    return result


# ---------------------------------------------------------------------------
# Markdown parser — split jd_text into corporate template sections
# ---------------------------------------------------------------------------
def _parse_jd_sections(text: str) -> dict:
    """Split markdown jd_text into the 6 corporate sections + extract sub-fields."""
    if not text:
        return {}
    import re as _re
    # Split on `## ` headings (Markdown H2)
    parts = _re.split(r'(?m)^##\s+', text.strip())
    sections = {}
    intro = parts[0].strip() if parts else ""
    for part in parts[1:]:
        lines = part.split("\n", 1)
        head = lines[0].strip()
        body = lines[1].strip() if len(lines) > 1 else ""
        # Normalise key
        head_norm = head.lower()
        if "role information" in head_norm or head_norm.startswith("i."):
            sections["role_info"] = body
        elif "role scope" in head_norm or head_norm.startswith("ii."):
            sections["scope"] = body
        elif "responsibilities" in head_norm or head_norm.startswith("iii."):
            sections["responsibilities"] = body
        elif "requirements" in head_norm or head_norm.startswith("iv."):
            sections["requirements"] = body
        elif "working conditions" in head_norm or head_norm.startswith("v."):
            sections["conditions"] = body
        elif "governance" in head_norm or head_norm.startswith("vi."):
            sections["governance"] = body
        else:
            sections.setdefault("other", "")
            sections["other"] += f"\n## {head}\n{body}"
    sections["_intro"] = intro

    # Extract Job Purpose paragraph from scope section
    scope = sections.get("scope", "")
    m = _re.search(r'(?:\*\*Job Purpose:?\*\*|Job Purpose:?)\s*(.+?)(?=\n\n|\Z)', scope, _re.DOTALL | _re.IGNORECASE)
    sections["job_purpose"] = (m.group(1).strip() if m else scope).strip()

    # Extract bullets from responsibilities
    resp = sections.get("responsibilities", "")
    bullets = []
    for line in resp.splitlines():
        s = line.strip()
        if not s or not s.startswith(("*", "-", "•")):
            continue
        cleaned = s.lstrip("*-• \t").strip()
        if not cleaned:
            continue
        # Convert leading **Bold:** Text → Bold: Text
        cleaned = _re.sub(r'^\*\*([^*]+?)\*\*\s*:?\s*', r'\1: ', cleaned)
        # Strip any leftover bold markers
        cleaned = cleaned.replace("**", "")
        # Collapse double colons / trim
        cleaned = _re.sub(r':\s*:', ':', cleaned).strip(' :')
        if cleaned:
            bullets.append(cleaned)
    sections["responsibility_bullets"] = bullets

    # Extract sub-fields from requirements
    reqs = sections.get("requirements", "")
    def _grab(label):
        m = _re.search(rf'(?:\*\*{label}:?\*\*|{label}:)\s*(.+?)(?=\n\*\*|\n##|\Z)', reqs, _re.DOTALL | _re.IGNORECASE)
        return m.group(1).strip() if m else None
    sections["min_education"] = _grab("Minimum Education")
    sections["pref_education"] = _grab("Preferred Education")
    sections["total_experience"] = _grab("Total Experience")
    sections["technical_skills"] = _grab("Technical Skills")

    # Extract sub-fields from conditions
    cond = sections.get("conditions", "")
    def _grab_cond(label):
        m = _re.search(rf'(?:\*\*{label}:?\*\*|{label}:)\s*(.+?)(?=\n\*\*|\n##|\Z)', cond, _re.DOTALL | _re.IGNORECASE)
        return m.group(1).strip() if m else None
    sections["travel"] = _grab_cond("Travel Requirement")
    sections["physical"] = _grab_cond("Physical/Specific Conditions") or _grab_cond("Physical")

    return sections


def _strip_md(s: str) -> str:
    """Strip markdown decorations for plain-text cells."""
    if not s:
        return ""
    import re as _re
    s = _re.sub(r'\*\*(.+?)\*\*', r'\1', s)
    s = _re.sub(r'^[\*\-•]\s*', '', s, flags=_re.MULTILINE)
    s = _re.sub(r'\n{3,}', '\n\n', s)
    return s.strip()


# ---------------------------------------------------------------------------
# Export — corporate xlsx template
# ---------------------------------------------------------------------------
@router.get("/{jd_id}/export.xlsx")
async def export_jd_xlsx(jd_id: str, request: Request):
    """Export JD as xlsx matching the corporate template layout."""
    get_current_user(request)
    jd_id = _resolve_jd_id(jd_id)

    with get_cursor() as cur:
        cur.execute("SELECT * FROM jd_repository WHERE id = %s", (jd_id,))
        row = cur.fetchone()
        if not row:
            raise HTTPException(status_code=404, detail="JD not found")
        cols = [desc[0] for desc in cur.description]
        jd = dict(zip(cols, row))

    import io
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from fastapi.responses import Response

    wb = Workbook()
    ws = wb.active
    ws.title = "Proposed JD_Final"
    ws.column_dimensions["A"].width = 32
    ws.column_dimensions["B"].width = 80

    bold = Font(bold=True, size=11)
    header_font = Font(bold=True, color="FFFFFFFF", size=11)
    header_fill = PatternFill("solid", fgColor="FF383832")
    section_fill = PatternFill("solid", fgColor="FF00FC40")
    thin = Side(style="thin", color="FF383832")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)
    wrap = Alignment(wrap_text=True, vertical="top")

    def put_header(row, txt):
        ws.cell(row=row, column=1, value=txt).font = header_font
        ws.cell(row=row, column=1).fill = header_fill
        ws.cell(row=row, column=2, value="").fill = header_fill
        ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=2)

    def put_section(row, txt):
        c = ws.cell(row=row, column=1, value=txt)
        c.font = bold
        c.fill = section_fill
        ws.cell(row=row, column=2, value="").fill = section_fill
        ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=2)

    def put_row(row, label, value):
        ws.cell(row=row, column=1, value=label).font = bold
        ws.cell(row=row, column=1).border = border
        ws.cell(row=row, column=1).alignment = wrap
        ws.cell(row=row, column=2, value=str(value or "")).alignment = wrap
        ws.cell(row=row, column=2).border = border

    parsed = _parse_jd_sections(jd.get("jd_text") or "")

    r = 1
    put_header(r, f"JOB DESCRIPTION — {jd.get('title', '')}"); r += 2

    put_section(r, "I. Role Information"); r += 1
    for label, key in [
        ("Position Title",            "title"),
        ("Job Code/ID",               "job_code"),
        ("Business Sector/Entity",    "business_sector"),
        ("Department/Function",       "department"),
        ("Grading/Ranking",           "grading"),
        ("Reporting To",              "reporting_to"),
        ("Location",                  "location"),
        ("Work Mode",                 "work_mode"),
        ("Employment Type",           "employment_type"),
    ]:
        put_row(r, label, jd.get(key)); r += 1
    r += 1

    put_section(r, "II. Role Scope"); r += 1
    job_purpose = jd.get("job_purpose") or _strip_md(parsed.get("job_purpose", ""))
    put_row(r, "Job Purpose", job_purpose); r += 2

    put_section(r, "III. Key Responsibilities"); r += 1
    bullets = parsed.get("responsibility_bullets") or []
    if bullets:
        for i, b in enumerate(bullets, 1):
            put_row(r, f"{i}.", _strip_md(b)); r += 1
    else:
        put_row(r, "Responsibilities", _strip_md(parsed.get("responsibilities", "")) or jd.get("jd_text", ""))
        r += 1
    r += 1

    put_section(r, "IV. Job Requirements"); r += 1
    put_row(r, "Minimum Education",
            jd.get("education_level") or _strip_md(parsed.get("min_education") or "")); r += 1
    put_row(r, "Preferred Education",
            jd.get("preferred_education") or _strip_md(parsed.get("pref_education") or "")); r += 1
    exp_text = (
        _strip_md(parsed.get("total_experience") or "")
        or (f"{jd.get('min_experience_years')}+ years" if jd.get("min_experience_years") else "")
    )
    put_row(r, "Total Experience", exp_text); r += 1
    skills_text = (
        ", ".join(jd.get("required_skills") or [])
        or _strip_md(parsed.get("technical_skills") or "")
    )
    put_row(r, "Technical Skills", skills_text); r += 2

    put_section(r, "V. Working Conditions"); r += 1
    put_row(r, "Travel Requirement",
            jd.get("travel_requirement") or _strip_md(parsed.get("travel") or "")); r += 1
    put_row(r, "Physical/Specific Conditions",
            jd.get("physical_conditions") or _strip_md(parsed.get("physical") or "")); r += 2

    put_section(r, "VI. Governance"); r += 1
    put_row(r, "Equal Opportunity", jd.get("equal_opportunity")); r += 1
    doc_ctrl = f"Version {jd.get('doc_version') or '1.0'} · Owner: {jd.get('doc_owner') or '—'} · Last Reviewed: {jd.get('doc_last_review') or '—'}"
    put_row(r, "Document Control", doc_ctrl); r += 1

    buf = io.BytesIO()
    wb.save(buf); buf.seek(0)
    fname = f"JD_{jd.get('title', 'export').replace(' ', '_')}_{jd_id}.xlsx"
    return Response(
        buf.getvalue(),
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f'attachment; filename="{fname}"'},
    )


# ---------------------------------------------------------------------------
# Export — corporate Word (.docx)
# ---------------------------------------------------------------------------
@router.get("/{jd_id}/export.docx")
async def export_jd_docx(jd_id: str, request: Request):
    """Export JD as a Word document matching the corporate template."""
    get_current_user(request)
    jd_id = _resolve_jd_id(jd_id)

    with get_cursor() as cur:
        cur.execute("SELECT * FROM jd_repository WHERE id = %s", (jd_id,))
        row = cur.fetchone()
        if not row:
            raise HTTPException(status_code=404, detail="JD not found")
        cols = [desc[0] for desc in cur.description]
        jd = dict(zip(cols, row))

    import io
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.table import WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from fastapi.responses import Response

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(1.5); section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1.8); section.right_margin = Cm(1.8)

    # Default style — Space Grotesk (matches frontend), fallback Calibri
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Brutalist palette (matching Excel + frontend)
    INK = RGBColor(0x38, 0x38, 0x32)
    SURFACE = RGBColor(0xFE, 0xFF, 0xD6)
    PRIMARY = RGBColor(0x00, 0x75, 0x18)
    GREEN = "00FC40"
    DARK = "383832"
    TAN = "F0F0E0"

    def shade_cell(cell, hex_color):
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), hex_color)
        tcPr.append(shd)

    def set_cell_borders(cell, color=DARK, sz=8):
        tc_pr = cell._tc.get_or_add_tcPr()
        existing = tc_pr.find(qn("w:tcBorders"))
        if existing is not None:
            tc_pr.remove(existing)
        borders = OxmlElement("w:tcBorders")
        for edge in ("top", "left", "bottom", "right"):
            b = OxmlElement(f"w:{edge}")
            b.set(qn("w:val"), "single")
            b.set(qn("w:sz"), str(sz))
            b.set(qn("w:color"), color)
            borders.append(b)
        tc_pr.append(borders)

    # ── TITLE BAR (dark fill, white text — mirrors Excel header) ──
    t_title = doc.add_table(rows=1, cols=1)
    title_cell = t_title.rows[0].cells[0]
    title_cell.text = ""
    p = title_cell.paragraphs[0]
    run = p.add_run(f"JOB DESCRIPTION — {jd.get('title','').upper()}")
    run.bold = True; run.font.size = Pt(16)
    run.font.color.rgb = SURFACE
    run.font.name = "Calibri"
    shade_cell(title_cell, DARK)
    set_cell_borders(title_cell, color=DARK, sz=16)
    title_cell.paragraphs[0].paragraph_format.space_before = Pt(4)
    title_cell.paragraphs[0].paragraph_format.space_after = Pt(4)
    doc.add_paragraph()

    def section_heading(text):
        # green fill bar like Excel
        t = doc.add_table(rows=1, cols=1)
        cell = t.rows[0].cells[0]
        cell.text = ""
        para = cell.paragraphs[0]
        para.paragraph_format.space_before = Pt(2)
        para.paragraph_format.space_after = Pt(2)
        r = para.add_run(text.upper())
        r.bold = True; r.font.size = Pt(13)
        r.font.color.rgb = INK
        shade_cell(cell, GREEN)
        set_cell_borders(cell, color=DARK, sz=12)

    def kv_table(rows):
        t = doc.add_table(rows=len(rows), cols=2)
        t.autofit = False
        # column widths
        for row in t.rows:
            row.cells[0].width = Cm(5.5)
            row.cells[1].width = Cm(11.5)
        for i, (k, v) in enumerate(rows):
            c0, c1 = t.rows[i].cells
            c0.text = ""; c1.text = ""
            r0 = c0.paragraphs[0].add_run(k)
            r0.bold = True; r0.font.size = Pt(11); r0.font.color.rgb = INK
            r1 = c1.paragraphs[0].add_run(str(v) if v not in (None, "") else "—")
            r1.font.size = Pt(11)
            shade_cell(c0, TAN)
            set_cell_borders(c0, color=DARK)
            set_cell_borders(c1, color=DARK)
            for c in (c0, c1):
                c.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        return t

    # I. Role Information
    section_heading("I. ROLE INFORMATION")
    kv_table([
        ("Position Title",         jd.get("title")),
        ("Job Code/ID",            jd.get("job_code")),
        ("Business Sector/Entity", jd.get("business_sector")),
        ("Department/Function",    jd.get("department")),
        ("Grading/Ranking",        jd.get("grading") or jd.get("seniority_level")),
        ("Reporting To",           jd.get("reporting_to")),
        ("Location",               jd.get("location")),
        ("Work Mode",              jd.get("work_mode")),
        ("Employment Type",        jd.get("employment_type")),
    ])

    parsed = _parse_jd_sections(jd.get("jd_text") or "")

    # II. Role Scope
    section_heading("II. ROLE SCOPE")
    p = doc.add_paragraph()
    p.add_run("Job Purpose: ").bold = True
    p.add_run(jd.get("job_purpose") or _strip_md(parsed.get("job_purpose", "")) or "—")

    # III. Key Responsibilities
    section_heading("III. KEY RESPONSIBILITIES")
    bullets = parsed.get("responsibility_bullets") or []
    if bullets:
        for b in bullets:
            para = doc.add_paragraph(style="List Bullet")
            para.add_run(_strip_md(b))
    else:
        # fallback — try plain non-markdown jd_text or responsibilities section
        plain = _strip_md(parsed.get("responsibilities", "") or jd.get("jd_text", ""))
        for line in plain.splitlines():
            line = line.strip()
            if line:
                doc.add_paragraph(line)

    # IV. Job Requirements
    section_heading("IV. JOB REQUIREMENTS")
    exp_text = (
        _strip_md(parsed.get("total_experience") or "")
        or (f"{jd.get('min_experience_years')}+ years" if jd.get("min_experience_years") else "—")
    )
    skills_text = (
        ", ".join(jd.get("required_skills") or [])
        or _strip_md(parsed.get("technical_skills") or "")
    )
    kv_table([
        ("Minimum Education",   jd.get("education_level") or _strip_md(parsed.get("min_education") or "")),
        ("Preferred Education", jd.get("preferred_education") or _strip_md(parsed.get("pref_education") or "")),
        ("Total Experience",    exp_text),
        ("Technical Skills",    skills_text),
    ])

    # V. Working Conditions
    section_heading("V. WORKING CONDITIONS")
    kv_table([
        ("Travel Requirement",        jd.get("travel_requirement") or _strip_md(parsed.get("travel") or "")),
        ("Physical/Specific Conditions", jd.get("physical_conditions") or _strip_md(parsed.get("physical") or "")),
    ])

    # VI. Governance
    section_heading("VI. GOVERNANCE")
    p = doc.add_paragraph()
    p.add_run("Equal Opportunity: ").bold = True
    p.add_run(jd.get("equal_opportunity") or "")
    p2 = doc.add_paragraph()
    p2.add_run("Document Control: ").bold = True
    p2.add_run(
        f"Version {jd.get('doc_version') or '1.0'}  ·  "
        f"Owner: {jd.get('doc_owner') or '—'}  ·  "
        f"Last Reviewed: {jd.get('doc_last_review') or '—'}"
    )

    buf = io.BytesIO()
    doc.save(buf); buf.seek(0)
    fname = f"JD_{(jd.get('title') or 'export').replace(' ', '_')}_{jd_id}.docx"
    return Response(
        buf.getvalue(),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{fname}"'},
    )


# ---------------------------------------------------------------------------
# Share / Visibility
# ---------------------------------------------------------------------------
class JDShareRequest(BaseModel):
    visibility: Optional[str] = None  # legacy: private | sector | global
    shared_sector: Optional[bool] = None
    shared_global: Optional[bool] = None


@router.post("/{jd_id}/share")
async def share_jd(jd_id: str, body: JDShareRequest, request: Request):
    user = _user_with_sector(get_current_user(request))
    jd_id = _resolve_jd_id(jd_id)
    # Resolve final flags from either new (shared_*) or legacy (visibility) payload
    if body.shared_sector is None and body.shared_global is None:
        # Legacy single-enum input
        v = (body.visibility or "private").lower()
        if v not in ("private", "sector", "global"):
            raise HTTPException(400, "visibility must be private | sector | global")
        shared_sector = (v == "sector")
        shared_global = (v == "global")
    else:
        shared_sector = bool(body.shared_sector)
        shared_global = bool(body.shared_global)

    with get_cursor() as cur:
        cur.execute("SELECT created_by, sector_id FROM jd_repository WHERE id = %s", (jd_id,))
        row = cur.fetchone()
        if not row:
            raise HTTPException(404, "JD not found")
        creator_id, jd_sector = row
        is_creator = creator_id == (user.get("user_id") or 0)
        is_admin   = user.get("role") in ("admin", "group_hr", "superadmin")

        if shared_global and not is_admin:
            raise HTTPException(403, "Only admin / group_hr can publish globally")
        if not (is_creator or is_admin):
            raise HTTPException(403, "Only the creator or admin can change visibility")

        # visibility legacy column reflects the highest level for back-compat
        if shared_global: legacy = "global"
        elif shared_sector: legacy = "sector"
        else: legacy = "private"

        new_sector = jd_sector or user.get("sector_id")
        cur.execute("""
            UPDATE jd_repository
               SET visibility = %s,
                   shared_sector = %s,
                   shared_global = %s,
                   sector_id = %s,
                   published_by = %s,
                   published_at = NOW(),
                   updated_at = NOW()
             WHERE id = %s
        """, (legacy, shared_sector, shared_global, new_sector, user.get("user_id"), jd_id))
        cur.execute(
            "INSERT INTO audit_log (user_id, action, resource_type, resource_id, details) VALUES (%s,%s,%s,%s,%s)",
            (user.get("user_id"), "JD_SHARE", "jd_repository", jd_id,
             json.dumps({"shared_sector": shared_sector, "shared_global": shared_global})),
        )
    return {"id": jd_id, "shared_sector": shared_sector, "shared_global": shared_global, "visibility": legacy}


# ---------------------------------------------------------------------------
# Scoring Weights
# ---------------------------------------------------------------------------
SCORING_PRESETS = {
    "engineering": {"weight_skills":40, "weight_experience":25, "weight_industry":15, "weight_education":10, "weight_certifications":10, "weight_culture":0},
    "sales":       {"weight_skills":25, "weight_experience":30, "weight_industry":20, "weight_education":5,  "weight_certifications":5,  "weight_culture":15},
    "design":      {"weight_skills":45, "weight_experience":20, "weight_industry":10, "weight_education":5,  "weight_certifications":0,  "weight_culture":20},
    "product":     {"weight_skills":30, "weight_experience":30, "weight_industry":20, "weight_education":10, "weight_certifications":0,  "weight_culture":10},
    "marketing":   {"weight_skills":30, "weight_experience":25, "weight_industry":25, "weight_education":10, "weight_certifications":5,  "weight_culture":5},
    "compliance":  {"weight_skills":20, "weight_experience":25, "weight_industry":15, "weight_education":10, "weight_certifications":25, "weight_culture":5},
    "healthcare":  {"weight_skills":25, "weight_experience":20, "weight_industry":15, "weight_education":15, "weight_certifications":20, "weight_culture":5},
    "finance":     {"weight_skills":25, "weight_experience":25, "weight_industry":20, "weight_education":10, "weight_certifications":15, "weight_culture":5},
}


def _normalize_weights(w: dict) -> dict:
    keys = ("weight_skills","weight_experience","weight_industry","weight_education","weight_certifications","weight_culture")
    total = sum(float(w.get(k, 0) or 0) for k in keys)
    if total <= 0:
        return {k: 0 for k in keys}
    return {k: round(float(w.get(k, 0) or 0) * 100.0 / total, 2) for k in keys}


class JDWeightsPatch(BaseModel):
    weight_skills: Optional[float] = None
    weight_experience: Optional[float] = None
    weight_industry: Optional[float] = None
    weight_education: Optional[float] = None
    weight_certifications: Optional[float] = None
    weight_culture: Optional[float] = None
    knockout_threshold: Optional[float] = None
    scoring_profile: Optional[str] = None
    normalize: bool = True
    apply_to_open_positions: bool = False


@router.patch("/{jd_id}/weights")
async def patch_jd_weights(jd_id: str, body: JDWeightsPatch, request: Request):
    user = get_current_user(request)
    jd_id = _resolve_jd_id(jd_id)
    with get_cursor() as _c:
        _c.execute("SELECT created_by FROM jd_repository WHERE id = %s", (jd_id,))
        _r = _c.fetchone()
    if _r is not None:
        ensure_owner_or_min(user, _r[0], "admin", request)
    fields = body.model_dump(exclude_none=True)
    apply_open = fields.pop("apply_to_open_positions", False)
    normalize = fields.pop("normalize", True)
    if not fields:
        raise HTTPException(400, "No weight fields provided")

    with get_cursor() as cur:
        cur.execute("SELECT * FROM jd_repository WHERE id = %s", (jd_id,))
        row = cur.fetchone()
        if not row:
            raise HTTPException(404, "JD not found")
        cols = [d[0] for d in cur.description]
        jd = dict(zip(cols, row))

        # Build merged weights and normalize
        merged = {k: jd.get(k) for k in ("weight_skills","weight_experience","weight_industry","weight_education","weight_certifications","weight_culture")}
        for k, v in fields.items():
            if k in merged:
                merged[k] = v
        if normalize:
            merged = _normalize_weights(merged)

        sets, params = [], []
        for k, v in merged.items():
            sets.append(f"{k} = %s"); params.append(v)
        for k in ("knockout_threshold", "scoring_profile"):
            if k in fields:
                sets.append(f"{k} = %s"); params.append(fields[k])
        sets.append("updated_at = NOW()")
        params.append(jd_id)
        cur.execute(f"UPDATE jd_repository SET {', '.join(sets)} WHERE id = %s", params)

        if apply_open:
            # Sync to all positions that haven't overridden + use this JD
            cur.execute("""
                UPDATE positions p SET
                    weight_skills = %s, weight_experience = %s, weight_industry = %s,
                    weight_education = %s, weight_certifications = %s, weight_culture = %s,
                    knockout_threshold = %s,
                    updated_at = NOW()
                WHERE p.weights_source_jd_id = %s
                  AND COALESCE(p.weights_overridden, FALSE) = FALSE
                  AND p.status = 'active'
            """, (
                merged["weight_skills"], merged["weight_experience"], merged["weight_industry"],
                merged["weight_education"], merged["weight_certifications"], merged["weight_culture"],
                fields.get("knockout_threshold", jd.get("knockout_threshold") or 0),
                jd_id,
            ))

        cur.execute(
            "INSERT INTO audit_log (user_id, action, resource_type, resource_id, details) VALUES (%s,%s,%s,%s,%s)",
            (user.get("user_id"), "JD_WEIGHTS_UPDATE", "jd_repository", jd_id, json.dumps(merged)),
        )
    try:
        from backend.agents.sync import enqueue_dedup
        enqueue_dedup("auto_rescore_jd_chain", {"jd_id": jd_id})
    except Exception as e:
        logger.warning(f"enqueue auto_rescore_jd_chain failed: {e}")
    return {"id": jd_id, **merged, "knockout_threshold": fields.get("knockout_threshold"), "scoring_profile": fields.get("scoring_profile")}


class JDPresetRequest(BaseModel):
    profile: str
    apply_to_open_positions: bool = False


@router.post("/{jd_id}/preset")
async def apply_jd_preset(jd_id: str, body: JDPresetRequest, request: Request):
    user = get_current_user(request)
    jd_id = _resolve_jd_id(jd_id)
    if body.profile not in SCORING_PRESETS:
        raise HTTPException(400, f"Unknown profile. Choose from: {', '.join(SCORING_PRESETS)}")
    preset = SCORING_PRESETS[body.profile]
    patch = JDWeightsPatch(**preset, scoring_profile=body.profile, apply_to_open_positions=body.apply_to_open_positions, normalize=False)
    return await patch_jd_weights(jd_id, patch, request)


@router.get("/meta/presets")
async def list_presets(request: Request):
    get_current_user(request)
    return {"presets": SCORING_PRESETS}


# ---------------------------------------------------------------------------
# JD weight LOCK
# ---------------------------------------------------------------------------
class JDLockRequest(BaseModel):
    weights_locked: Optional[bool] = None
    weights_locked_dims: Optional[list[str]] = None


@router.patch("/{jd_id}/lock")
async def patch_jd_lock(jd_id: str, body: JDLockRequest, request: Request):
    user = get_current_user(request)
    jd_id = _resolve_jd_id(jd_id)
    with get_cursor() as _c:
        _c.execute("SELECT created_by FROM jd_repository WHERE id = %s", (jd_id,))
        _r = _c.fetchone()
    if _r is not None:
        ensure_owner_or_min(user, _r[0], "admin", request)
    sets, params = [], []
    if body.weights_locked is not None:
        sets.append("weights_locked = %s"); params.append(body.weights_locked)
    if body.weights_locked_dims is not None:
        sets.append("weights_locked_dims = %s::jsonb"); params.append(json.dumps(body.weights_locked_dims))
    if not sets:
        raise HTTPException(400, "No fields")
    params.append(jd_id)
    with get_cursor() as cur:
        cur.execute(f"UPDATE jd_repository SET {', '.join(sets)} WHERE id = %s RETURNING id", params)
        if not cur.fetchone():
            raise HTTPException(404, "JD not found")
        cur.execute(
            "INSERT INTO audit_log (user_id, action, resource_type, resource_id, details) VALUES (%s,%s,%s,%s,%s)",
            (user.get("user_id"), "JD_LOCK_UPDATE", "jd_repository", jd_id, json.dumps(body.model_dump(exclude_none=True))),
        )
    return {"id": jd_id, **body.model_dump(exclude_none=True)}


# ---------------------------------------------------------------------------
# Competencies for JD
# ---------------------------------------------------------------------------
class JDCompetencyEntry(BaseModel):
    competency_id: int
    required_level: int
    importance: Optional[str] = "required"
    weight: Optional[float] = 1.0
    notes: Optional[str] = None


@router.get("/{jd_id}/competencies")
async def list_jd_competencies(jd_id: str, request: Request):
    get_current_user(request)
    jd_id = _resolve_jd_id(jd_id)
    with get_cursor() as cur:
        cur.execute("""
            SELECT jc.competency_id, c.key, c.label, c.category,
                   jc.required_level, jc.importance, jc.weight, jc.notes
            FROM jd_competencies jc
            JOIN competencies c ON c.id = jc.competency_id
            WHERE jc.jd_id = %s
            ORDER BY jc.importance, c.label
        """, (jd_id,))
        cols = [d[0] for d in cur.description]
        rows = [dict(zip(cols, r)) for r in cur.fetchall()]
    return {"competencies": rows, "total": len(rows)}


@router.put("/{jd_id}/competencies")
async def replace_jd_competencies(jd_id: str, request: Request):
    user = get_current_user(request)
    jd_id = _resolve_jd_id(jd_id)
    with get_cursor() as _c:
        _c.execute("SELECT created_by FROM jd_repository WHERE id = %s", (jd_id,))
        _r = _c.fetchone()
    if _r is not None:
        ensure_owner_or_min(user, _r[0], "admin", request)
    body = await request.json()
    if not isinstance(body, list):
        raise HTTPException(400, "Expected JSON array")
    with get_cursor() as cur:
        cur.execute("SELECT id FROM jd_repository WHERE id = %s", (jd_id,))
        if not cur.fetchone():
            raise HTTPException(404, "JD not found")
        cur.execute("DELETE FROM jd_competencies WHERE jd_id = %s", (jd_id,))
        for row in body:
            cid = row.get("competency_id")
            lvl = row.get("required_level")
            if cid is None or lvl is None:
                continue
            try:
                lvl = int(lvl)
            except Exception:
                continue
            lvl = max(1, min(5, lvl))
            cur.execute("""
                INSERT INTO jd_competencies (jd_id, competency_id, required_level, importance, weight, notes)
                VALUES (%s, %s, %s, %s, %s, %s)
                ON CONFLICT (jd_id, competency_id) DO UPDATE SET
                    required_level = EXCLUDED.required_level,
                    importance = EXCLUDED.importance,
                    weight = EXCLUDED.weight,
                    notes = EXCLUDED.notes
            """, (jd_id, cid, lvl, row.get("importance") or "required",
                  row.get("weight") or 1.0, row.get("notes")))
        cur.execute(
            "INSERT INTO audit_log (user_id, action, resource_type, resource_id, details) VALUES (%s,%s,%s,%s,%s)",
            (user.get("user_id"), "JD_COMP_REPLACE", "jd_repository", jd_id, json.dumps({"count": len(body)})),
        )
    return {"jd_id": jd_id, "count": len(body), "message": "saved"}


def _internal_extract_jd_competencies(jd_id: int) -> dict:
    """Defensive helper — auto-extract JD competencies from jd_text via LLM."""
    with get_cursor() as cur:
        cur.execute("SELECT jd_text FROM jd_repository WHERE id = %s", (jd_id,))
        row = cur.fetchone()
        if not row:
            return {"extracted": [], "created": 0, "updated": 0, "error": "JD not found"}
        jd_text = row[0] or ""
        cur.execute("""
            SELECT c.id, c.key, c.label, c.description
            FROM competencies c
            JOIN competency_frameworks f ON f.id = c.framework_id
            WHERE c.active = TRUE AND f.is_default = TRUE
        """)
        lib = [{"id": r[0], "key": r[1], "label": r[2], "description": r[3]} for r in cur.fetchall()]

    if not jd_text.strip() or not lib:
        return {"extracted": [], "created": 0, "updated": 0}

    key_to_id = {c["key"]: c["id"] for c in lib}
    library_brief = [{"key": c["key"], "label": c["label"], "description": (c["description"] or "")[:120]} for c in lib]

    prompt = f"""You are tagging a JD with competencies from this exact library:
{json.dumps(library_brief)}

Decide which the JD requires + level (1-5) + importance (required|preferred|nice) + brief evidence.
Skip noise. Return ONLY a JSON array.
[{{"key":"...","required_level":4,"importance":"required","evidence":"..."}}]

JD text:
{jd_text[:8000]}
"""
    extracted: list = []
    try:
        result = llm_call(prompt, model=LITE_MODEL, temperature=0.1, max_tokens=1500)
        if result:
            m = re.search(r'\[[\s\S]*\]', result)
            if m:
                extracted = json.loads(m.group()) or []
    except Exception as e:
        logger.warning(f"JD comp extract LLM failed for jd {jd_id}: {e}")
        extracted = []

    created = 0
    updated = 0
    with get_cursor() as cur:
        for item in extracted:
            if not isinstance(item, dict):
                continue
            key = item.get("key")
            cid = key_to_id.get(key)
            if not cid:
                continue
            try:
                lvl = int(item.get("required_level") or 3)
            except Exception:
                lvl = 3
            lvl = max(1, min(5, lvl))
            importance = item.get("importance") or "required"
            if importance not in ("required", "preferred", "nice"):
                importance = "required"
            evidence = (item.get("evidence") or "")[:400]
            cur.execute("""
                INSERT INTO jd_competencies (jd_id, competency_id, required_level, importance, weight, notes)
                VALUES (%s, %s, %s, %s, 1.0, %s)
                ON CONFLICT (jd_id, competency_id) DO UPDATE SET
                    required_level = EXCLUDED.required_level,
                    importance = EXCLUDED.importance,
                    notes = EXCLUDED.notes
                RETURNING (xmax = 0) AS inserted
            """, (jd_id, cid, lvl, importance, evidence))
            r = cur.fetchone()
            if r and r[0]:
                created += 1
            else:
                updated += 1
    return {"extracted": extracted, "created": created, "updated": updated}


@router.post("/{jd_id}/competencies/auto-extract", dependencies=[Depends(require_role("recruiter"))])
async def auto_extract_jd_competencies(jd_id: str, request: Request):
    get_current_user(request)
    jd_id = _resolve_jd_id(jd_id)
    result = _internal_extract_jd_competencies(jd_id)
    return {"jd_id": jd_id, **result}


def _sync_position_competencies_from_jd(jd_id: int, position_id: int) -> int:
    """Idempotent: copy JD competencies → position_competencies (source='jd')."""
    with get_cursor() as cur:
        cur.execute("""
            INSERT INTO position_competencies (position_id, competency_id, required_level, importance, weight, source, notes)
            SELECT %s, jc.competency_id, jc.required_level, jc.importance, jc.weight, 'jd', jc.notes
            FROM jd_competencies jc
            WHERE jc.jd_id = %s
            ON CONFLICT (position_id, competency_id) DO UPDATE SET
                required_level = EXCLUDED.required_level,
                importance = EXCLUDED.importance,
                weight = EXCLUDED.weight,
                source = CASE WHEN position_competencies.source = 'override'
                              THEN position_competencies.source ELSE 'jd' END
            RETURNING competency_id
        """, (position_id, jd_id))
        return len(cur.fetchall())


# ============== JD EMBEDDINGS (mig 040) ==============

@router.post("/embeddings/backfill")
async def backfill_jd_embeddings(request: Request):
    """Embed any JD missing jd_embedding. Admin/superadmin only."""
    user = get_current_user(request)
    if user.get("role") not in ("admin", "superadmin"):
        raise HTTPException(status_code=403, detail="admin only")
    import asyncio as _aio
    from backend.agents.ingest import _embed_jd, _safe
    with get_cursor() as cur:
        cur.execute("SELECT id, title, jd_text FROM jd_repository WHERE jd_embedding IS NULL AND status=%s", ("active",))
        rows = cur.fetchall()
    queued = 0
    for r in rows:
        _aio.create_task(_safe(_embed_jd(r[0], r[1] or "", r[2] or ""), f"jd-backfill {r[0]}"))
        queued += 1
    return {"queued": queued, "message": f"{queued} JDs queued for embedding"}


@router.get("/search/semantic")
async def jd_semantic_search(q: str, limit: int = 10, request: Request = None):
    """Cosine-similarity search across JD bodies."""
    get_current_user(request)
    if not q.strip():
        return {"results": []}
    from backend.core.config import embed_text
    vec = await embed_text(q.strip()[:2000])
    if not vec:
        raise HTTPException(status_code=503, detail="embedding unavailable")
    with get_cursor() as cur:
        cur.execute("""
            SELECT id, public_id, title, department, seniority_level,
                   1 - (jd_embedding <=> %s::vector) AS similarity
            FROM jd_repository
            WHERE jd_embedding IS NOT NULL AND status = %s
            ORDER BY jd_embedding <=> %s::vector
            LIMIT %s
        """, (str(vec), "active", str(vec), max(1, min(50, limit))))
        cols = [d[0] for d in cur.description]
        results = [dict(zip(cols, r)) for r in cur.fetchall()]
    return {"query": q, "results": results, "count": len(results)}


# ---------------------------------------------------------------------------
# CSV export
# ---------------------------------------------------------------------------
@router.get("/export.csv")
async def export_jds_csv(request: Request):
    """Export JD repository as CSV."""
    get_current_user(request)
    buf = _io.StringIO()
    w = _csv.writer(buf)
    w.writerow(["public_id", "title", "department", "seniority", "visibility",
                "status", "created_at", "owner_id"])
    with get_cursor() as cur:
        cur.execute("""
            SELECT public_id, title, department, seniority_level, visibility,
                   status, created_at, created_by
            FROM jd_repository
            ORDER BY created_at DESC
        """)
        for r in cur.fetchall():
            w.writerow([("" if v is None else v) for v in r])
    return Response(
        content=buf.getvalue(),
        media_type="text/csv",
        headers={"Content-Disposition": 'attachment; filename="jds.csv"'},
    )
